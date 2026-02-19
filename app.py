import bcrypt
import os
from dotenv import load_dotenv

load_dotenv()
import certifi
import pandas as pd
import secrets
import string
import sys
import smtplib
import re
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime, timedelta, timezone
import pytz
from flask import Flask, request, jsonify, send_file, send_from_directory, session
from flask_cors import CORS
from pymongo import MongoClient
from gridfs import GridFS
from bson import ObjectId
from io import BytesIO
from functools import wraps
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import _Cell

# Fix Unicode encoding for Windows console
if sys.platform == 'win32':
    try:
        if hasattr(sys.stdout, 'reconfigure'):
            sys.stdout.reconfigure(encoding='utf-8', errors='replace')
            sys.stderr.reconfigure(encoding='utf-8', errors='replace')
        else:
            import codecs
            if hasattr(sys.stdout, 'buffer'):
                sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
                sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')
    except (AttributeError, ValueError):
        pass

# --- App Initialization ---
app = Flask(__name__)
SECRET_KEY = os.getenv('SECRET_KEY')
if not SECRET_KEY:
    raise RuntimeError("❌ SECRET_KEY not set in environment variables. Application cannot start.")
app.secret_key = SECRET_KEY
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Session configuration
app.config['SESSION_COOKIE_SAMESITE'] = 'None'
app.config['SESSION_COOKIE_SECURE'] = True

# CORS and TLS configuration
os.environ.setdefault('SSL_CERT_FILE', certifi.where())
CORS(app, resources={r"/api/*": {
    "origins": [
        r"^https?://localhost(:\d+)?$",
        r"^https?://127\.0\.0\.1(:\d+)?$",
        r"^https://.*\.netlify\.app$",
        r"^https://sps-ksou\.onrender\.com$",
        r"^null$",
        r"^file://.*$"
    ],
    "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    "allow_headers": ["Content-Type", "Authorization"],
    "supports_credentials": True
}}) 
def _cache_headers(response):
    ct = response.headers.get('Content-Type', '')
    if 'text/html' in ct:
        response.headers['Cache-Control'] = 'no-store'
    return response
app.after_request(_cache_headers)

@app.route('/', methods=['GET'])
def root():
    return send_file(os.path.join(BASE_DIR, 'index.html'))

@app.route('/dashboard', methods=['GET'])
def serve_dashboard():
    return send_file(os.path.join(BASE_DIR, 'dashboard.html'))

@app.route('/studentprogression', methods=['GET'])
def serve_studentprogression():
    return send_file(os.path.join(BASE_DIR, 'studentprogression.html'))

@app.route('/reports', methods=['GET'])
def serve_reports():
    return send_file(os.path.join(BASE_DIR, 'reports.html'))

@app.route('/settings', methods=['GET'])
def serve_settings():
    return send_file(os.path.join(BASE_DIR, 'settings.html'))

@app.route('/<path:filepath>', methods=['GET'])
def serve_static(filepath):
    full_path = os.path.join(BASE_DIR, filepath)
    if os.path.isfile(full_path):
        if filepath.endswith('.html'):
            return send_file(full_path)
        return send_from_directory(BASE_DIR, filepath)
    return jsonify({"message": "Not Found"}), 404

@app.errorhandler(404)
def handle_404(e):
    if request.path.startswith('/api/'):
        return jsonify({"message": "Not Found"}), 404
    return send_file(os.path.join(BASE_DIR, 'index.html'))

@app.route('/healthz', methods=['GET'])
def healthz():
    return "ok", 200

# --- Configuration ---
MONGO_URI = os.getenv('MONGO_URI')
DB_NAME = "StudentProgressionDB"

# --- Database Connection ---
try:
    client = MongoClient(MONGO_URI)
    db = client[DB_NAME]
    client.admin.command('ismaster')
    print("✅ Successfully connected to MongoDB Atlas!")
    
    # Migration: Standardize Extracurricular Schema
    def migrate_extracurricular_schema():
        from pymongo import UpdateOne
        
        try:
            collection = db.extracurricular_records
            updates = []
            
            for record in collection.find({}):
                update_fields = {}
                unset_fields = {}
                
                # Normalize field names
                if "activity_name" in record and not record.get("event_name"):
                    update_fields["event_name"] = record.get("activity_name")
                    unset_fields["activity_name"] = ""
                    
                if "organization" in record and not record.get("organizer"):
                    update_fields["organizer"] = record.get("organization")
                    unset_fields["organization"] = ""
                    
                if "organizer_name" in record and not record.get("organizer"):
                    update_fields["organizer"] = record.get("organizer_name")
                    unset_fields["organizer_name"] = ""
                    
                if "participants" in record and not record.get("number_of_participants"):
                    update_fields["number_of_participants"] = record.get("participants")
                    unset_fields["participants"] = ""
                    
                if "number_of_students" in record and not record.get("number_of_participants"):
                    update_fields["number_of_participants"] = record.get("number_of_students")
                    unset_fields["number_of_students"] = ""
                    
                if update_fields or unset_fields:
                    updates.append(
                        UpdateOne(
                            {"_id": record["_id"]},
                            {"$set": update_fields, "$unset": unset_fields}
                        )
                    )
            
            if updates:
                collection.bulk_write(updates)
                
        except Exception as e:
            print(f"❌ Error during schema migration: {e}")

    # Run schema migration
    migrate_extracurricular_schema()

    # --- SaaS Initialization ---
    # Ensure default college exists
    if 'colleges' not in db.list_collection_names():
        db.create_collection('colleges')
        # print("✅ Created 'colleges' collection.")

    if not db.colleges.find_one({"name": "Gharda Institute of Technology"}):
        db.colleges.insert_one({
            "name": "Gharda Institute of Technology",
            "status": "active",
            "subscription_active": True,
            "subscription_expiry": None,
            "max_users_per_branch": 2,
            "created_at": datetime.now(timezone.utc)
        })
        # print("✅ Seeded default college: Gharda Institute of Technology")

    # Ensure Super Admin exists (Safe Production Implementation)
    admin_email = os.getenv('ADMIN_EMAIL')
    admin_password = os.getenv('ADMIN_PASSWORD')
    
    if not admin_email or not admin_password:
        print("⚠ ADMIN credentials not set in environment variables. Admin user creation skipped.")
    else:
        # Step A: Check if ANY super_admin exists
        existing_super_admin = db.users.find_one({"role": "super_admin"})

        if existing_super_admin:
            # Step B: Super admin exists
            if existing_super_admin['email'] == admin_email:
                # Update password if needed
                if not bcrypt.checkpw(admin_password.encode('utf-8'), existing_super_admin['password']):
                    hashed_pw = bcrypt.hashpw(admin_password.encode('utf-8'), bcrypt.gensalt())
                    db.users.update_one(
                        {"_id": existing_super_admin['_id']},
                        {"$set": {"password": hashed_pw}}
                    )
                    print(f"✅ Admin password updated for: {admin_email}")
            else:
                # Email does NOT match. Reassign super admin to environment email
                hashed_pw = bcrypt.hashpw(admin_password.encode('utf-8'), bcrypt.gensalt())
                db.users.update_one(
                    {"_id": existing_super_admin['_id']},
                    {"$set": {
                        "email": admin_email,
                        "password": hashed_pw,
                        "full_name": "Super Admin" # Ensure name is consistent
                    }}
                )
                print("🔁 Super admin reassigned to environment ADMIN_EMAIL")
        else:
            # Step C: NO super_admin exists -> Create new
            hashed_pw = bcrypt.hashpw(admin_password.encode('utf-8'), bcrypt.gensalt())
            db.users.insert_one({
                "email": admin_email,
                "password": hashed_pw,
                "role": "super_admin",
                "college": None,
                "branch": None,
                "full_name": "Super Admin",
                "email_verified": True,
                "created_at": datetime.now(timezone.utc)
            })
            print(f"✅ Created Super Admin: {admin_email}")

except Exception as e:
    print(f"❌ DATABASE ERROR: Could not connect to MongoDB. Full error: {e}")
    sys.exit(1)

# --- SMTP Configuration ---
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587
SMTP_EMAIL = os.getenv('SMTP_EMAIL')
SMTP_PASSWORD = os.getenv('SMTP_PASSWORD')

# --- Utility Functions ---
def normalize_name(name):
    """Normalize student name: remove leading special chars, trim, upper case, collapse spaces."""
    if not isinstance(name, str):
        if pd.isna(name):
            return ""
        name = str(name)
    
    # Remove leading special characters (/, -, .) and spaces
    name = re.sub(r'^[/\-.\s]+', '', name)
    
    # Standardize format
    name = name.strip().upper()
    name = re.sub(r'\s+', ' ', name)
    return name

def convert_objectid_to_str(obj):
    """Recursively convert ObjectId instances to strings for JSON serialization."""
    if isinstance(obj, ObjectId):
        return str(obj)
    elif isinstance(obj, dict):
        return {key: convert_objectid_to_str(value) for key, value in obj.items()}
    elif isinstance(obj, list):
        return [convert_objectid_to_str(item) for item in obj]
    elif isinstance(obj, datetime):
        return obj.isoformat()
    else:
        return obj

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if request.method == 'OPTIONS':
            return jsonify(status='ok'), 200
        if 'user_id' not in session:
            return jsonify({"message": "Authentication required", "authenticated": False}), 401
        return f(*args, **kwargs)
    return decorated_function

def cleanup_pending_deletes(college, branch, force=False):
    """Clean up stale pending deletes (safety net for failed finalizations)."""
    try:
        query = {
            "college": college, 
            "branch": branch, 
            "pending_delete": True
        }
        
        # If not forced, only delete old ones (grace period)
        if not force:
            # 30 second grace period to be safe
            threshold = datetime.now(timezone.utc) - timedelta(seconds=30)
            query["delete_requested_at"] = {"$lt": threshold}
        
        fs = GridFS(db)
        
        # 1. Result Files
        result_files = list(db.result_files.find(query))
        affected_batches = set()
        for f in result_files:
            if f.get('file_id'):
                try: fs.delete(f['file_id'])
                except: pass
            db.result_files.delete_one({"_id": f["_id"]})
            if f.get('batch_id'):
                affected_batches.add(str(f['batch_id']))
        
        # Recompute dashboard for affected batches
        for bid in affected_batches:
            try:
                # Recompute dashboard summary for the affected batch
                recompute_dashboard_summary(batch_id=str(bid), college=college, branch=branch)
            except Exception as e:
                print(f"⚠️ Warning: Could not recompute dashboard summary for batch {bid} during cleanup: {e}")

        # 2. Intake Files
        intake_files = list(db.intake_files.find(query))
        for f in intake_files:
            if f.get('file_id'):
                try: fs.delete(f['file_id'])
                except: pass
            db.intake_files.delete_one({"_id": f["_id"]})

        # 3. Extracurricular Files
        extra_files = list(db.extracurricular_files.find(query))
        for f in extra_files:
            if f.get('file_id'):
                try: fs.delete(f['file_id'])
                except: pass
            db.extracurricular_files.delete_one({"_id": f["_id"]})

        # 4. Extracurricular Records
        extra_records = list(db.extracurricular_records.find(query))
        for r in extra_records:
            # Delete evidence files
            if "evidence_file_ids" in r and r["evidence_file_ids"]:
                for fid in r["evidence_file_ids"]:
                    try: fs.delete(ObjectId(fid))
                    except: pass
            elif "evidence_file_id" in r and r["evidence_file_id"]:
                try: fs.delete(ObjectId(r["evidence_file_id"]))
                except: pass
            db.extracurricular_records.delete_one({"_id": r["_id"]})

        # 5. Placement Files
        place_files = list(db.placement_outcome_files.find(query))
        for f in place_files:
            if f.get('file_id'):
                try: fs.delete(f['file_id'])
                except: pass
            db.placement_outcome_files.delete_one({"_id": f["_id"]})

        # 6. Placement Outcomes (Records)
        db.placement_outcomes.delete_many(query)

        # 7. Reports
        reports = list(db.reports.find(query))
        for r in reports:
            if r.get('file_id'):
                try: fs.delete(r['file_id'])
                except: pass
            db.reports.delete_one({"_id": r["_id"]})
            
        # print(f"🧹 Cleaned up pending deletes for {college} - {branch}")
            
    except Exception as e:
        print(f"⚠️ Error in cleanup_pending_deletes: {e}")

def send_otp_email(to_email, otp, otp_type='password_reset'):
    try:
        msg = MIMEMultipart()
        msg['From'] = SMTP_EMAIL
        msg['To'] = to_email
        
        if otp_type == 'registration':
            msg['Subject'] = "Your Registration OTP – Student Progression System"
            body = f"Your OTP for account registration is: {otp}\n\nThis OTP is valid for 5 minutes."
        else:
            msg['Subject'] = "Your Password Reset OTP – Student Progression System"
            body = f"Your OTP for password reset is: {otp}\n\nThis OTP is valid for 5 minutes."
        
        msg.attach(MIMEText(body, 'plain'))
        
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(SMTP_EMAIL, SMTP_PASSWORD)
        text = msg.as_string()
        server.sendmail(SMTP_EMAIL, to_email, text)
        server.quit()
        return True
    except Exception as e:
        print(f"SMTP Error: {e}")
        return False

def check_and_update_subscription(college_name):
    """
    Checks if a college's subscription has expired.
    If expired, updates the database and returns the updated college object.
    """
    # Use "name" field as per existing schema
    college = db.colleges.find_one({"name": college_name})
    if not college:
        return None

    expiry = college.get("subscription_expiry")
    if expiry:
        # Using IST timezone for subscription validation (India-first SaaS model)
        ist = pytz.timezone("Asia/Kolkata")
        today = datetime.now(ist).date()
        expiry_date = None

        if isinstance(expiry, str):
            try:
                # Handle YYYY-MM-DD or ISO format
                if 'T' in expiry:
                    expiry_date = datetime.fromisoformat(expiry.replace('Z', '+00:00')).date()
                else:
                    expiry_date = datetime.strptime(expiry, "%Y-%m-%d").date()
            except ValueError:
                pass 
        elif isinstance(expiry, datetime):
            expiry_date = expiry.date()

        # Check expiry and update if needed
        # We only update if it's currently active to avoid redundant writes
        if expiry_date and today > expiry_date and college.get("subscription_active", True):
            db.colleges.update_one(
                {"name": college_name},
                {"$set": {"subscription_active": False}}
            )
            college["subscription_active"] = False

    return college

# --- Auth API Endpoints ---

@app.route('/api/auth/register', methods=['POST'])
def register_user():
    data = request.get_json()
    email = data.get('email', '').strip().lower()
    password = data.get('password', '')
    branch = data.get('branch', '').strip()
    college = data.get('college', '').strip()
    role = data.get('role', 'faculty') # Default to faculty

    if not email or not password or not branch:
        return jsonify({"message": "Email, password, and branch are required"}), 400
    
    if len(password) < 6:
        return jsonify({"message": "Password must be at least 6 characters"}), 400

    # --- SaaS Validations ---
    
    # 1. Validate College Status
    college_doc = check_and_update_subscription(college)
    
    if not college_doc:
        return jsonify({"message": "Invalid college selected. Please contact support."}), 400
        
    if college_doc.get('status') != 'active':
        return jsonify({"message": "College is inactive."}), 403
        
    # 2. Check Subscription
    if not college_doc.get('subscription_active', True):
        return jsonify({
            "success": False,
            "message": "Subscription expired. Please contact administrator."
        }), 403

    # 3. Validate Email Domain
    allowed_domains = college_doc.get('allowed_domains', [])
    if allowed_domains:
        domain = email.split('@')[-1]
        if domain not in allowed_domains:
            return jsonify({"message": f"Registration restricted to official domains: {', '.join(allowed_domains)}"}), 400
    else:
        domain = email.split('@')[-1]
        blocked_domains = ['gmail.com', 'yahoo.com', 'outlook.com', 'hotmail.com', 'rediffmail.com']
        if domain in blocked_domains:
            return jsonify({"message": "Public email domains are not allowed. Please use your official college email."}), 400

    # Check if user exists
    existing_user = db.users.find_one({"email": email})

    # 4. Check User Limit
    if not existing_user:
        max_users = college_doc.get('max_users_per_branch', 2)
        
        # Count existing FACULTY users
        current_user_count = db.users.count_documents({
            "college": college, 
            "branch": branch,
            "role": "faculty" 
        })
        
        if current_user_count >= max_users:
            return jsonify({"message": f"Maximum user limit ({max_users}) reached for {branch} department."}), 400

    if existing_user:
        if existing_user.get('email_verified', False):
            return jsonify({"message": "An account with this email already exists. Each email can be linked to only one college and one branch."}), 400
        # If not verified, we proceed to update the existing record instead of creating new one

    hashed_pw = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())
    
    # Generate OTP for email verification
    otp = ''.join(secrets.choice(string.digits) for _ in range(6))
    expiry = datetime.now(timezone.utc) + timedelta(minutes=5)
    
    db.otps.update_one(
        {"email": email},
        {"$set": {"otp": otp, "expiry": expiry}},
        upsert=True
    )
    
    user_doc = {
        "email": email,
        "password": hashed_pw,
        "branch": branch,
        "college": college,
        "role": role,
        "email_verified": False,
        "created_at": datetime.now(timezone.utc)
    }
    
    if existing_user:
        db.users.update_one({"email": email}, {"$set": user_doc})
    else:
        db.users.insert_one(user_doc)
    
    # Try sending email
    email_sent = send_otp_email(email, otp, otp_type='registration')
    
    if not email_sent:
        # print(f"Failed to send email. OTP for {email} is {otp}")
        # Return success (dev mode fallback)
        pass
    
    return jsonify({"message": "Registration successful. Please verify your email.", "email": email}), 201

@app.route('/api/auth/verify-email', methods=['POST'])
def verify_email():
    data = request.get_json()
    email = data.get('email', '').strip().lower()
    otp = data.get('otp', '').strip()
    
    record = db.otps.find_one({"email": email})
    if not record:
        return jsonify({"message": "Invalid request"}), 400
        
    if record['otp'] != otp:
        return jsonify({"message": "Invalid OTP"}), 400
        
    if record['expiry'].replace(tzinfo=timezone.utc) < datetime.now(timezone.utc):
        return jsonify({"message": "OTP expired"}), 400
        
    # Mark user as verified
    db.users.update_one({"email": email}, {"$set": {"email_verified": True}})
    db.otps.delete_one({"email": email})
    
    return jsonify({"message": "Email verified successfully"}), 200

@app.route('/api/auth/login', methods=['POST'])
def login_user_new():
    data = request.get_json()
    email = data.get('email', '').strip().lower()
    password = data.get('password', '')
    
    user = db.users.find_one({"email": email})
    
    if not user or not bcrypt.checkpw(password.encode('utf-8'), user['password']):
        return jsonify({"message": "Invalid credentials"}), 401
        
    if not user.get('email_verified', False):
         return jsonify({"message": "Email not verified. Please verify your account."}), 403

    # --- SaaS Subscription Check for Login ---
    if user.get('role') != 'super_admin' and user.get('college'):
        college_doc = check_and_update_subscription(user['college'])
        
        if not college_doc:
             return jsonify({"success": False, "message": "College not found."}), 400
             
        if college_doc.get('status') != 'active':
            return jsonify({"success": False, "message": "College is inactive."}), 403
        
        if not college_doc.get('subscription_active', True):
             return jsonify({
                 "success": False,
                 "message": "Subscription expired. Please contact administrator."
             }), 403

    # Set Session
    # Remove specific console-only debugs but keep meaningful logs
    pass
    
    session['user_id'] = str(user['_id'])
    session['email'] = user['email']
    session['branch'] = user['branch']
    session['college'] = user.get('college', '')
    session['role'] = user.get('role', 'faculty')
    
    # Run cleanup for this user's context
    if user.get('college') and user.get('branch'):
        # Run in background or just call it (it's fast enough)
        try:
            cleanup_pending_deletes(user.get('college'), user.get('branch'))
        except Exception as e:
            # print(f"Cleanup failed: {e}")
            pass

    return jsonify({
        "success": True,
        "user": {
            "email": user['email'],
            "branch": user['branch'],
            "college": user.get('college', ''),
            "role": user.get('role', 'faculty')
        }
    }), 200

@app.route('/api/auth/logout', methods=['POST'])
def logout_user():
    session.clear()
    return jsonify({"message": "Logged out successfully"}), 200

@app.route('/api/auth/forgot-password', methods=['POST'])
def forgot_password():
    data = request.get_json()
    email = data.get('email', '').strip().lower()
    
    user = db.users.find_one({"email": email})
    if not user:
        return jsonify({"message": "Email not found"}), 404
        
    otp = ''.join(secrets.choice(string.digits) for _ in range(6))
    expiry = datetime.now(timezone.utc) + timedelta(minutes=5)
    
    db.otps.update_one(
        {"email": email},
        {"$set": {"otp": otp, "expiry": expiry}},
        upsert=True
    )
    
    # Try sending email
    if send_otp_email(email, otp, otp_type='password_reset'):
        return jsonify({"message": "OTP sent to email"}), 200
    else:
        # Log failure
        # print(f"Failed to send email. OTP for {email} is {otp}")
        return jsonify({"message": "Failed to send email (Check server logs for OTP in dev mode)"}), 500

@app.route('/api/auth/verify-otp', methods=['POST'])
def verify_otp():
    data = request.get_json()
    email = data.get('email', '').strip().lower()
    otp = data.get('otp', '').strip()
    
    record = db.otps.find_one({"email": email})
    if not record:
        return jsonify({"message": "Invalid request"}), 400
        
    if record['otp'] != otp:
        return jsonify({"message": "Invalid OTP"}), 400
        
    if record['expiry'].replace(tzinfo=timezone.utc) < datetime.now(timezone.utc):
        return jsonify({"message": "OTP expired"}), 400
        
    return jsonify({"message": "OTP verified"}), 200

@app.route('/api/auth/reset-password', methods=['POST'])
def reset_password():
    data = request.get_json()
    email = data.get('email', '').strip().lower()
    otp = data.get('otp', '').strip()
    new_password = data.get('new_password', '')
    
    if len(new_password) < 6:
        return jsonify({"message": "Password must be at least 6 characters"}), 400

    # Verify OTP again to be safe
    record = db.otps.find_one({"email": email})
    if not record or record['otp'] != otp:
        return jsonify({"message": "Invalid OTP"}), 400
        
    hashed_pw = bcrypt.hashpw(new_password.encode('utf-8'), bcrypt.gensalt())
    db.users.update_one({"email": email}, {"$set": {"password": hashed_pw, "email_verified": True}})
    db.otps.delete_one({"email": email})
    
    return jsonify({"message": "Password reset successfully"}), 200

# --- User Profile & Account Management Endpoints ---

@app.route('/api/user/profile', methods=['GET', 'OPTIONS'])
@login_required
def get_user_profile():
    """Fetch logged-in user's profile details."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
        
    user_id = session.get('user_id')
    user = db.users.find_one({"_id": ObjectId(user_id)})
    
    if not user:
        return jsonify({"message": "User not found"}), 404
        
    # Cleanup pending deletes on profile load
    college = user.get("college")
    branch = user.get("branch")
    if college and branch:
        cleanup_pending_deletes(college, branch)
        
    return jsonify({
        "email": user.get("email"),
        "college": user.get("college"),
        "branch": user.get("branch"),
        "role": user.get("role", "Faculty"),  # Default to Faculty if not set
        "created_at": user.get("created_at")
    }), 200

@app.route('/api/user/delete', methods=['POST', 'OPTIONS'])
@login_required
def delete_account():
    """Permanently delete user account and all associated data."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
        
    data = request.get_json()
    password = data.get('password')
    
    if not password:
        return jsonify({"message": "Password is required to confirm deletion"}), 400
        
    user_id = session.get('user_id')
    user = db.users.find_one({"_id": ObjectId(user_id)})
    
    if not user:
        return jsonify({"message": "User not found"}), 404
        
    # Verify password
    # user['password'] is stored as bytes in some implementations or string. 
    # Usually bcrypt.hashpw returns bytes. Stored in DB as string (decoded).
    stored_pw = user.get('password')
    if isinstance(stored_pw, str):
        stored_pw = stored_pw.encode('utf-8')
        
    if not bcrypt.checkpw(password.encode('utf-8'), stored_pw):
        return jsonify({"message": "Incorrect password"}), 401
        
    college = user.get('college')
    branch = user.get('branch')
    
    if not college or not branch:
        # Fallback if profile is incomplete, just delete user
        db.users.delete_one({"_id": ObjectId(user_id)})
        session.clear()
        return jsonify({"message": "Account deleted (profile was incomplete)"}), 200
        
    try:
        # UPDATED LOGIC: Only delete the user record.
        # Department-level data (batches, results, etc.) MUST persist for the college+branch.
        # This allows re-registration or multiple users for the same department.
        
        # NOTE: Department data deletion must be a separate, explicit admin action.
        
        # 4. Delete User (ONLY THIS IS EXECUTED)
        db.users.delete_one({"_id": ObjectId(user_id)})
        
        # 5. Clear Session
        session.clear()
        
        print(f"⚠️ ACCOUNT DELETED (User Only): {user.get('email')} ({college} - {branch})")
        
        return jsonify({"message": "Account deleted successfully. Department data has been retained."}), 200
        
    except Exception as e:
        print(f"❌ Error deleting account: {e}")
        return jsonify({"message": f"Server error during deletion: {e}"}), 500


# --- Result File Management APIs ---

def build_master_and_summary(batch_id=None, college=None, branch=None):
    """Build master DataFrame and per-semester summary from latest GridFS excel files."""
    fs = GridFS(db)
    sem_keys = [f"sem{i}" for i in range(1, 9)]
    query_base = {"file_type": "excel", "deleted": {"$ne": True}, "pending_delete": {"$ne": True}}
    
    if college:
        query_base["college"] = college
    if branch:
        query_base["branch"] = branch
        
    if batch_id:
        try:
            batch_obj_id = ObjectId(batch_id)
            query_base["batch_id"] = batch_obj_id
        except Exception:
            pass
    
    # Lookup semester files (handles variations like "sem1", "Sem 1")
    latest_docs = {}
    for sem in sem_keys:
        sem_num = sem.replace("sem", "")
        regex_pattern = f"^sem\\s*{sem_num}$"
        
        doc = db.result_files.find_one(
            {**query_base, "semester": {"$regex": regex_pattern, "$options": "i"}},
            sort=[("uploaded_at", -1)]
        )
        latest_docs[sem] = doc

    import pandas as pd
    
    # Check if openpyxl is available
    try:
        import openpyxl
        excel_engine = 'openpyxl'
    except ImportError:
        print("⚠️ Warning: openpyxl not installed. Excel file reading may fail. Install with: pip install openpyxl")
        excel_engine = None
    
    master = pd.DataFrame(columns=["Name", "NORMALIZED_NAME", "Role"] + [f"Sem{i}" for i in range(1, 9)])
    
    # Calculate semester summaries using hybrid logic (File count vs Master Sheet)
    cleared_set: set = set() # Tracks students eligible for next semester (Passed previous)
    sem_file_names = {} # Cache names found in each semester file for "Left" vs "Backlog" distinction

    # Build Master Sheet
    for sem_index in range(1, 9):
        sem = f"sem{sem_index}"
        try:
            doc = latest_docs.get(sem)
            if not doc:
                sem_file_names[sem] = set()
                continue

            grid_id = doc["file_id"]
            file_obj = fs.get(grid_id)
            data = file_obj.read()
            import io
            
            # Use openpyxl engine if available
            if excel_engine:
                raw = pd.read_excel(io.BytesIO(data), header=None, engine=excel_engine)
            else:
                raw = pd.read_excel(io.BytesIO(data), header=None)

            # Flexible header aliases
            name_aliases = {"NAME", "STUDENT NAME", "NAME OF STUDENT"}
            grade_aliases = {"CGPA", "SGPA", "GPA", "GRADE"}

            header_row = None
            for i, row in raw.iterrows():
                row_up = [str(c).strip().upper() for c in row.tolist()]
                if any(alias in row_up for alias in name_aliases) and any(alias in row_up for alias in grade_aliases):
                    header_row = i
                    break
            
            if excel_engine:
                df = pd.read_excel(io.BytesIO(data), skiprows=header_row if header_row is not None else 0, engine=excel_engine)
            else:
                df = pd.read_excel(io.BytesIO(data), skiprows=header_row if header_row is not None else 0)
                
            df.columns = [str(c).strip().upper() for c in df.columns]
            
            name_col = next((c for c in df.columns if c in name_aliases), None)
            cgpa_col = next((c for c in df.columns if c in grade_aliases), None)
            
            if name_col is None or cgpa_col is None:
                sem_file_names[sem] = set()
                continue # Skip if columns not found

            if name_col != "NAME":
                df.rename(columns={name_col: "NAME"}, inplace=True)

            df = df[["NAME", cgpa_col]].dropna(subset=["NAME"]).copy()
            df["NORMALIZED_NAME"] = df["NAME"].apply(normalize_name)
            df[cgpa_col] = pd.to_numeric(df[cgpa_col], errors='coerce')

            sem_map = dict(zip(df["NORMALIZED_NAME"], df[cgpa_col]))
            names_in_file = set(df["NORMALIZED_NAME"])

            # Resolve aliases (renamed students)
            if sem_index > 1:
                 prev_sem_key = f"sem{sem_index-1}"
                 prev_names = sem_file_names.get(prev_sem_key, set())
                 
                 new_names = names_in_file - prev_names
                 missing_names = prev_names - names_in_file
                 
                 if new_names and missing_names:
                     resolved_map = {} # new -> old
                     import difflib
                     
                     available_missing = set(missing_names)
                     
                     for new_name in new_names:
                         best_match = None
                         
                         new_tokens = set(new_name.split())
                         
                         for old_name in available_missing:
                             # Sequence Matcher (Threshold 0.85)
                             ratio = difflib.SequenceMatcher(None, new_name, old_name).ratio()
                             if ratio > 0.85:
                                 best_match = old_name
                                 break 
                             
                             # 2. Token Set Match (Name Reordering / Addition / Deletion)
                             # Rule: Share at least 3 tokens OR (Share 2 tokens AND total tokens <= 3)
                             old_tokens = set(old_name.split())
                             common = new_tokens.intersection(old_tokens)
                             
                             is_token_match = False
                             if len(common) >= 3:
                                 is_token_match = True
                             elif len(common) >= 2 and (len(new_tokens) <= 3 or len(old_tokens) <= 3):
                                 is_token_match = True
                                 
                             if is_token_match:
                                 best_match = old_name
                                 break
                        
                         if best_match:
                             resolved_map[new_name] = best_match
                             available_missing.remove(best_match)
                     
                     # Apply Resolution
                     if resolved_map:
                         pass
                         for new_n, old_n in resolved_map.items():
                             if new_n in sem_map:
                                 # Transfer grade to OLD name
                                 sem_map[old_n] = sem_map[new_n]
                                 del sem_map[new_n]
                                 
                                 # Update names_in_file to reflect OLD name (preserving identity)
                                 names_in_file.remove(new_n)
                                 names_in_file.add(old_n)

            sem_file_names[sem] = names_in_file

            # Update Master Sheet based on Sem Index
            if sem_index == 1:
                df_unique = df.drop_duplicates(subset=["NORMALIZED_NAME"], keep='last')
                df_out = pd.DataFrame({
                    "Name": df_unique["NAME"], 
                    "NORMALIZED_NAME": df_unique["NORMALIZED_NAME"], 
                    "Role": "Regular"
                })
                for i in range(1, 9):
                    df_out[f"Sem{i}"] = pd.NA
                df_out[f"Sem{sem_index}"] = df_out["NORMALIZED_NAME"].map(sem_map)
                
                master = pd.concat([master, df_out], ignore_index=True)
                cleared_set = {n for n in names_in_file if n in sem_map and pd.notna(sem_map[n])}
            
            elif sem_index == 2:
                eligible = cleared_set
                cleared_now = {n for n in eligible if n in sem_map and pd.notna(sem_map[n])}
                master.loc[master["NORMALIZED_NAME"].isin(cleared_now), f"Sem{sem_index}"] = master["NORMALIZED_NAME"].map(sem_map)
                cleared_set = cleared_now
            
            elif sem_index == 3:
                known_names = set(master["NORMALIZED_NAME"]) 
                dse_names = [n for n in names_in_file if n not in known_names]
                if dse_names:
                    dse_subset = df[df["NORMALIZED_NAME"].isin(dse_names)].drop_duplicates(subset=["NORMALIZED_NAME"], keep='last')
                    dse_rows = pd.DataFrame({
                        "Name": dse_subset["NAME"], 
                        "NORMALIZED_NAME": dse_subset["NORMALIZED_NAME"], 
                        "Role": "DSE"
                    })
                    for i in range(1, 9):
                        dse_rows[f"Sem{i}"] = pd.NA
                    dse_rows[f"Sem{sem_index}"] = dse_rows["NORMALIZED_NAME"].map(sem_map)
                    master = pd.concat([master, dse_rows], ignore_index=True)
                
                eligible_regular = set(master.loc[master["Role"] == "Regular", "NORMALIZED_NAME"]).intersection(cleared_set)
                regular_cleared = {n for n in eligible_regular if n in sem_map and pd.notna(sem_map[n])}
                dse_cleared = {n for n in dse_names if n in sem_map and pd.notna(sem_map[n])}
                master.loc[master["NORMALIZED_NAME"].isin(regular_cleared | dse_cleared), f"Sem{sem_index}"] = master["NORMALIZED_NAME"].map(sem_map)
                cleared_set = regular_cleared | dse_cleared
            
            else:
                eligible = cleared_set
                cleared_now = {n for n in eligible if n in sem_map and pd.notna(sem_map[n])}
                master.loc[master["NORMALIZED_NAME"].isin(cleared_now), f"Sem{sem_index}"] = master["NORMALIZED_NAME"].map(sem_map)
                cleared_set = cleared_now

        except Exception as e:
            print(f"⚠️ Error processing {sem}: {str(e)}")
            sem_file_names[sem] = set()

    # --- PHASE 2: Calculate Summary (HYBRID: File Total + Master Status) ---
    summary_rows = []
    
    for sem_index in range(1, 9):
        sem = f"sem{sem_index}"
        col_name = f"Sem{sem_index}"
        
        # 1. TOTAL (Source of Truth = Semester File)
        names_in_file = sem_file_names.get(sem, set())
        
        # If no file uploaded, summary is 0
        if not names_in_file:
            summary_rows.append({"semester": sem, "total": 0, "without_backlog": 0, "with_backlog": 0, "left": 0, "avg_cgpa": None})
            continue
            
        total_appeared = len(names_in_file)
        
        # 2. STATUS (Source of Truth = Master Sheet Cell)
        # We look up ONLY the students present in the file
        # Filter master to rows where NORMALIZED_NAME is in this file
        relevant_rows = master[master["NORMALIZED_NAME"].isin(names_in_file)]
        
        # Without Backlog: Master Cell is NON-EMPTY (Passed)
        without_backlog_count = relevant_rows[col_name].notna().sum()
        
        # With Backlog: FORCED BALANCE (Total - Without)
        # Logic: If they are in the file (Total) and didn't pass (Without Backlog), they MUST have a backlog.
        # This guarantees TOTAL = WITHOUT + WITH.
        with_backlog_count = total_appeared - without_backlog_count
        
        # 3. LEFT (Informational Only - Semester-wise)
        # Rule: Appeared in N-1 AND NOT in N
        left_count = 0
        if sem_index > 1:
            prev_sem = f"sem{sem_index-1}"
            prev_names = sem_file_names.get(prev_sem, set())
            # Note: DSE students appear in Sem 3. They are not in Sem 2.
            # So (Sem2 - Sem3) correctly calculates Regulars who left.
            # DSEs don't trigger "Left" because they weren't in Sem 2.
            left_names = prev_names - names_in_file
            left_count = len(left_names)



            
        # 4. Averages
        avg_cgpa = relevant_rows.loc[relevant_rows[col_name].notna(), col_name].astype(float).mean()
        
        summary_rows.append({
            "semester": sem,
            "total": int(total_appeared),
            "total_students": int(total_appeared), # Standardized alias
            "without_backlog": int(without_backlog_count),
            "passed_students": int(without_backlog_count), # Standardized alias
            "with_backlog": int(with_backlog_count),
            "left": int(left_count),
            "left_students": int(left_count), # Standardized alias
            "avg_cgpa": round(float(avg_cgpa), 2) if pd.notna(avg_cgpa) else None
        })

    # Remove NORMALIZED_NAME column before returning to preserve original structure
    if "NORMALIZED_NAME" in master.columns:
        master.drop(columns=["NORMALIZED_NAME"], inplace=True)

    return master, summary_rows


def _adjust_summary_for_export(master, summary_rows):
    """Apply export-only adjustments inspired by result_parser.py without affecting dashboard data.

    - Align Sem3 metrics to Sem4 student list if both are available
    - Set Sem4 "left" to 0 after alignment
    - Append combined averages (Sem1&Sem2, Sem3&Sem4, Sem5&Sem6, Sem7&Sem8)
    - Sort DSE students by first token of name (surname proxy) in the master sheet
    """
    import math
    # Clone rows to avoid mutating the dashboard data by reference
    rows = [dict(r) for r in summary_rows]

    # Helper map for easy access
    sem_to_row = { (r.get("semester") or "").lower(): r for r in rows }

    # Align Sem3 stats based on Sem4 name list if both exist
    # DISABLED: This legacy logic conflicts with the new cumulative backlog/left tracking.
    # The new logic correctly handles Sem3 and Sem4 independently with cumulative state.
    # try:
    #     sem3_key, sem4_key = "sem3", "sem4"
    #     if sem3_key in sem_to_row and sem4_key in sem_to_row:
    #         r3 = sem_to_row[sem3_key]
    #         r4 = sem_to_row[sem4_key]
    #         # Align Sem3 total to Sem4 total (names in Sem4 file), preserving Sem3 pass count
    #         sem3_total = int(r4.get("total") or 0)
    #         sem3_without_backlog = int(r3.get("without_backlog") or 0)
    #         sem3_with_backlog = max(0, sem3_total - sem3_without_backlog)
    #         r3["total"] = sem3_total
    #         r3["with_backlog"] = sem3_with_backlog
    #         # Set Sem4 left to 0 as per parser logic
    #         r4["left"] = 0
    # except Exception:
    #     # Fail open; keep original if adjustment fails
    #     pass

    # Append combined averages
    def get_avg(label):
        row = sem_to_row.get(label)
        return row.get("avg_cgpa") if row else None

    combined_pairs = [("sem1", "sem2"), ("sem3", "sem4"), ("sem5", "sem6"), ("sem7", "sem8")]
    for a, b in combined_pairs:
        a_avg, b_avg = get_avg(a), get_avg(b)
        if a_avg is not None and b_avg is not None:
            combined_avg = round((float(a_avg) + float(b_avg)) / 2, 2)
            rows.append({
                "semester": f"{a.upper()}&{b.upper()} Combined Avg",
                "total": None,
                "without_backlog": None,
                "with_backlog": None,
                "left": None,
                "avg_cgpa": combined_avg,
            })

    # Sort DSE students by first token (surname proxy) after regulars in master
    try:
        if not master.empty and "Role" in master.columns and "Name" in master.columns:
            regular_df = master[master["Role"] == "Regular"]
            dse_df = master[master["Role"] == "DSE"].copy()
            if not dse_df.empty:
                dse_df["_surname_key"] = dse_df["Name"].astype(str).str.strip().str.split().str[0]
                dse_df = dse_df.sort_values(by="_surname_key").drop(columns=["_surname_key"])
                master = pd.concat([regular_df, dse_df], ignore_index=True)
    except Exception:
        pass

    return master, rows

def recompute_dashboard_summary(batch_id=None, college=None, branch=None):
    """Compute and store dashboard summary in MongoDB."""
    master, summary_rows = build_master_and_summary(batch_id=batch_id, college=college, branch=branch)
    # Apply same alignment and combined averages so dashboard matches export
    try:
        master, summary_rows = _adjust_summary_for_export(master, summary_rows)
    except Exception:
        pass
    
    # Determine summary_id based on batch_id
    if batch_id:
        try:
            batch_obj_id = ObjectId(batch_id)
            summary_id = f"summary_{batch_id}"
        except Exception:
            # Fallback if batch_id is invalid (shouldn't happen with valid calls)
            # If no batch_id, we might need a unique ID for college+branch
            summary_id = f"summary_{college}_{branch}" if college and branch else "summary"
    else:
        summary_id = f"summary_{college}_{branch}" if college and branch else "summary"
    
    # If no data, clear
    if not any((row.get('total') or 0) > 0 for row in summary_rows):
        doc = {"_id": summary_id, "updated_at": datetime.now(timezone.utc), "semesters": [], "kpis": {}}
        if batch_id:
            try:
                doc["batch_id"] = ObjectId(batch_id)
            except:
                pass
        # Add metadata
        if college: doc["college"] = college
        if branch: doc["branch"] = branch
            
        db.dashboard_summary.replace_one({"_id": summary_id}, doc, upsert=True)
        return
        
    # Build KPI data
    kpis = {
        "sanctioned": 0, # Manual input needed?
        "admitted": 0,
        "dse": 0,
        "completed": 0
    }
    
    # Admitted: count of students in Sem1 or Sem2
    # This logic is approximate. Better to use Intake data if available.
    # We will try to fetch intake data later if needed, but for now rely on master sheet.
    
    # KPIs example (can be extended):
    # Preserve existing KPI overrides if present (from admissions upload)
    existing = db.dashboard_summary.find_one({"_id": summary_id}) or {}
    kpis = existing.get("kpis", {})
    
    # Restore intake KPIs from source file if missing (Source of Truth Recovery)
    if batch_id and ("sanctioned_intake" not in kpis or "total_admitted" not in kpis):
        try:
            latest_intake = db.intake_files.find_one(
                {"batch_id": ObjectId(batch_id), "deleted": {"$ne": True}, "pending_delete": {"$ne": True}},
                sort=[("uploaded_at", -1)]
            )
            if latest_intake and "kpis" in latest_intake:
                file_kpis = latest_intake["kpis"]
                if file_kpis.get("sanctioned_intake") is not None:
                    kpis["sanctioned_intake"] = file_kpis["sanctioned_intake"]
                if file_kpis.get("total_admitted") is not None:
                    kpis["total_admitted"] = file_kpis["total_admitted"]
                # Note: We do NOT restore DSE/Completed from intake file here anymore, 
                # because Master Sheet (Results) is the source of truth for these values.
        except Exception as e:
            print(f"⚠️ Error restoring intake KPIs: {e}")

    # Always refresh computed fields available from results (Source of Truth: Master Sheet)
    # Fix for KPI Card Sync Issue: Always recalculate DSE and Completed from master sheet
    dse_calc = 0
    completed_calc = 0
    
    if not master.empty:
        if "Role" in master.columns:
            dse_calc = int(master.loc[master["Role"] == "DSE"].shape[0])
        if "Sem8" in master.columns:
            completed_calc = int(master.dropna(subset=["Sem8"]).shape[0])
            
    pass
    
    # Update KPIs with calculated values (Prioritize actual results over intake/cached values)
    kpis["dse"] = dse_calc
    kpis["successfully_completed"] = completed_calc
    
    # Calculate Second Year API (Sem 3-4) and Third Year API (Sem 5-6)
    sem3_cgpa = None
    sem4_cgpa = None
    sem5_cgpa = None
    sem6_cgpa = None
    
    for row in summary_rows:
        sem_name = row.get("semester", "").lower()
        avg_cgpa = row.get("avg_cgpa")
        if sem_name == "sem3" and avg_cgpa is not None:
            try:
                sem3_cgpa = float(avg_cgpa)
            except (ValueError, TypeError):
                pass
        elif sem_name == "sem4" and avg_cgpa is not None:
            try:
                sem4_cgpa = float(avg_cgpa)
            except (ValueError, TypeError):
                pass
        elif sem_name == "sem5" and avg_cgpa is not None:
            try:
                sem5_cgpa = float(avg_cgpa)
            except (ValueError, TypeError):
                pass
        elif sem_name == "sem6" and avg_cgpa is not None:
            try:
                sem6_cgpa = float(avg_cgpa)
            except (ValueError, TypeError):
                pass
    
    # Second Year API
    if sem3_cgpa is not None and sem4_cgpa is not None:
        kpis["second_year_api"] = round((sem3_cgpa + sem4_cgpa) / 2, 2)
    elif sem3_cgpa is not None:
        kpis["second_year_api"] = round(sem3_cgpa, 2)
    elif sem4_cgpa is not None:
        kpis["second_year_api"] = round(sem4_cgpa, 2)
    else:
        kpis["second_year_api"] = None
    
    # Third Year API
    if sem5_cgpa is not None and sem6_cgpa is not None:
        kpis["third_year_api"] = round((sem5_cgpa + sem6_cgpa) / 2, 2)
    elif sem5_cgpa is not None:
        kpis["third_year_api"] = round(sem5_cgpa, 2)
    elif sem6_cgpa is not None:
        kpis["third_year_api"] = round(sem6_cgpa, 2)
    else:
        kpis["third_year_api"] = None
    
    # Calculate Overall API (average of all semesters with data)
    cgpa_values = []
    for row in summary_rows:
        avg_cgpa = row.get("avg_cgpa")
        if avg_cgpa is not None:
            try:
                cgpa_float = float(avg_cgpa)
                if cgpa_float > 0:
                    cgpa_values.append(cgpa_float)
            except (ValueError, TypeError):
                pass
    
    if cgpa_values:
        kpis["academic_performance_index"] = round(sum(cgpa_values) / len(cgpa_values), 2)
    else:
        kpis["academic_performance_index"] = None

    # Calculate Extracurricular Participation
    if batch_id:
        try:
            batch_obj_id = ObjectId(batch_id)
            
            # Helper to sum participants
            def get_participation(query):
                pipeline = [
                    {"$match": query},
                    {"$group": {
                        "_id": None, 
                        "total": {"$sum": {"$ifNull": ["$number_of_participants", {"$ifNull": ["$number_of_students", {"$ifNull": ["$count_participants", 0]}]}]}}
                    }}
                ]
                res = list(db.extracurricular_records.aggregate(pipeline))
                return res[0]["total"] if res else 0

            # Sports
            kpis["sports_participation"] = get_participation({
                "batch_id": batch_obj_id,
                "$or": [{"output_type": "Sports"}, {"category": "sports"}]
            })
            
            # Technical
            kpis["technical_participation"] = get_participation({
                "batch_id": batch_obj_id,
                "$or": [{"output_type": "Technical"}, {"category": "technical"}]
            })
            
            # Cultural
            kpis["cultural_participation"] = get_participation({
                "batch_id": batch_obj_id,
                "$or": [{"output_type": "Cultural"}, {"category": "cultural"}, {"category": "student activities"}]
            })
            
            # Internships
            kpis["internship_participation"] = get_participation({
                "batch_id": batch_obj_id,
                "$or": [{"output_type": "Internship"}, {"level": "Internship"}, {"category": "internships"}]
            })
            
            # Courses
            kpis["course_participation"] = get_participation({
                "batch_id": batch_obj_id,
                "$or": [{"output_type": "Courses"}, {"category": "courses"}, {"category": "certifications"}]
            })
            
            # Industrial Visits (IV)
            kpis["iv_participation"] = get_participation({
                "batch_id": batch_obj_id,
                "$or": [{"output_type": "Industrial Visit"}, {"category": "industrial_visits"}]
            })
            
        except Exception as e:
            print(f"Error calculating extracurricular KPIs: {e}")

    doc = {"_id": summary_id, "updated_at": datetime.now(timezone.utc), "semesters": summary_rows, "kpis": kpis}
    if batch_id:
        try:
            batch_obj_id = ObjectId(batch_id)
            doc["batch_id"] = batch_obj_id
        except Exception:
            pass
    db.dashboard_summary.replace_one({"_id": summary_id}, doc, upsert=True)


@app.route('/api/upload-admissions-kpis', methods=['POST', 'OPTIONS'])
def upload_admissions_kpis():
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    if 'file' not in request.files:
        return jsonify({"message": "No file part"}), 400
    file = request.files['file']
    batch_id = request.form.get('batch_id')
    
    if not batch_id:
        return jsonify({"message": "Batch is required"}), 400
    
    # Validate batch_id
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
             return jsonify({"message": "User college or branch not found"}), 400

        batch_obj_id = ObjectId(batch_id)
        batch = db.batches.find_one({"_id": batch_obj_id})
        if not batch:
            return jsonify({"message": "Invalid batch ID"}), 400
            
        # Verify college and branch ownership
        if batch.get('college') != college or batch.get('branch') != branch:
            return jsonify({"message": "Unauthorized to access this batch"}), 403

        # Check for duplicate file (Option B: Duplicate Upload Handling)
        existing_file = db.intake_files.find_one({
            "batch_id": batch_obj_id,
            "filename": file.filename
        })
        warning_msg = ""
        if existing_file:
            warning_msg = " Warning: Data from a file with this name may already exist. Re-uploading can cause duplicate entries."

    except Exception:
        return jsonify({"message": "Invalid batch ID format"}), 400
    
    try:
        import pandas as pd
        # Check if openpyxl is available for Excel files
        try:
            import openpyxl
        except ImportError:
            return jsonify({
                "message": "Missing required dependency 'openpyxl'. Please install it using: pip install openpyxl"
            }), 500
        
        # Read content once so we can both parse and store
        file_content = file.read()
        df = pd.read_excel(BytesIO(file_content), engine='openpyxl')
        # Normalize headers
        df.columns = [str(c).strip().lower() for c in df.columns]
        # Flexible header aliases (case-insensitive; we already lower-cased)
        header_aliases = {
            'sanctioned_intake': ['sanctioned intake', 'sanctioned_intake', 'intake'],
            'total_admitted': ['total admitted', 'total_admitted', 'admitted'],
            'dse': ['dse', 'dse admitted', 'direct second year', 'dse intake', 'dse_overridden'],
            'successfully_completed': [
                'successfully passed without backlog',
                'successfully completed',
                'successfully completed (w/o kt)',
                'successfully completed without backlog',
                'without backlog',
                'passed without backlog'
            ]
        }

        def pick(col_keys):
            for key in col_keys:
                if key in df.columns:
                    return key
            return None

        row = df.iloc[0]
        sanctioned_col = pick(header_aliases['sanctioned_intake'])
        total_admitted_col = pick(header_aliases['total_admitted'])
        dse_col = pick(header_aliases['dse'])
        success_completed_col = pick(header_aliases['successfully_completed'])

        # Build overrides, treat missing optional columns gracefully
        overridden = {
            'sanctioned_intake': int(row[sanctioned_col]) if sanctioned_col and pd.notna(row[sanctioned_col]) else None,
            'total_admitted': int(row[total_admitted_col]) if total_admitted_col and pd.notna(row[total_admitted_col]) else None,
            'dse': int(row[dse_col]) if dse_col and pd.notna(row[dse_col]) else None,
            'successfully_completed': int(row[success_completed_col]) if success_completed_col and pd.notna(row[success_completed_col]) else None
        }

        if not sanctioned_col and not total_admitted_col and not dse_col and not success_completed_col:
            return jsonify({"message": "No recognized KPI headers found in the first row."}), 400
        # Merge into existing summary for this batch
        summary_id = f"summary_{batch_id}"
        doc = db.dashboard_summary.find_one({"_id": summary_id}) or {"kpis": {}, "batch_id": batch_obj_id, "college": college, "branch": branch}
        kpis = doc.get('kpis', {})
        for k, v in overridden.items():
            if v is not None:
                kpis[k] = v
        
        # Store the source file into GridFS so it can be downloaded later
        fs = GridFS(db)
        grid_id = fs.put(BytesIO(file_content), filename=file.filename, file_type='intake', college=college, branch=branch, uploaded_at=datetime.now(timezone.utc))
        file_size = len(file_content)

        # Store the file metadata for upload tracking (with GridFS id)
        db.intake_files.insert_one({
            "file_id": grid_id,
            "filename": file.filename,
            "batch_id": batch_obj_id,
            "college": college,
            "branch": branch,
            "uploaded_at": datetime.now(timezone.utc),
            "kpis": overridden,
            "file_size": file_size
        })
        
        # Store batch-specific dashboard summary
        summary_id = f"summary_{batch_id}"
        db.dashboard_summary.replace_one({"_id": summary_id}, {"_id": summary_id, "batch_id": batch_obj_id, "updated_at": datetime.now(timezone.utc), "semesters": doc.get('semesters', []), "kpis": kpis}, upsert=True)
        return jsonify({"message": "Admissions KPIs uploaded successfully" + warning_msg}), 200
    except Exception as e:
        return jsonify({"message": f"Error parsing KPI file: {e}"}), 500




@app.route('/api/dashboard-summary', methods=['GET'])
def get_dashboard_summary():
    college = session.get('college')
    branch = session.get('branch')
    
    if not college or not branch:
        missing = []
        if not college: missing.append("college")
        if not branch: missing.append("branch")
        return jsonify({"message": f"User profile incomplete: missing {', '.join(missing)}. Please logout and login again."}), 400

    # Run cleanup for this user's context
    cleanup_pending_deletes(college, branch)

    batch_id = request.args.get('batch_id')
    
    # Determine summary_id based on batch_id
    if batch_id:
        try:
            batch_obj_id = ObjectId(batch_id)
            # Verify batch exists and belongs to college + branch
            batch = db.batches.find_one({"_id": batch_obj_id, "college": college, "branch": branch})
            if not batch:
                 # If batch provided but not authorized, return empty or error? 
                 # Dashboard usually expects data. If unauthorized, maybe empty.
                 # But usually 403 is better. However, frontend might crash.
                 # Let's return empty summary structure.
                 return jsonify({
                    "kpis": {
                        "sanctioned_intake": 0,
                        "total_admitted": 0,
                        "dse": 0,
                        "successfully_completed": 0
                    },
                    "semesters": []
                }), 200

            summary_id = f"summary_{batch_id}"
        except Exception:
            summary_id = f"summary_{college}_{branch}"
    else:
        summary_id = f"summary_{college}_{branch}"
    
    # Check if intake files exist for this batch (only if batch_id provided)
    if batch_id:
        try:
            batch_obj_id = ObjectId(batch_id)
            intake_files_count = db.intake_files.count_documents({"batch_id": batch_obj_id, "college": college})
            
            # If no intake files exist, clear intake-related KPIs from dashboard summary
            if intake_files_count == 0:
                summary_doc = db.dashboard_summary.find_one({"_id": summary_id})
                if summary_doc:
                    kpis = summary_doc.get('kpis', {})
                    # Clear intake-related KPIs
                    kpis.pop('sanctioned_intake', None)
                    kpis.pop('total_admitted', None)
                    
                    # Update the dashboard summary
                    db.dashboard_summary.update_one(
                        {"_id": summary_id},
                        {"$set": {
                            "kpis": kpis,
                            "updated_at": datetime.now(timezone.utc)
                        }}
                    )
        except Exception:
            pass  # If batch_id is invalid, continue normally
    
    try:
        recompute_dashboard_summary(batch_id=batch_id, college=college, branch=branch)
    except Exception:
        pass
    doc = db.dashboard_summary.find_one({"_id": summary_id})
    if not doc:
        # Return empty summary if not found
        return jsonify({
            "kpis": {
                "sanctioned_intake": 0,
                "total_admitted": 0,
                "dse": 0,
                "successfully_completed": 0
            },
            "semesters": []
        }), 200
    # Convert ObjectId to string for JSON serialization
    doc = convert_objectid_to_str(doc)

    return jsonify(doc), 200


@app.route('/api/export-result-summary', methods=['GET'])
@login_required
def export_result_summary():
    """Generate and download the combined result summary Excel (master + summary)."""
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
            print(f"❌ User profile incomplete in session: {session}")
            missing = []
            if not college: missing.append("college")
            if not branch: missing.append("branch")
            return jsonify({"message": f"User profile incomplete: missing {', '.join(missing)}. Please logout and login again."}), 400

        batch_id = request.args.get('batch_id', '')
        
        if not batch_id:
            return jsonify({"message": "Batch ID is required"}), 400
        
        # Validate batch_id
        try:
            batch_obj_id = ObjectId(batch_id)
            batch = db.batches.find_one({"_id": batch_obj_id})
            if not batch:
                return jsonify({"message": "Invalid batch ID"}), 400
            
            # Verify college and branch ownership
            if batch.get('college') != college or batch.get('branch') != branch:
                return jsonify({"message": "Unauthorized to access this batch"}), 403

            batch_range = batch.get('batch_range', 'Result')
        except Exception:
            return jsonify({"message": "Invalid batch ID format"}), 400
        
        import pandas as pd
        master, summary_rows = build_master_and_summary(batch_id=batch_id, college=college, branch=branch)
        # Apply export-only adjustments to reflect result_parser.py behavior
        master, summary_rows = _adjust_summary_for_export(master, summary_rows)
        # Build summary dataframe
        summary_df = pd.DataFrame(summary_rows)
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            master.to_excel(writer, sheet_name='Master', index=False)
            summary_df.to_excel(writer, sheet_name='Summary', index=False)
        output.seek(0)
        from flask import Response
        return Response(
            output.read(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={'Content-Disposition': f'attachment; filename="Result_Summary_{batch_range.replace(" ", "_")}.xlsx"'}
        )
    except Exception as e:
        return jsonify({"message": f"Error generating export: {e}"}), 500
@app.route('/api/upload-result-file', methods=['POST', 'OPTIONS'])
@login_required
def upload_result_file():
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    if 'file' not in request.files:
        return jsonify({"message": "No file part"}), 400
    
    file = request.files['file']
    semester = request.form.get('semester')
    file_type = request.form.get('file_type', 'excel')
    batch_id = request.form.get('batch_id')
    
    if not semester:
        return jsonify({"message": "Semester is required"}), 400
    
    if not batch_id:
        return jsonify({"message": "Batch is required"}), 400
    
    # Validate batch_id
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
            print(f"❌ get_batches: User profile incomplete in session: {session}")
            missing = []
            if not college: missing.append("college")
            if not branch: missing.append("branch")
            return jsonify({"message": f"User profile incomplete: missing {', '.join(missing)}. Please logout and login again."}), 400

        batch_obj_id = ObjectId(batch_id)
        batch = db.batches.find_one({"_id": batch_obj_id})
        if not batch:
            return jsonify({"message": "Invalid batch ID"}), 400
        
        # Verify college and branch ownership
        if batch.get('college') != college or batch.get('branch') != branch:
            return jsonify({"message": "Unauthorized to access this batch"}), 403

        # Check for duplicate file (Option B: Duplicate Upload Handling)
        existing_file = db.result_files.find_one({
            "batch_id": batch_obj_id,
            "semester": semester,
            "filename": file.filename
        })
        warning_msg = ""
        if existing_file:
             warning_msg = f" Warning: Data from a file with this name may already exist for {semester}. Re-uploading can cause duplicate entries."

    except Exception:
        return jsonify({"message": "Invalid batch ID format"}), 400
    
    if file.filename == '':
        return jsonify({"message": "No file selected"}), 400
    
    try:
        # Store file in MongoDB GridFS
        fs = GridFS(db)
        
        # Generate unique filename
        import uuid
        file_id = str(uuid.uuid4())
        filename = f"{semester}_{file_id}_{file.filename}"
        
        # Read file content for validation and size
        file_content = file.read()
        file_size = len(file_content)
        
        # --- Strict Validation ---
        try:
            import io
            import pandas as pd
            
            # Check if openpyxl is available
            try:
                import openpyxl
                excel_engine = 'openpyxl'
            except ImportError:
                excel_engine = None

            if excel_engine:
                raw = pd.read_excel(io.BytesIO(file_content), header=None, engine=excel_engine)
            else:
                raw = pd.read_excel(io.BytesIO(file_content), header=None)
            
            name_aliases = {"NAME", "STUDENT NAME", "NAME OF STUDENT"}
            grade_aliases = {"CGPA", "SGPA", "GPA", "GRADE"}
            
            header_found = False
            potential_missing_grade = False
            potential_missing_name = False
            
            for i, row in raw.iterrows():
                row_str = [str(c).strip().upper() for c in row.tolist()]
                
                has_name = any(alias in row_str for alias in name_aliases)
                has_grade = any(alias in row_str for alias in grade_aliases)
                
                if has_name and has_grade:
                    header_found = True
                    break
                
                if has_name:
                    potential_missing_grade = True
                elif has_grade:
                    potential_missing_name = True
            
            if not header_found:
                if potential_missing_grade:
                    return jsonify({"message": "CGPA/Grade column not found. Expected: CGPA / SGPA / GPA / Grade"}), 400
                elif potential_missing_name:
                    return jsonify({"message": "Name column not found. Expected: Name / Student Name / Name of Student"}), 400
                else:
                    return jsonify({"message": "Name column not found. Expected: Name / Student Name / Name of Student"}), 400
                    
        except Exception as e:
             return jsonify({"message": f"Invalid file content: {str(e)}"}), 400
        # -------------------------

        # Reset file pointer
        file.seek(0)

        
        # Store file
        file_id = fs.put(
            file, 
            filename=filename, 
            semester=semester, 
            file_type=file_type, 
            uploaded_at=datetime.now(timezone.utc),
            college=college,
            branch=branch
        )
        
        # Store metadata in collection
        db.result_files.insert_one({
            "file_id": file_id,
            "filename": file.filename,
            "semester": semester,
            "file_type": file_type,
            "batch_id": batch_obj_id,
            "college": college,
            "branch": branch,
            "uploaded_at": datetime.now(timezone.utc),
            "file_size": file_size
        })
        
        # Recompute dashboard summary after upload (with batch_id)
        try:
            recompute_dashboard_summary(batch_id=str(batch_obj_id), college=college, branch=branch)
            # Remove debug logs for dashboard recompute
            pass
        except Exception as e:
            print(f"⚠️ Error recomputing dashboard summary: {str(e)}")
            import traceback
            traceback.print_exc()

        return jsonify({"message": f"File uploaded successfully for {semester.upper()}" + warning_msg}), 200
        
    except Exception as e:
        return jsonify({"message": f"Error uploading file: {str(e)}"}), 500

@app.route('/api/upload-gazette-file', methods=['POST', 'OPTIONS'])
@login_required
def upload_gazette_file():
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    if 'file' not in request.files:
        return jsonify({"message": "No file part"}), 400
    
    file = request.files['file']
    semester = request.form.get('semester')
    batch_id = request.form.get('batch_id')
    
    if not semester:
        return jsonify({"message": "Semester is required"}), 400
    
    if not batch_id:
        return jsonify({"message": "Batch is required"}), 400
    
    # Validate batch_id
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
            missing = []
            if not college: missing.append("college")
            if not branch: missing.append("branch")
            return jsonify({"message": f"User profile incomplete: missing {', '.join(missing)}. Please logout and login again."}), 400

        batch_obj_id = ObjectId(batch_id)
        batch = db.batches.find_one({"_id": batch_obj_id})
        if not batch:
            return jsonify({"message": "Invalid batch ID"}), 400
        
        # Verify college and branch ownership
        if batch.get('college') != college or batch.get('branch') != branch:
            return jsonify({"message": "Unauthorized to access this batch"}), 403

    except Exception:
        return jsonify({"message": "Invalid batch ID format"}), 400
    
    if file.filename == '':
        return jsonify({"message": "No file selected"}), 400
    
    try:
        # Store file in MongoDB GridFS
        fs = GridFS(db)
        
        # Generate unique filename
        import uuid
        file_id = str(uuid.uuid4())
        filename = f"{semester}_{file_id}_{file.filename}"
        
        # Read file content for size calculation
        file_content = file.read()
        file_size = len(file_content)
        
        # Reset file pointer
        file.seek(0)
        
        # Store file
        file_id = fs.put(file, filename=filename, semester=semester, file_type='pdf', college=college, branch=branch, uploaded_at=datetime.now(timezone.utc))
        
        # Store metadata in collection
        db.result_files.insert_one({
            "file_id": file_id,
            "filename": file.filename,
            "semester": semester,
            "file_type": 'pdf',
            "batch_id": batch_obj_id,
            "college": college,
            "branch": branch,
            "uploaded_at": datetime.now(timezone.utc),
            "file_size": file_size
        })
        
        # Gazette does not impact summary; skip recompute
        return jsonify({"message": f"Gazette file uploaded successfully for {semester.upper()}"}), 200
        
    except Exception as e:
        return jsonify({"message": f"Error uploading file: {str(e)}"}), 500

@app.route('/api/get-result-files', methods=['GET', 'OPTIONS'])
@login_required
def get_result_files():
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    semester = request.args.get('semester', '')
    batch_id = request.args.get('batch_id', '')
    
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
            missing = []
            if not college: missing.append("college")
            if not branch: missing.append("branch")
            return jsonify({"message": f"User profile incomplete: missing {', '.join(missing)}. Please logout and login again."}), 400

        # Run cleanup for this user's context
        cleanup_pending_deletes(college, branch)

        # Get all batch IDs for this branch to filter files
        branch_batches = list(db.batches.find({"college": college, "branch": branch}, {"_id": 1}))
        branch_batch_ids = [b["_id"] for b in branch_batches]

        query = {"college": college, "branch": branch}
        if semester:
            query['semester'] = semester
        
        # Filter by batch_id if provided
        if batch_id:
            try:
                batch_obj_id = ObjectId(batch_id)
                # Verify batch belongs to user's branch
                if batch_obj_id not in branch_batch_ids:
                    return jsonify({"message": "Unauthorized access to batch"}), 403
                query['batch_id'] = batch_obj_id
            except Exception:
                return jsonify({"message": "Invalid batch ID format"}), 400
        else:
            # If no specific batch, only show files from batches belonging to user's branch
            query['batch_id'] = {"$in": branch_batch_ids}
        
        # Exclude soft-deleted files
        query['deleted'] = {"$ne": True}
        query['pending_delete'] = {"$ne": True}
        
        files = list(db.result_files.find(query).sort('uploaded_at', -1))
        
        # Add batch_range to each file for display
        for file in files:
            if file.get('batch_id'):
                batch = db.batches.find_one({"_id": file['batch_id']})
                if batch:
                    file['batch_range'] = batch.get('batch_range', 'N/A')
                else:
                    file['batch_range'] = 'N/A'
            else:
                file['batch_range'] = 'N/A'
        
        # Convert ObjectId to string for JSON serialization
        files = convert_objectid_to_str(files)
        
        return jsonify(files), 200
        
    except Exception as e:
        return jsonify({"message": f"Error retrieving files: {str(e)}"}), 500

@app.route('/api/get-intake-files', methods=['GET', 'OPTIONS'])
@login_required
def get_intake_files():
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    batch_id = request.args.get('batch_id', '')
    
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
            missing = []
            if not college: missing.append("college")
            if not branch: missing.append("branch")
            return jsonify({"message": f"User profile incomplete: missing {', '.join(missing)}. Please logout and login again."}), 400

        # Run cleanup for this user's context
        cleanup_pending_deletes(college, branch)

        # Get all batch IDs for this branch to filter files
        branch_batches = list(db.batches.find({"college": college, "branch": branch}, {"_id": 1}))
        branch_batch_ids = [b["_id"] for b in branch_batches]

        query = {"college": college, "branch": branch}
        
        # Filter by batch_id if provided
        if batch_id:
            try:
                batch_obj_id = ObjectId(batch_id)
                if batch_obj_id not in branch_batch_ids:
                    return jsonify({"message": "Unauthorized access to batch"}), 403
                query['batch_id'] = batch_obj_id
            except Exception:
                return jsonify({"message": "Invalid batch ID format"}), 400
        else:
            query['batch_id'] = {"$in": branch_batch_ids}

        # Exclude soft-deleted files
        query['deleted'] = {"$ne": True}
        query['pending_delete'] = {"$ne": True}

        files = list(db.intake_files.find(query).sort('uploaded_at', -1))
        
        # Add batch_range to each file for display
        for file in files:
            if file.get('batch_id'):
                batch = db.batches.find_one({"_id": file['batch_id']})
                if batch:
                    file['batch_range'] = batch.get('batch_range', 'N/A')
                else:
                    file['batch_range'] = 'N/A'
            else:
                file['batch_range'] = 'N/A'
        
        # Convert ObjectId to string for JSON serialization
        files = convert_objectid_to_str(files)
        
        return jsonify(files), 200
        
    except Exception as e:
        return jsonify({"message": f"Error retrieving intake files: {str(e)}"}), 500

@app.route('/api/download-intake-file/<file_id>', methods=['GET', 'OPTIONS'])
@login_required
def download_intake_file(file_id):
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
             return jsonify({"message": "User college or branch not found"}), 400

        fs = GridFS(db)
        file_doc = db.intake_files.find_one({"_id": ObjectId(file_id)})
        if not file_doc:
            return jsonify({"message": "File not found"}), 404
            
        # Verify ownership via batch
        batch_id = file_doc.get('batch_id')
        if batch_id:
            batch = db.batches.find_one({"_id": batch_id})
            if batch and batch.get('branch') != branch:
                return jsonify({"message": "Unauthorized access"}), 403

        grid_id = file_doc.get('file_id')
        if not grid_id:
            return jsonify({"message": "File data not available"}), 404
        file_data = fs.get(grid_id)
        if not file_data:
            return jsonify({"message": "File data not found"}), 404
        from flask import Response
        return Response(
            file_data.read(),
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={'Content-Disposition': f'attachment; filename="{file_doc.get("filename", "intake.xlsx")}"'}
        )
    except Exception as e:
        return jsonify({"message": f"Error downloading intake file: {str(e)}"}), 500

@app.route('/api/delete-intake-file/<file_id>', methods=['DELETE', 'OPTIONS'])
@login_required
def delete_intake_file(file_id):
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
             return jsonify({"message": "User college or branch not found"}), 400

        # Get the file document before deleting to know which batch it belongs to
        file_doc = db.intake_files.find_one({"_id": ObjectId(file_id)})
        if not file_doc:
            return jsonify({"message": "Intake file not found"}), 404
        
        batch_id = file_doc.get('batch_id')
        
        # Verify ownership
        if batch_id:
            batch = db.batches.find_one({"_id": batch_id})
            if batch and (batch.get('college') != college or batch.get('branch') != branch):
                return jsonify({"message": "Unauthorized access"}), 403
        
        finalize = (request.args.get('finalize') or '').lower() == 'true'
        if not finalize:
            # Soft delete - Pending Delete State
            db.intake_files.update_one(
                {"_id": ObjectId(file_id)},
                {"$set": {
                    "pending_delete": True,
                    "delete_requested_at": datetime.now(timezone.utc)
                }}
            )
            return jsonify({"message": "Intake file marked for deletion", "pending_delete": True}), 200
        
        # Permanent delete
        result = db.intake_files.delete_one({"_id": ObjectId(file_id)})
        if result.deleted_count == 0:
            return jsonify({"message": "Intake file not found"}), 404
        
        # Check if there are any remaining intake files for this batch
        if batch_id:
            pass
        
        return jsonify({"message": "Intake file reference removed successfully"}), 200
        
    except Exception as e:
        return jsonify({"message": f"Error deleting intake file: {str(e)}"}), 500

@app.route('/api/restore-intake-file/<file_id>', methods=['POST', 'OPTIONS'])
@login_required
def restore_intake_file(file_id):
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    try:
        college = session.get('college')
        branch = session.get('branch')
        if not college or not branch:
             return jsonify({"message": "User college or branch not found"}), 400
        
        file_doc = db.intake_files.find_one({"_id": ObjectId(file_id)})
        if not file_doc:
            return jsonify({"message": "Intake file not found"}), 404
        
        batch_id = file_doc.get('batch_id')
        if batch_id:
            batch = db.batches.find_one({"_id": batch_id})
            if batch and (batch.get('college') != college or batch.get('branch') != branch):
                return jsonify({"message": "Unauthorized access"}), 403
        
        db.intake_files.update_one(
            {"_id": ObjectId(file_id)},
            {"$unset": {"pending_delete": "", "delete_requested_at": "", "deleted": ""}}
        )
        return jsonify({"message": "Intake file restored"}), 200
    except Exception as e:
        return jsonify({"message": f"Error restoring intake file: {str(e)}"}), 500

@app.route('/api/finalize-all-pending', methods=['POST', 'OPTIONS'])
@login_required
def finalize_all_pending():
    """
    Endpoint to force-finalize all pending deletes for the current user.
    This is called by the frontend on page unload/refresh to ensure no pending deletes are left behind.
    """
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
        
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
             return jsonify({"message": "User college or branch not found"}), 400
             
        # Force cleanup of all pending deletes
        cleanup_pending_deletes(college, branch, force=True)
        
        return jsonify({"message": "All pending deletes finalized"}), 200
    except Exception as e:
        print(f"Error in finalize_all_pending: {str(e)}")
        return jsonify({"message": f"Error finalizing pending deletes: {str(e)}"}), 500

@app.route('/api/download-result-file/<file_id>', methods=['GET', 'OPTIONS'])
@login_required
def download_result_file(file_id):
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
             return jsonify({"message": "User college or branch not found"}), 400

        fs = GridFS(db)
        
        file_doc = db.result_files.find_one({"_id": ObjectId(file_id)})
        if not file_doc:
            return jsonify({"message": "File not found"}), 404
        
        # Verify ownership via batch
        batch_id = file_doc.get('batch_id')
        if batch_id:
            batch = db.batches.find_one({"_id": batch_id})
            if batch and (batch.get('college') != college or batch.get('branch') != branch):
                return jsonify({"message": "Unauthorized access"}), 403

        gridfs_file_id = file_doc['file_id']
        file_data = fs.get(gridfs_file_id)
        if not file_data:
            return jsonify({"message": "File data not found"}), 404
        
        # Determine content type based on file extension
        filename = file_doc['filename'].lower()
        if filename.endswith('.pdf'):
            mimetype = 'application/pdf'
        elif filename.endswith('.xlsx'):
            mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        elif filename.endswith('.xls'):
            mimetype = 'application/vnd.ms-excel'
        else:
            mimetype = 'application/octet-stream'
        
        # Return file data for download
        from flask import Response
        return Response(
            file_data.read(),
            mimetype=mimetype,
            headers={
                'Content-Disposition': f'attachment; filename="{file_doc["filename"]}"'
            }
        )
        
    except Exception as e:
        return jsonify({"message": f"Error downloading file: {str(e)}"}), 500

@app.route('/api/delete-result-file/<file_id>', methods=['DELETE', 'OPTIONS'])
@login_required
def delete_result_file(file_id):
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
             return jsonify({"message": "User college or branch not found"}), 400

        fs = GridFS(db)
        
        # Check if file exists in metadata
        file_doc = db.result_files.find_one({"_id": ObjectId(file_id)})
        if not file_doc:
            return jsonify({"message": "File not found"}), 404
        
        # Get batch_id before deletion for dashboard recomputation
        batch_id = file_doc.get('batch_id')
        
        # Verify ownership
        if batch_id:
            batch = db.batches.find_one({"_id": batch_id})
            if batch and (batch.get('college') != college or batch.get('branch') != branch):
                return jsonify({"message": "Unauthorized access"}), 403
        
        finalize = (request.args.get('finalize') or '').lower() == 'true'
        if not finalize:
            # Soft delete - Pending Delete State
            db.result_files.update_one(
                {"_id": ObjectId(file_id)},
                {"$set": {
                    "pending_delete": True,
                    "delete_requested_at": datetime.now(timezone.utc)
                }}
            )
            
            # NOTE: We do NOT recompute dashboard here anymore. 
            # Recomputation only happens on final deletion.

            return jsonify({"message": "File marked for deletion", "pending_delete": True}), 200
        
        # Permanent delete
        try:
            fs.delete(file_doc['file_id'])
        except Exception:
            pass
        db.result_files.delete_one({"_id": ObjectId(file_id)})
        
        # Recompute dashboard to reflect data removal - ONLY on final delete
        if batch_id:
           try:
               recompute_dashboard_summary(batch_id=str(batch_id), college=college, branch=branch)
           except Exception as e:
               print(f"⚠️ Warning: Could not recompute dashboard summary for batch {batch_id}: {e}")
        
        return jsonify({"message": "File reference removed successfully"}), 200

        
    except Exception as e:
        return jsonify({"message": f"Error deleting file: {str(e)}"}), 500

@app.route('/api/restore-result-file/<file_id>', methods=['POST', 'OPTIONS'])
@login_required
def restore_result_file(file_id):
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
             return jsonify({"message": "User college or branch not found"}), 400
        
        file_doc = db.result_files.find_one({"_id": ObjectId(file_id)})
        if not file_doc:
            return jsonify({"message": "File not found"}), 404
        
        batch_id = file_doc.get('batch_id')
        if batch_id:
            batch = db.batches.find_one({"_id": batch_id})
            if batch and (batch.get('college') != college or batch.get('branch') != branch):
                return jsonify({"message": "Unauthorized access"}), 403
        
        db.result_files.update_one(
            {"_id": ObjectId(file_id)},
            {"$unset": {"pending_delete": "", "delete_requested_at": "", "deleted": ""}}
        )
        return jsonify({"message": "File restored successfully"}), 200
    except Exception as e:
        return jsonify({"message": f"Error restoring file: {str(e)}"}), 500

# Duplicate endpoint removed

# --- Extracurricular Module APIs ---

# Initialize indexes for extracurricular_records collection
def init_extracurricular_indexes():
    """Create indexes for extracurricular_records collection."""
    try:
        db.extracurricular_records.create_index([("academic_year", 1)])
        db.extracurricular_records.create_index([("event_date", 1)])
        db.extracurricular_records.create_index([("category", 1)])
        # Silent on success
        pass
    except Exception as e:
        print(f"⚠️ Warning: Could not create extracurricular indexes: {e}")

# Initialize indexes on startup
init_extracurricular_indexes()

# Validation constants
EXTRACURRICULAR_CATEGORIES = ["Sports", "Technical", "Cultural", "Student Activities"]  # Legacy support
EXTRACURRICULAR_OUTPUT_TYPES = ["Sports", "Cultural", "Technical", "Internship", "Courses", "Industrial Visit"]

# Backward compatibility helper: normalize old output_type names to new ones
def normalize_output_type(output_type):
    """Normalize output_type for backward compatibility."""
    if not output_type:
        return output_type
    output_type = str(output_type).strip()
    # Map old names to new names
    mappings = {
        "College-Level Events": "Cultural",
        "Technical Events": "Technical",
        "Internships": "Internship",
        "Courses / Certifications": "Courses"
    }
    return mappings.get(output_type, output_type)

def normalize_choice(value, choices):
    if not value:
        return value
    value_str = str(value).strip()
    lower_val = value_str.lower()
    for c in choices:
        if lower_val == str(c).strip().lower():
            return c
    return value_str

def validate_extracurricular_record(data, is_bulk=False):
    """Validate a single extracurricular record (supports both old and new structure).
    
    Returns: (is_valid, error_message, validated_data)
    """
    errors = []
    validated = {}
    
    # Check if this is new output-type based structure
    output_type = data.get("output_type", "").strip()
    
    # Legacy structure (for backward compatibility)
    if not output_type and "category" in data:
        # Old validation logic for backward compatibility
        required_fields = ["academic_year", "event_date", "category", "count_participants"]
        for field in required_fields:
            if field not in data or (isinstance(data[field], str) and not data[field].strip()):
                errors.append(f"Missing required field: {field}")
            else:
                validated[field] = data[field]

        if errors:
            return False, "; ".join(errors), None
        
        category_raw = str(validated["category"]).strip()
        validated["category"] = category_raw # Removed strict normalization
        
        try:
            event_date = pd.to_datetime(validated["event_date"]).date()
            validated["event_date"] = event_date.isoformat()
        except Exception:
            errors.append("Invalid event_date format. Expected YYYY-MM-DD")
        
        try:
            count_participants = int(validated["count_participants"])
            if count_participants < 0:
                errors.append("count_participants must be non-negative")
            else:
                validated["count_participants"] = count_participants
        except (ValueError, TypeError):
            errors.append("count_participants must be a non-negative integer")
        
        if "count_won" in data and data["count_won"] is not None:
            try:
                count_won = int(data["count_won"])
                if count_won < 0:
                    errors.append("count_won must be non-negative")
                elif count_won > validated["count_participants"]:
                    errors.append("count_won cannot exceed count_participants")
                else:
                    validated["count_won"] = count_won
            except (ValueError, TypeError):
                errors.append("count_won must be a non-negative integer")
        else:
            validated["count_won"] = 0
        
        if "level" in data:
            validated["level"] = str(data["level"]).strip() if data["level"] else None
        if "type" in data:
            validated["type"] = str(data["type"]).strip() if data["type"] else None
        if "organizer" in data:
            validated["organizer"] = str(data["organizer"]).strip() if data["organizer"] else None
        if "remarks" in data:
            validated["remarks"] = str(data["remarks"]).strip() if data["remarks"] else None
        
        if errors:
            return False, "; ".join(errors), None
        
        return True, None, validated
    
    # New output-type based structure
    if not output_type:
        errors.append("Missing required field: output_type")
        return False, "; ".join(errors), None
    
    # Normalize "College-Level Events" to "Cultural" for backward compatibility if needed, but allow custom
    output_type = normalize_output_type(output_type)
    
    # Removed strict EXTRACURRICULAR_OUTPUT_TYPES check to allow free text
    
    validated["output_type"] = output_type
    
    # Common required field
    if "academic_year" not in data or (isinstance(data["academic_year"], str) and not data["academic_year"].strip()):
        errors.append("Missing required field: academic_year")
    else:
        validated["academic_year"] = str(data["academic_year"]).strip()
    
    # Output-type specific validation - relaxed for free text
    # We still check specific fields if they exist, but we don't strictly enforce structure based on type
    
    # Participants
    participants_raw = data.get("number_of_participants", data.get("number_of_students"))
    if participants_raw not in (None, ""):
        try:
            validated["number_of_participants"] = int(float(participants_raw))
            if validated["number_of_participants"] < 0:
                errors.append("number_of_participants must be non-negative")
        except (ValueError, TypeError):
             # If not an integer, maybe it's free text describing participants? 
             # For now, let's keep it numeric as it's a count, but be lenient if missing?
             # User asked for "free-text inputs", but a count is usually a number.
             # I'll stick to numeric for "number_of_..." fields to avoid DB issues if schema expects int.
            errors.append("number_of_participants must be an integer")
    else:
        # Relax mandatory requirement? Or keep it? User said "retain required fields".
        # Assuming count is required.
        errors.append("Missing required field: number_of_participants/number_of_students")

    # Wins (Optional)
    if "number_of_wins" in data and data["number_of_wins"] not in (None, ""):
        try:
            validated["number_of_wins"] = int(float(data["number_of_wins"]))
        except (ValueError, TypeError):
            errors.append("number_of_wins must be an integer")

    # Collect all other potential fields loosely
    optional_fields = [
        "sports_name", "level", "organizer", "remarks", "proof_link_or_file", "evidence_link",
        "activity_name", "category", "custom_category", "date", "event_name",
        "event_type", "organization", "outcome", "achievement", "organization_name", 
        "internship_domain", "internship_title", "duration", "mode", "stipend",
        "course_name", "platform", "certification_status",
        "industrial_visit_details", "number_of_faculty",
        "event_category", "cultural_custom_category", "has_stipend", "stipend_amount"
    ]
    
    for key in optional_fields:
        if key in data and data[key] is not None:
             validated[key] = str(data[key]).strip()

    # Technical Validation (Achievement/Outcome is mandatory)
    if output_type == "Technical":
        if not validated.get("achievement") and not validated.get("outcome"):
             errors.append("Missing required field: achievement")

    # Stipend amount (numeric check)
    if "stipend_amount" in data and data["stipend_amount"] not in (None, ""):
        try:
            validated["stipend_amount"] = float(data["stipend_amount"])
        except (ValueError, TypeError):
            errors.append("stipend_amount must be a number")
            
    # Names array
    validated["student_names"] = []
    names = data.get("student_names")
    
    if names:
        if isinstance(names, str):
            if names.strip():
                validated["student_names"] = [n.strip() for n in names.split(",") if n.strip()]
        elif isinstance(names, list):
             validated["student_names"] = [str(n).strip() for n in names if str(n).strip()]
    
    # Additional specific required fields based on "loose" type matching or just generic requirements
    # If type is Internship, organization_name is usually key.
    # To support free text, we can't strictly say "If Internship then X is required" if the user typed "Internship Summer".
    # So we will rely on the UI to enforce required fields for known types, but backend will be permissive 
    # except for the absolute basics (academic_year, number_of_participants).
    
    if errors:
        return False, "; ".join(errors), None
        
    # Standardize field names (Aliases -> Canonical)
    # Canonical: number_of_participants
    if "count_participants" in validated:
        validated["number_of_participants"] = validated.pop("count_participants")
    elif "participants" in validated:
        validated["number_of_participants"] = validated.pop("participants")
    elif "number_of_students" in validated:
        validated["number_of_participants"] = validated.pop("number_of_students")
        
    # Canonical: number_of_wins
    if "count_won" in validated:
        validated["number_of_wins"] = validated.pop("count_won")
    elif "wins" in validated:
        validated["number_of_wins"] = validated.pop("wins")

    # Canonical: event_name (for generic 'type' fields)
    if "type" in validated and "event_name" not in validated:
        validated["event_name"] = validated.get("type")
    
    # Canonical: event_name aliases
    if "event" in validated and "event_name" not in validated:
        validated["event_name"] = validated.pop("event")
    elif "activity_name" in validated and "event_name" not in validated:
        validated["event_name"] = validated.pop("activity_name")
    elif "activity" in validated and "event_name" not in validated:
        validated["event_name"] = validated.pop("activity")

    # Canonical: organizer aliases
    if "organizer_name" in validated and "organizer" not in validated:
        validated["organizer"] = validated.pop("organizer_name")
    elif "organization" in validated and "organizer" not in validated:
        validated["organizer"] = validated.pop("organization")

    # Ensure canonical fields exist with defaults if missing
    if "number_of_participants" not in validated:
        validated["number_of_participants"] = 0
    if "number_of_wins" not in validated:
        validated["number_of_wins"] = 0
        
    # --- STRICT SCHEMA ENFORCEMENT ---
    # Final cleanup of any lingering legacy keys to ensure clean DB storage
    validated.pop("activity_name", None)
    validated.pop("organization", None)
    validated.pop("organizer_name", None)
    validated.pop("participants", None)
    validated.pop("number_of_students", None)
    validated.pop("count_participants", None)

    return True, None, validated

# Helper functions for bulk upload (shared between Extracurricular and Placement)
def split_names(val):
    if val not in (None, ""):
        val_str = str(val)
        if val_str.strip():
            return [n.strip() for n in val_str.split(",") if n.strip()]
    return []

def get_first_value(d, keys):
    for key in keys:
        if key in d and d[key] not in (None, ""):
            return d[key]
    return None

def get_evidence_link(d):
    return get_first_value(d, [
        "evidence_link", "evidence link", 
        "proof_link_or_file", "proof link or file", 
        "report_or_photos", "report or photos",
        "link", "proof", "evidence"
    ]) or ""

def get_student_names(d):
    val = get_first_value(d, [
        "student_names", "student names",
        "name_of_students", "name of students",
        "names", "students",
        "student_name", "student name",
        "name of student", "name of student(s)",
        "student list", "list of students",
        "participants", "participant names"
    ])
    if val not in (None, ""):
        val_str = str(val)
        if val_str.strip():
                return [n.strip() for n in val_str.split(",") if n.strip()]
    return []

@app.route('/api/extracurricular/upload-bulk', methods=['POST', 'OPTIONS'])
@login_required
def upload_extracurricular_bulk():
    """Bulk upload extracurricular records via CSV/Excel with validation report."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    if 'file' not in request.files:
        return jsonify({"message": "No file part"}), 400
    
    file = request.files['file']
    batch_id = request.form.get('batch_id')
    output_type = normalize_output_type((request.form.get('output_type') or '').strip())
    
    if not batch_id:
        return jsonify({"message": "Batch is required"}), 400
        
    # Removed strict EXTRACURRICULAR_OUTPUT_TYPES check to allow free text
    
    # Validate batch_id and ownership
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
             return jsonify({"message": "User college or branch not found in session"}), 400

        # Try as ObjectId first, then as batch_range
        batch = None
        batch_obj_id = None
        
        if ObjectId.is_valid(batch_id):
            batch_obj_id = ObjectId(batch_id)
            batch = db.batches.find_one({"_id": batch_obj_id, "college": college, "branch": branch})
            
        if not batch:
            # Try finding by batch_range string
            batch = db.batches.find_one({"batch_range": batch_id, "college": college, "branch": branch})
            if batch:
                batch_obj_id = batch["_id"]

        if not batch:
            return jsonify({"message": f"Batch '{batch_id}' not found or unauthorized access"}), 400
            
    except Exception as e:
        return jsonify({"message": f"Error validating batch: {str(e)}"}), 400
    
    if file.filename == '':
        return jsonify({"message": "No file selected"}), 400
    
    try:
        # Check for duplicate file (Option B: Duplicate Upload Handling)
        existing_file = db.extracurricular_files.find_one({
            "batch_id": batch_obj_id,
            "filename": file.filename,
            "output_type": output_type
        })
        warning_msg = ""
        if existing_file:
             warning_msg = f" Warning: Data from a file with this name may already exist for {output_type}. Re-uploading can cause duplicate entries."

        # Read file content
        file_content = file.read()
        file_size = len(file_content)
        file.seek(0)
        
        # Store file in GridFS for tracking
        fs = GridFS(db)
        grid_file_id = fs.put(
            BytesIO(file_content),
            filename=file.filename,
            file_type='extracurricular_bulk',
            uploaded_at=datetime.now(timezone.utc),
            batch_id=batch_obj_id,
            college=college,
            branch=branch
        )
        
        # Reset file pointer for reading
        file.seek(0)
        
        # Determine file type and read
        filename_lower = file.filename.lower()
        if filename_lower.endswith('.csv'):
            df = pd.read_csv(BytesIO(file_content))
        elif filename_lower.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(BytesIO(file_content))
        else:
            return jsonify({"message": "Unsupported file format. Use CSV or Excel."}), 400
        
        # Normalize column names to lowercase, strip, and replace newlines with spaces
        df.columns = [str(c).replace('\n', ' ').strip().lower() for c in df.columns]
        rows_to_validate = []
        for idx, row in df.iterrows():
            row_dict = {k: (str(v).strip() if isinstance(v, str) else v) for k, v in row.to_dict().items()}
            row_dict = {k: v for k, v in row_dict.items() if pd.notna(v)}
            row_dict['output_type'] = output_type

            if output_type == "Sports":
                # Expected columns per template
                participants_val = get_first_value(
                    row_dict,
                    [
                        "number_of_participants",
                        "number of participants",
                        "number_of_students",
                        "number of students",
                        "count_participants",
                        "count participants"
                    ]
                )
                # Robust name mapping
                sports_name_val = get_first_value(
                    row_dict,
                    [
                        "sports_name", "sport_name", "sport", 
                        "activity_name", "event_name", "name",
                        "sport name", "sports name", "activity name", "event name"
                    ]
                )
                payload = {
                    "academic_year": row_dict.get("academic_year"),
                    "output_type": "Sports",
                    "sports_name": sports_name_val or "",
                    "level": row_dict.get("level", ""),
                    "number_of_participants": participants_val,
                    "number_of_wins": row_dict.get("number_of_wins", 0),
                    "student_names": get_student_names(row_dict),
                    "organizer": row_dict.get("organizer", ""),
                    "remarks": row_dict.get("remarks", ""),
                    "evidence_link": get_evidence_link(row_dict)
                }
                rows_to_validate.append(payload)
            elif output_type == "Cultural":
                participants_val = get_first_value(
                    row_dict,
                    [
                        "number_of_participants",
                        "number of participants",
                        "number_of_students",
                        "number of students",
                        "count_participants",
                        "count participants"
                    ]
                )
                event_category_val = get_first_value(
                    row_dict, 
                    [
                        "event_category", "event category", "category",
                        "activity_category", "activity category", "type"
                    ]
                )
                payload = {
                    "academic_year": row_dict.get("academic_year"),
                    "output_type": "Cultural",
                    "activity_name": row_dict.get("activity_name", ""),
                    "event_category": event_category_val or "",
                    "custom_category": row_dict.get("custom_category", ""),
                    "date": row_dict.get("date", ""),
                    "number_of_participants": participants_val,
                    "student_names": get_student_names(row_dict),
                    "organizer": row_dict.get("organizer", ""),
                    "remarks": row_dict.get("remarks", ""),
                    "evidence_link": get_evidence_link(row_dict)
                }
                rows_to_validate.append(payload)
            elif output_type == "Technical":
                participants_val = get_first_value(
                    row_dict,
                    [
                        "number_of_participants",
                        "number of participants",
                        "number_of_students",
                        "number of students",
                        "count_participants",
                        "count participants"
                    ]
                )
                event_name_val = get_first_value(
                    row_dict,
                    [
                        "event_name", "activity_name", "title", "name",
                        "competition_name", "hackathon_name"
                    ]
                )
                organization_val = get_first_value(
                    row_dict,
                    [
                        "organization", "organization_name", "organization name",
                        "organizer", "organizing_body", "organizing body",
                        "institute", "college"
                    ]
                )
                payload = {
                    "academic_year": row_dict.get("academic_year"),
                    "output_type": "Technical",
                    "event_name": event_name_val or "",
                    "event_type": row_dict.get("event_type", ""),
                    "organization": organization_val or "",
                    "organizer": organization_val or "", # Map to organizer for Report
                    "level": row_dict.get("level", ""),
                    "number_of_participants": participants_val,
                    "student_names": get_student_names(row_dict),
                    "achievement": row_dict.get("achievement", row_dict.get("outcome", "")),
                    "remarks": row_dict.get("remarks", ""),
                    "evidence_link": get_evidence_link(row_dict)
                }
                rows_to_validate.append(payload)
            elif output_type == "Internship":
                students_val = get_first_value(
                    row_dict,
                    [
                        "number_of_students",
                        "number of students",
                        "number_of_participants",
                        "number of participants",
                        "count_participants",
                        "count participants"
                    ]
                )
                payload = {
                    "academic_year": row_dict.get("academic_year"),
                    "output_type": "Internship",
                    "internship_title": get_first_value(row_dict, ["internship_title", "internship title", "internship_domain", "internship domain", "domain", "title"]),
                    "organization_name": get_first_value(row_dict, ["organization_name", "organization name", "company", "organization"]),
                    "mode": row_dict.get("mode", ""),
                    "duration": row_dict.get("duration", ""),
                    "number_of_students": students_val,
                    "student_names": get_student_names(row_dict),
                    "stipend": row_dict.get("stipend", ""),
                    "stipend_amount": get_first_value(row_dict, ["stipend_amount", "stipend amount", "amount"]),
                    "evidence_link": get_evidence_link(row_dict)
                }
                rows_to_validate.append(payload)
            elif output_type == "Courses":
                mode_val = row_dict.get("mode", "")
                students_val = get_first_value(
                    row_dict,
                    [
                        "number_of_students",
                        "number of students",
                        "number_of_participants",
                        "number of participants",
                        "count_participants",
                        "count participants"
                    ]
                )
                payload = {
                    "academic_year": row_dict.get("academic_year"),
                    "output_type": "Courses",
                    "mode": mode_val,
                    "course_name": get_first_value(row_dict, ["course_name", "course name", "course", "name"]),
                    "platform": row_dict.get("platform", ""),
                    "duration": row_dict.get("duration", ""),
                    "number_of_students": students_val,
                    "student_names": get_student_names(row_dict),
                    "certification_status": get_first_value(row_dict, ["certification_status", "certification status", "status"]),
                    "evidence_link": get_evidence_link(row_dict)
                }
                rows_to_validate.append(payload)
            elif output_type == "Industrial Visit":
                students_val = get_first_value(
                    row_dict,
                    [
                        "total_students",
                        "total students",
                        "number_of_students",
                        "number of students",
                        "number_of_participants",
                        "number of participants",
                        "count_participants",
                        "count participants"
                    ]
                )
                faculty_val = get_first_value(
                    row_dict,
                    [
                        "number_of_faculty",
                        "number of faculty",
                        "no_of_faculty",
                        "no of faculty",
                        "faculty_count",
                        "faculty count",
                        "faculty"  # Fallback if user uses old template but puts number
                    ]
                )
                payload = {
                    "academic_year": row_dict.get("academic_year"),
                    "output_type": "Industrial Visit",
                    "industrial_visit_details": get_first_value(row_dict, ["industrial_visit_details", "industrial visit details", "details", "visit details"]),
                    "date": row_dict.get("date", ""),
                    "number_of_students": students_val,
                    "number_of_faculty": faculty_val,
                    "evidence_link": get_evidence_link(row_dict)
                }
                rows_to_validate.append(payload)
            else:
                # Generic handler for custom types
                participants_val = get_first_value(
                    row_dict,
                    [
                        "number_of_participants",
                        "number of participants",
                        "number_of_students",
                        "number of students",
                        "count_participants",
                        "count participants",
                        "total_students",
                        "total students"
                    ]
                )
                payload = {
                    "academic_year": row_dict.get("academic_year"),
                    "output_type": output_type,
                    "number_of_participants": participants_val,
                    "student_names": get_student_names(row_dict),
                    "remarks": row_dict.get("remarks", ""),
                    "evidence_link": get_evidence_link(row_dict)
                }
                # Capture all other fields that are not standard
                standard_keys = ["academic_year", "output_type", "number_of_participants", "student_names", "remarks", "proof_link_or_file", "evidence_link"]
                for k, v in row_dict.items():
                    if k not in standard_keys and k not in ["number of participants", "number of students", "count_participants", "total_students", "count participants", "total students"]:
                        payload[k] = v
                        
                rows_to_validate.append(payload)
        
        # Validation report
        valid_records = []
        invalid_records = []
        
        for idx, row_dict in enumerate(rows_to_validate):
            is_valid, error_msg, validated_data = validate_extracurricular_record(row_dict, is_bulk=True)
            
            if is_valid:
                valid_records.append({
                    "row": idx + 2,  # +2 because 0-indexed and header row
                    "data": validated_data
                })
            else:
                invalid_records.append({
                    "row": idx + 2,
                    "data": row_dict,
                    "error": error_msg
                })
        
        # Return validation report (preview before commit)
        return jsonify({
            "message": f"Validation complete. {len(valid_records)} valid, {len(invalid_records)} invalid records." + warning_msg,
            "warning": warning_msg.strip(),
            "valid_count": len(valid_records),
            "invalid_count": len(invalid_records),
            "valid_records": valid_records,  # Return all valid records for commit
            "valid_records_preview": valid_records[:10],  # Preview first 10 for display
            "invalid_records": invalid_records,
            "errors": invalid_records,
            "file_id": str(grid_file_id),  # Return file_id for tracking
            "filename": file.filename,  # Return filename for tracking
            "batch_id": batch_id,  # Return batch_id for commit
            "preview": True
        }), 200
        
    except Exception as e:
        return jsonify({"message": f"Error processing file: {str(e)}"}), 500


@app.route('/api/extracurricular/template', methods=['GET', 'OPTIONS'])
def download_extracurricular_template():
    """Download the bulk upload template for a specific Extracurricular Output Type."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    try:
        output_type = (request.args.get('output_type') or '').strip()
        output_type = normalize_output_type(output_type)
        if output_type not in EXTRACURRICULAR_OUTPUT_TYPES:
            return jsonify({"message": f"Invalid or missing output_type. Use one of: {', '.join(EXTRACURRICULAR_OUTPUT_TYPES)}"}), 400

        # Define per-type columns (1-to-1 with single-entry form fields)
        templates = {
            "Sports": [
                "academic_year", "sports_name", "level",
                "number_of_participants", "number_of_wins",
                "student_names", "organizer", "remarks", "evidence_link"
            ],
            "Cultural": [
                "academic_year", "activity_name", "category",
                "number_of_participants", "student_names",
                "organizer", "remarks", "evidence_link"
            ],
            "Technical": [
                "academic_year", "event_name", "event_type", "organization",
                "level", "number_of_participants", "student_names",
                "achievement", "remarks", "evidence_link"
            ],
            "Internship": [
                "academic_year", "internship_title", "organization_name",
                "mode", "duration", "number_of_students", "student_names",
                "stipend", "stipend_amount", "evidence_link"
            ],
            "Courses": [
                "academic_year", "mode", "course_name", "platform", "duration",
                "number_of_students", "student_names",
                "evidence_link"
            ],
            "Industrial Visit": [
                "academic_year", "industrial_visit_details", "date",
                "total_students", "number_of_faculty", "evidence_link"
            ],
        }

        columns = templates[output_type]
        df = pd.DataFrame(columns=columns)

        output = BytesIO()
        writer = pd.ExcelWriter(output, engine='xlsxwriter')
        df.to_excel(writer, index=False, sheet_name='Template')

        workbook = writer.book
        inst_sheet = workbook.add_worksheet('Instructions')
        inst_sheet.write(0, 0, f'Extracurricular Bulk Template - {output_type}')
        inst_sheet.write(2, 0, 'General Rules:')
        inst_sheet.write(3, 0, '1) Use separate template per Output Type.')
        inst_sheet.write(4, 0, '2) Names fields are optional; comma-separated if provided.')
        inst_sheet.write(5, 0, '3) Only academic_year and participant counts are required (per spec).')
        # Removed fixed value instructions for Cultural/Other categories to allow free text
        writer.close()
        output.seek(0)

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'Extracurricular_{output_type}_Template.xlsx'
        )
    except Exception as e:
        return jsonify({"message": f"Error generating template: {str(e)}"}), 500

@app.route('/api/get-extracurricular-files', methods=['GET', 'OPTIONS'])
@login_required
def get_extracurricular_files():
    """Get list of uploaded extracurricular files (both bulk and individual uploads)."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    batch_id = request.args.get('batch_id', '')
    college = session.get('college')
    branch = session.get('branch')
    if not college or not branch:
        return jsonify({"message": "User college or branch not found"}), 400
    
    try:
        files_list = []
        
        # Build query for bulk files
        bulk_query = {}
        branch_batch_ids = []

        if batch_id:
            try:
                batch_obj_id = ObjectId(batch_id)
                # Verify ownership
                batch = db.batches.find_one({"_id": batch_obj_id, "college": college, "branch": branch})
                if not batch:
                    return jsonify({"message": "Invalid batch ID or unauthorized access"}), 400
                bulk_query['batch_id'] = batch_obj_id
            except Exception:
                return jsonify({"message": "Invalid batch ID format"}), 400
        else:
            # Find all batches for this branch
            branch_batches = list(db.batches.find({"college": college, "branch": branch}, {"_id": 1}))
            branch_batch_ids = [b["_id"] for b in branch_batches]
            bulk_query['batch_id'] = {"$in": branch_batch_ids}
        # Exclude soft-deleted bulk files
        bulk_query['deleted'] = {"$ne": True}
        bulk_query['pending_delete'] = {"$ne": True}
        
        # Get bulk upload files from extracurricular_files collection
        bulk_files = list(db.extracurricular_files.find(bulk_query).sort('uploaded_at', -1))
        for file in bulk_files:
            batch_range = 'N/A'
            if file.get('batch_id'):
                batch = db.batches.find_one({"_id": file['batch_id']})
                if batch:
                    batch_range = batch.get('batch_range', 'N/A')
            
            files_list.append({
                "_id": str(file["_id"]),
                "file_id": str(file.get("file_id", "")),
                "filename": file.get("filename", "Unknown"),
                "uploaded_at": file.get("uploaded_at", datetime.now(timezone.utc)).isoformat() if isinstance(file.get("uploaded_at"), datetime) else str(file.get("uploaded_at", "")),
                "record_count": file.get("record_count", 0),
                "file_type": "bulk",
                "batch_range": batch_range
            })
        
        # Get individual record files from extracurricular_records collection
        # Find records that have evidence_file_ids
        records_query = {
            "$or": [
                {"evidence_file_ids": {"$exists": True, "$ne": []}},
                {"evidence_file_id": {"$exists": True, "$ne": None}}
            ]
        }
        
        # Filter by batch_id if provided
        if batch_id:
            records_query['batch_id'] = ObjectId(batch_id)
        else:
            records_query['batch_id'] = {"$in": branch_batch_ids}
        
        # Exclude soft-deleted records
        records_query["deleted"] = {"$ne": True}
        records_query["pending_delete"] = {"$ne": True}
        records_with_files = list(db.extracurricular_records.find(records_query).sort('created_at', -1))
        
        fs = GridFS(db)
        seen_file_ids = set()
        
        for record in records_with_files:
            file_ids = []
            if "evidence_file_ids" in record and record["evidence_file_ids"]:
                file_ids = record["evidence_file_ids"]
            elif "evidence_file_id" in record and record["evidence_file_id"]:
                file_ids = [record["evidence_file_id"]]
            
            for file_id_str in file_ids:
                try:
                    file_id = ObjectId(file_id_str)
                    if str(file_id) in seen_file_ids:
                        continue
                    seen_file_ids.add(str(file_id))
                    
                    # Try to get file info from GridFS
                    try:
                        grid_file = fs.get(file_id)
                        # Skip evidence files marked deleted in GridFS metadata
                        fdoc = db.fs.files.find_one({"_id": file_id}) or {}
                        if fdoc.get("deleted") is True:
                            continue
                        
                        # Get batch_range for individual records
                        batch_range = 'N/A'
                        if record.get('batch_id'):
                            batch = db.batches.find_one({"_id": record['batch_id']})
                            if batch:
                                batch_range = batch.get('batch_range', 'N/A')
                        
                        files_list.append({
                            "_id": str(record["_id"]),  # Use record ID as identifier
                            "file_id": str(file_id),
                            "filename": grid_file.filename or f"Evidence_{str(file_id)[:8]}",
                            "uploaded_at": record.get("created_at", datetime.now(timezone.utc)).isoformat() if isinstance(record.get("created_at"), datetime) else str(record.get("created_at", "")),
                            "record_count": 1,
                            "file_type": "individual",
                            "output_type": record.get("output_type", "Legacy"),
                            "academic_year": record.get("academic_year", ""),
                            "batch_range": batch_range
                        })
                    except Exception as e:
                        print(f"⚠️ Could not retrieve GridFS file {file_id}: {e}")
                        continue
                except Exception as e:
                    print(f"⚠️ Invalid file ID {file_id_str}: {e}")
                    continue
        
        # Sort by uploaded_at descending
        files_list.sort(key=lambda x: x.get("uploaded_at", ""), reverse=True)
        
        return jsonify(files_list), 200
    except Exception as e:
        return jsonify({"message": f"Error retrieving extracurricular files: {str(e)}"}), 500

@app.route('/api/download-extracurricular-file/<file_id>', methods=['GET', 'OPTIONS'])
def download_extracurricular_file(file_id):
    """Download an extracurricular bulk upload file."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    # Allow public access via ID for report links
    try:
        fs = GridFS(db)
        file_obj_id = ObjectId(file_id)
        
        if not fs.exists(file_obj_id):
             return jsonify({"message": "File not found"}), 404
             
        file_doc = fs.get(file_obj_id)
        
        from flask import send_file
        return send_file(
            file_doc,
            download_name=file_doc.filename,
            as_attachment=True
        )
    except Exception as e:
        return jsonify({"message": f"Error downloading file: {str(e)}"}), 500

@app.route('/api/delete-extracurricular-file/<file_id>', methods=['DELETE', 'OPTIONS'])
def delete_extracurricular_file(file_id):
    """Soft delete or finalize delete an extracurricular file (bulk or individual uploads)."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    college = session.get('college')
    branch = session.get('branch')
    if not college or not branch:
        return jsonify({"message": "User college or branch not found"}), 400

    try:
        file_type = request.args.get('file_type', 'bulk')
        gridfs_file_id = request.args.get('file_id')
        finalize = (request.args.get('finalize') or '').lower() == 'true'
        
        if file_type == 'individual':
            record = db.extracurricular_records.find_one({"_id": ObjectId(file_id)})
            if not record:
                return jsonify({"message": "Record not found"}), 404
            
            # Verify ownership
            batch = db.batches.find_one({"_id": record['batch_id']})
            if not batch or batch.get('college') != college or batch.get('branch') != branch:
                 return jsonify({"message": "Unauthorized access"}), 403
            
            if not finalize:
                if gridfs_file_id:
                    try:
                        db.fs.files.update_one(
                            {"_id": ObjectId(gridfs_file_id)},
                            {"$set": {
                                "pending_delete": True,
                                "delete_requested_at": datetime.now(timezone.utc)
                            }}
                        )
                    except Exception as e:
                        print(f"⚠️ Warning: Could not mark GridFS file for deletion: {e}")
                return jsonify({"message": "Evidence file marked for deletion", "pending_delete": True}), 200
            else:
                if gridfs_file_id:
                    try:
                        fs = GridFS(db)
                        fs.delete(ObjectId(gridfs_file_id))
                    except Exception as e:
                        print(f"⚠️ Warning: Could not delete GridFS file: {e}")
                if "evidence_file_ids" in record and record["evidence_file_ids"]:
                    updated_ids = [fid for fid in record["evidence_file_ids"] if fid != gridfs_file_id]
                    db.extracurricular_records.update_one(
                        {"_id": ObjectId(file_id)},
                        {"$set": {"evidence_file_ids": updated_ids}}
                    )
                elif "evidence_file_id" in record and record["evidence_file_id"] == gridfs_file_id:
                    db.extracurricular_records.update_one(
                        {"_id": ObjectId(file_id)},
                        {"$unset": {"evidence_file_id": ""}}
                    )
                return jsonify({"message": "Evidence file permanently deleted", "finalized": True}), 200
        else:
            file_meta = db.extracurricular_files.find_one({"_id": ObjectId(file_id)})
            if not file_meta:
                return jsonify({"message": "File not found"}), 404
            
            # Verify ownership
            batch = db.batches.find_one({"_id": file_meta['batch_id']})
            if not batch or batch.get('college') != college or batch.get('branch') != branch:
                 return jsonify({"message": "Unauthorized access"}), 403
            
            if not finalize:
                db.extracurricular_files.update_one(
                    {"_id": ObjectId(file_id)},
                    {"$set": {
                        "pending_delete": True,
                        "delete_requested_at": datetime.now(timezone.utc)
                    }}
                )
                return jsonify({"message": "Extracurricular file marked for deletion", "pending_delete": True}), 200
            else:
                if file_meta.get("file_id"):
                    try:
                        fs = GridFS(db)
                        fs.delete(file_meta["file_id"])
                    except Exception as e:
                        print(f"⚠️ Warning: Could not delete GridFS file: {e}")
                result = db.extracurricular_files.delete_one({"_id": ObjectId(file_id)})
                if result.deleted_count == 0:
                    return jsonify({"message": "File not found"}), 404
                return jsonify({"message": "Extracurricular file permanently deleted", "finalized": True}), 200
        
    except Exception as e:
        return jsonify({"message": f"Error deleting file: {str(e)}"}), 500


@app.route('/api/restore-extracurricular-file/<file_id>', methods=['POST', 'OPTIONS'])
@login_required
def restore_extracurricular_file(file_id):
    """Restore a soft-deleted extracurricular file (bulk or evidence)."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    college = session.get('college')
    branch = session.get('branch')
    if not college or not branch:
        return jsonify({"message": "User college or branch not found"}), 400
    
    file_type = request.args.get('type', 'bulk') # 'bulk' or 'individual'
    
    try:
        if file_type == 'individual':
            # Evidence file stored in GridFS (fs.files)
            gridfs_file_id = request.args.get('file_id')
            if not gridfs_file_id:
                return jsonify({"message": "GridFS file ID required for individual file restore"}), 400
            
            # Verify ownership via record
            record = db.extracurricular_records.find_one({"_id": ObjectId(file_id)})
            if not record:
                return jsonify({"message": "Record not found"}), 404
            
            batch = db.batches.find_one({"_id": record['batch_id']})
            if not batch or batch.get('college') != college or batch.get('branch') != branch:
                 return jsonify({"message": "Unauthorized access"}), 403
            
            db.fs.files.update_one(
                {"_id": ObjectId(gridfs_file_id)},
                {"$unset": {"pending_delete": "", "delete_requested_at": "", "deleted": ""}}
            )
            return jsonify({"message": "Evidence file restored successfully"}), 200
            
        else:
            # Bulk upload file record
            file_meta = db.extracurricular_files.find_one({"_id": ObjectId(file_id)})
            if not file_meta:
                return jsonify({"message": "File not found"}), 404
            
            batch = db.batches.find_one({"_id": file_meta['batch_id']})
            if not batch or batch.get('college') != college or batch.get('branch') != branch:
                 return jsonify({"message": "Unauthorized access"}), 403
            
            db.extracurricular_files.update_one(
                {"_id": ObjectId(file_id)},
                {"$unset": {"pending_delete": "", "delete_requested_at": "", "deleted": ""}}
            )
            return jsonify({"message": "Extracurricular file restored successfully"}), 200
            
    except Exception as e:
        return jsonify({"message": f"Error restoring file: {str(e)}"}), 500

@app.route('/api/extracurricular/upload-commit', methods=['POST', 'OPTIONS'])
@login_required
def commit_extracurricular_bulk():
    """Commit validated records from bulk upload."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    data = request.get_json()
    # Support both formats: direct records array or wrapped in valid_records
    # If 'records' is provided, it might contain wrappers from valid_records
    raw_records = data.get('records', [])
    if not raw_records and 'valid_records' in data:
        raw_records = data.get('valid_records', [])
    
    records = []
    for item in raw_records:
        if isinstance(item, dict) and 'data' in item and isinstance(item['data'], dict):
            records.append(item['data'])
        else:
            records.append(item)
    
    if not records:
        return jsonify({"message": "No records to commit"}), 400
    
    # Get batch_id from request
    batch_id = data.get('batch_id')
    if not batch_id:
        return jsonify({"message": "Batch is required"}), 400
    
    # Validate batch_id and ownership
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
             return jsonify({"message": "User college or branch not found in session"}), 400

        batch = None
        batch_obj_id = None
        
        if ObjectId.is_valid(batch_id):
            batch_obj_id = ObjectId(batch_id)
            batch = db.batches.find_one({"_id": batch_obj_id, "college": college, "branch": branch})
        
        if not batch:
            # Try finding by batch_range string
            batch = db.batches.find_one({"batch_range": batch_id, "college": college, "branch": branch})
            if batch:
                batch_obj_id = batch["_id"]

        if not batch:
            return jsonify({"message": "Invalid batch ID or unauthorized access"}), 400
    except Exception:
        return jsonify({"message": "Invalid batch ID format"}), 400
    
    try:
        fs = GridFS(db)
        inserted_count = 0
        errors = []
        
        # Store file metadata if file_id is provided
        file_id = data.get('file_id')
        filename = data.get('filename', 'extracurricular_bulk.csv')
        if file_id:
            # Verify file exists in GridFS
            try:
                grid_file = fs.get(ObjectId(file_id))
                file_size = grid_file.length
                
                # Store file metadata in collection
                db.extracurricular_files.insert_one({
                    "file_id": ObjectId(file_id),
                    "filename": filename,
                    "batch_id": batch_obj_id,
                    "college": college,
                    "branch": branch,
                    "uploaded_at": datetime.now(timezone.utc),
                    "file_size": file_size,
                    "record_count": len(records)
                })
            except Exception as e:
                print(f"⚠️ Warning: Could not track file metadata: {e}")
        
        for record_data in records:
            # Re-validate before insert
            is_valid, error_msg, validated_data = validate_extracurricular_record(record_data, is_bulk=False)
            
            if not is_valid:
                errors.append({"record": record_data, "error": error_msg})
                continue
            
            # Handle evidence file if present
            evidence_file_id = None
            if "evidence_file" in record_data and record_data["evidence_file"]:
                # If evidence_file is a file object or base64, store in GridFS
                # For now, we'll store the reference as-is (assuming it's already a file_id or URL)
                evidence_file_id = record_data["evidence_file"]
            
            # Create document
            doc = {
                "academic_year": validated_data["academic_year"],
                "batch_id": batch_obj_id,
                "college": college,
                "branch": branch,
                "created_at": datetime.now(timezone.utc),
                "verified": False,
                "approved": False
            }
            
            # Legacy structure
            if "category" in validated_data and "output_type" not in validated_data:
                doc.update({
                    "category": validated_data["category"],
                    "event_date": validated_data.get("event_date"),
                    "level": validated_data.get("level"),
                    "type": validated_data.get("type"),
                    "count_participants": validated_data.get("count_participants"),
                    "count_won": validated_data.get("count_won", 0),
                    "organizer": validated_data.get("organizer"),
                    "remarks": validated_data.get("remarks"),
                    "evidence_file_id": evidence_file_id
                })
            # New structure
            elif "output_type" in validated_data:
                doc["output_type"] = validated_data["output_type"]
                # Store all validated fields
                for key, value in validated_data.items():
                    if key not in ["academic_year", "output_type"]:
                        doc[key] = value

                if doc["output_type"] == "Cultural":
                    participants = doc.get("number_of_participants")
                    if participants is not None and "number_of_students" not in doc:
                        doc["number_of_students"] = participants
                elif doc["output_type"] == "Technical":
                    participants = doc.get("number_of_participants")
                    if participants is not None and "number_of_students" not in doc:
                        doc["number_of_students"] = participants
                elif doc["output_type"] == "Sports":
                    participants = doc.get("number_of_participants", doc.get("number_of_students"))
                    if participants is not None:
                        doc["number_of_participants"] = participants

                # Bulk upload currently doesn't support individual evidence files
                doc["evidence_file_ids"] = []
                if evidence_file_id:
                     doc["evidence_file_ids"].append(evidence_file_id)
            
            db.extracurricular_records.insert_one(doc)
            inserted_count += 1
        
        return jsonify({
            "message": f"Successfully inserted {inserted_count} records",
            "inserted_count": inserted_count,
            "count": inserted_count,
            "errors": errors
        }), 200
        
    except Exception as e:
        return jsonify({"message": f"Error committing records: {str(e)}"}), 500


@app.route('/api/get-placement-files', methods=['GET', 'OPTIONS'])
@login_required
def get_placement_files():
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200

    batch_id = request.args.get('batch_id', '')
    college = session.get('college')
    branch = session.get('branch')
    if not college or not branch:
        return jsonify({"message": "User college or branch not found"}), 400

    try:
        query = {}
        if batch_id:
            try:
                batch_obj_id = ObjectId(batch_id)
                # Verify ownership
                batch = db.batches.find_one({"_id": batch_obj_id, "college": college, "branch": branch})
                if not batch:
                    return jsonify({"message": "Invalid batch ID or unauthorized access"}), 400
                query['batch_id'] = batch_obj_id
            except Exception:
                return jsonify({"message": "Invalid batch ID format"}), 400
        else:
            # Filter by all batches for this college and branch
            branch_batches = list(db.batches.find({"college": college, "branch": branch}, {"_id": 1}))
            branch_batch_ids = [b["_id"] for b in branch_batches]
            query['batch_id'] = {"$in": branch_batch_ids}

        # Exclude soft-deleted files
        query['deleted'] = {"$ne": True}
        query['pending_delete'] = {"$ne": True}
        files = list(db.placement_outcome_files.find(query).sort('uploaded_at', -1))

        for file in files:
            batch_range = 'N/A'
            if file.get('batch_id'):
                batch = db.batches.find_one({"_id": file['batch_id']})
                if batch:
                    batch_range = batch.get('batch_range', 'N/A')
            file['batch_range'] = batch_range

        files = convert_objectid_to_str(files)

        return jsonify(files), 200
    except Exception as e:
        return jsonify({"message": f"Error retrieving placement files: {str(e)}"}), 500


@app.route('/api/download-placement-file/<file_id>', methods=['GET', 'OPTIONS'])
def download_placement_file(file_id):
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200

    # Allow public access via ID for report links
    try:
        fs = GridFS(db)
        
        # Strategy 1: Check if it's a bulk upload file record (placement_outcome_files)
        # This is used by the frontend "Uploaded Files" list
        file_doc = db.placement_outcome_files.find_one({"_id": ObjectId(file_id)})
        
        if file_doc:
            # It's a bulk upload record
            grid_id = file_doc.get('file_id')
            if not grid_id:
                return jsonify({"message": "File data not available"}), 404
            
            grid_file = fs.get(grid_id)
            filename = file_doc.get("filename", "placement_outcomes.xlsx")
        else:
            # Strategy 2: Check if it's a direct GridFS file ID (evidence file)
            # This is used by the generated DOCX report links
            try:
                grid_id = ObjectId(file_id)
                if fs.exists(grid_id):
                    grid_file = fs.get(grid_id)
                    filename = grid_file.filename or "placement_evidence.pdf"
                else:
                     return jsonify({"message": "File not found"}), 404
            except Exception:
                 return jsonify({"message": "File not found"}), 404

        if not grid_file:
            return jsonify({"message": "File data not found"}), 404

        from flask import Response
        # Determine mimetype based on filename extension
        import mimetypes
        mimetype, _ = mimetypes.guess_type(filename)
        if not mimetype:
            mimetype = 'application/octet-stream'
            
        return Response(
            grid_file.read(),
            mimetype=mimetype,
            headers={'Content-Disposition': f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        return jsonify({"message": f"Error downloading placement file: {str(e)}"}), 500


@app.route('/api/delete-placement-file/<file_id>', methods=['DELETE', 'OPTIONS'])
def delete_placement_file(file_id):
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200

    college = session.get('college')
    branch = session.get('branch')
    if not college or not branch:
        return jsonify({"message": "User college or branch not found"}), 400

    try:
        fs = GridFS(db)
        file_doc = db.placement_outcome_files.find_one({"_id": ObjectId(file_id)})
        if not file_doc:
            return jsonify({"message": "File not found"}), 404

        # Verify ownership
        batch = db.batches.find_one({"_id": file_doc.get('batch_id')})
        if not batch or batch.get('college') != college or batch.get('branch') != branch:
             return jsonify({"message": "Unauthorized access"}), 403

        finalize = (request.args.get('finalize') or '').lower() == 'true'
        if not finalize:
            db.placement_outcome_files.update_one(
                {"_id": ObjectId(file_id)},
                {"$set": {
                    "pending_delete": True,
                    "delete_requested_at": datetime.now(timezone.utc)
                }}
            )
            return jsonify({"message": "Placement file marked for deletion", "pending_delete": True}), 200
        
        grid_id = file_doc.get('file_id')
        if grid_id:
            try:
                fs.delete(grid_id)
            except Exception:
                pass
        
        db.placement_outcome_files.delete_one({"_id": ObjectId(file_id)})
        
        return jsonify({"message": "Placement file permanently deleted"}), 200
    except Exception as e:
        return jsonify({"message": f"Error deleting placement file: {str(e)}"}), 500

@app.route('/api/restore-placement-file/<file_id>', methods=['POST', 'OPTIONS'])
@login_required
def restore_placement_file(file_id):
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    try:
        college = session.get('college')
        branch = session.get('branch')
        if not college or not branch:
            return jsonify({"message": "User college or branch not found"}), 400
        
        file_doc = db.placement_outcome_files.find_one({"_id": ObjectId(file_id)})
        if not file_doc:
            return jsonify({"message": "File not found"}), 404
        
        batch = db.batches.find_one({"_id": file_doc.get('batch_id')})
        if not batch or batch.get('college') != college or batch.get('branch') != branch:
             return jsonify({"message": "Unauthorized access"}), 403
        
        db.placement_outcome_files.update_one(
            {"_id": ObjectId(file_id)},
            {"$unset": {"pending_delete": "", "delete_requested_at": "", "deleted": ""}}
        )
        return jsonify({"message": "Placement file restored successfully"}), 200
    except Exception as e:
        return jsonify({"message": f"Error restoring placement file: {str(e)}"}), 500
@app.route('/api/extracurricular', methods=['POST', 'OPTIONS'])
@login_required
def add_extracurricular_event():
    """Add a single extracurricular event (supports both old and new structure)."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    # Get batch_id from request
    batch_id = None
    if request.content_type and 'application/json' in request.content_type:
        data = request.get_json() or {}
        batch_id = data.get('batch_id')
    else:
        batch_id = request.form.get('batch_id')
    
    if not batch_id:
        return jsonify({"message": "Batch is required"}), 400
    
    # Validate batch_id and ownership
    try:
        college = session.get('college')
        branch = session.get('branch')
        if not college or not branch:
            missing = []
            if not college: missing.append("college")
            if not branch: missing.append("branch")
            return jsonify({"message": f"User profile incomplete: missing {', '.join(missing)}. Please logout and login again."}), 400

        # Try as ObjectId first, then as batch_range
        batch = None
        batch_obj_id = None
        
        if ObjectId.is_valid(batch_id):
            batch_obj_id = ObjectId(batch_id)
            batch = db.batches.find_one({"_id": batch_obj_id, "college": college, "branch": branch})
        
        if not batch:
            # Try finding by batch_range string
            batch = db.batches.find_one({"batch_range": batch_id, "college": college, "branch": branch})
            if batch:
                batch_obj_id = batch["_id"]

        if not batch:
            return jsonify({"message": f"Batch '{batch_id}' not found or unauthorized access"}), 400
            
    except Exception as e:
        return jsonify({"message": f"Error validating batch: {str(e)}"}), 400
    
    # Handle both JSON and form-data
    if request.content_type and 'application/json' in request.content_type:
        data = request.get_json() or {}
    else:
        # Form data - handle both old and new structure
        data = {}
        # Common fields
        if request.form.get('academic_year'):
            data["academic_year"] = request.form.get('academic_year')
        if request.form.get('output_type'):
            data["output_type"] = request.form.get('output_type')
        if request.form.get('student_names'):
            data["student_names"] = request.form.get('student_names')
        
        # Legacy fields
        if request.form.get('event_date'):
            data["event_date"] = request.form.get('event_date')
        if request.form.get('category'):
            data["category"] = request.form.get('category')
        if request.form.get('level'):
            data["level"] = request.form.get('level')
        if request.form.get('type'):
            data["type"] = request.form.get('type')
        if request.form.get('count_participants'):
            data["count_participants"] = request.form.get('count_participants')
        if request.form.get('count_won'):
            data["count_won"] = request.form.get('count_won')
        if request.form.get('organizer'):
            data["organizer"] = request.form.get('organizer')
        if request.form.get('remarks'):
            data["remarks"] = request.form.get('remarks')
        
        # New structure fields - Sports
        if request.form.get('sport_name'):
            data["sports_name"] = request.form.get('sport_name')
        elif request.form.get('sports_name'):
            data["sports_name"] = request.form.get('sports_name')
        if request.form.get('event_name'):
            data["event_name"] = request.form.get('event_name')
        if request.form.get('number_of_participants'):
            data["number_of_participants"] = request.form.get('number_of_participants')
        if request.form.get('number_of_wins'):
            data["number_of_wins"] = request.form.get('number_of_wins')
        
        # Cultural
        if request.form.get('event_category'):
            data["event_category"] = request.form.get('event_category')
        if request.form.get('cultural_custom_category'):
            data["cultural_custom_category"] = request.form.get('cultural_custom_category')
        if request.form.get('number_of_students'):
            data["number_of_students"] = request.form.get('number_of_students')
        
        # Technical
        if request.form.get('event_type'):
            data["event_type"] = request.form.get('event_type')
        if request.form.get('achievement'):
            data["achievement"] = request.form.get('achievement')
        
        # Internship
        if request.form.get('internship_domain'):
            data["internship_domain"] = request.form.get('internship_domain')
        if request.form.get('organization_name'):
            data["organization_name"] = request.form.get('organization_name')
        if request.form.get('mode'):
            data["mode"] = request.form.get('mode')
        if request.form.get('duration'):
            data["duration"] = request.form.get('duration')
        if request.form.get('has_stipend'):
            data["has_stipend"] = request.form.get('has_stipend')
        if request.form.get('stipend_amount'):
            data["stipend_amount"] = request.form.get('stipend_amount')
        if request.form.get('number_of_students'):
            data["number_of_students"] = request.form.get('number_of_students')
        
        # Courses
        if request.form.get('platform'):
            data["platform"] = request.form.get('platform')
        if request.form.get('course_name'):
            data["course_name"] = request.form.get('course_name')
        if request.form.get('number_of_students'):
            data["number_of_students"] = request.form.get('number_of_students')
        
        # Industrial Visit
        if request.form.get('industrial_visit_details'):
            data["industrial_visit_details"] = request.form.get('industrial_visit_details')
        if request.form.get('iv_date'):
            data["date"] = request.form.get('iv_date')
        elif request.form.get('date'):
            data["date"] = request.form.get('date')
        if request.form.get('total_students'):
            data["number_of_students"] = request.form.get('total_students')
        if request.form.get('number_of_faculty'):
            data["number_of_faculty"] = request.form.get('number_of_faculty')
        elif request.form.get('faculty'):
            data["number_of_faculty"] = request.form.get('faculty')
    
    # Validate
    is_valid, error_msg, validated_data = validate_extracurricular_record(data, is_bulk=False)
    
    if not is_valid:
        return jsonify({"message": error_msg}), 400
    
    try:
        # Handle multiple evidence files if uploaded
        evidence_file_ids = []
        if 'evidence_files' in request.files:
            files = request.files.getlist('evidence_files')
            fs = GridFS(db)
            for evidence_file in files:
                if evidence_file.filename:
                    file_id = fs.put(
                        evidence_file,
                        filename=evidence_file.filename,
                        file_type='extracurricular_evidence',
                        uploaded_at=datetime.now(timezone.utc)
                    )
                    evidence_file_ids.append(str(file_id))
        elif 'evidence_file' in request.files:
            # Legacy single file support
            evidence_file = request.files['evidence_file']
            if evidence_file.filename:
                fs = GridFS(db)
                file_id = fs.put(
                    evidence_file,
                    filename=evidence_file.filename,
                    file_type='extracurricular_evidence',
                    uploaded_at=datetime.now(timezone.utc)
                )
                evidence_file_ids.append(str(file_id))
        elif data.get('evidence_file_id'):
            # Legacy single file ID
            evidence_file_ids = [data.get('evidence_file_id')]
        
        # Create document - handle both old and new structure
        doc = {
            "academic_year": validated_data["academic_year"],
            "batch_id": batch_obj_id,
            "college": college,
            "branch": branch,  # Add branch field
            "created_at": datetime.now(timezone.utc),
            "verified": False,
            "approved": False
        }
        
        # Legacy structure
        if "category" in validated_data:
            doc.update({
                "category": validated_data["category"],
                "event_date": validated_data.get("event_date"),
                "level": validated_data.get("level"),
                "type": validated_data.get("type"),
                "count_participants": validated_data.get("count_participants"),
                "count_won": validated_data.get("count_won", 0),
                "organizer": validated_data.get("organizer"),
                "remarks": validated_data.get("remarks"),
                "evidence_file_id": evidence_file_ids[0] if evidence_file_ids else None
            })
        # New structure
        elif "output_type" in validated_data:
            doc["output_type"] = validated_data["output_type"]
            # Store all validated fields
            for key, value in validated_data.items():
                if key not in ["academic_year", "output_type"]:
                    doc[key] = value
            doc["evidence_file_ids"] = evidence_file_ids if evidence_file_ids else []
            
            # Ensure number_of_students is populated for types that prefer it
            if doc["output_type"] in ["Internship", "Courses", "Industrial Visit", "Cultural", "Technical"]:
                if "number_of_participants" in doc and "number_of_students" not in doc:
                    doc["number_of_students"] = doc["number_of_participants"]
        
        result = db.extracurricular_records.insert_one(doc)
        
        return jsonify({
            "message": "Record added successfully",
            "id": str(result.inserted_id)
        }), 201
        
    except Exception as e:
        return jsonify({"message": f"Error adding record: {str(e)}"}), 500

@app.route('/api/extracurricular/records-list', methods=['GET', 'OPTIONS'])
@login_required
def get_extracurricular_records_list():
    """Get extracurricular records list for Recent Records table."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    output_type = request.args.get('output_type')
    academic_year = request.args.get('academic_year')
    batch_id = request.args.get('batch_id')
    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 50))
    
    college = session.get('college')
    branch = session.get('branch')
    if not college or not branch:
        return jsonify({"message": "User college or branch not found"}), 400

    # Run cleanup for this user's context
    cleanup_pending_deletes(college, branch)

    try:
        query = {}
        if output_type:
            query["output_type"] = output_type
        if academic_year:
            query["academic_year"] = academic_year
        if batch_id:
            try:
                batch_obj_id = ObjectId(batch_id)
                # Verify ownership
                batch = db.batches.find_one({"_id": batch_obj_id, "college": college, "branch": branch})
                if not batch:
                     # Instead of 400, maybe just return empty or error? 
                     # Let's return error to be explicit
                     return jsonify({"message": "Invalid batch ID or unauthorized access"}), 400
                query["batch_id"] = batch_obj_id
            except Exception:
                pass
        else:
            # Filter by all batches for this college and branch
            branch_batches = list(db.batches.find({"college": college, "branch": branch}, {"_id": 1}))
            branch_batch_ids = [b["_id"] for b in branch_batches]
            query['batch_id'] = {"$in": branch_batch_ids}
        
        # Exclude soft-deleted records
        query['deleted'] = {"$ne": True}
        query['pending_delete'] = {"$ne": True}
        query['pending_delete'] = {"$ne": True}
        query['pending_delete'] = {"$ne": True}
        query['pending_delete'] = {"$ne": True}
        
        skip = (page - 1) * per_page
        
        # Get total count
        total = db.extracurricular_records.count_documents(query)
        
        # Get records
        records = list(
            db.extracurricular_records
            .find(query)
            .sort("created_at", -1)
            .skip(skip)
            .limit(per_page)
        )
        
        # Format records
        formatted_records = []
        for doc in records:
            # Get batch information
            batch_range = "N/A"
            if doc.get("batch_id"):
                try:
                    batch = db.batches.find_one({"_id": doc["batch_id"]})
                    if batch:
                        batch_range = batch.get("batch_range", "N/A")
                except Exception:
                    pass
            
            # Normalize output_type for display (backward compatibility)
            raw_output_type = doc.get("output_type", doc.get("category", "Legacy"))
            normalized_display_type = normalize_output_type(raw_output_type) if raw_output_type else "Legacy"
            
            record = {
                "id": str(doc["_id"]),
                "output_type": normalized_display_type,
                "academic_year": doc.get("academic_year", ""),
                "batch_range": batch_range,
                "student_names": doc.get("student_names", []),
                "created_at": doc.get("created_at", datetime.now(timezone.utc)).isoformat() if isinstance(doc.get("created_at"), datetime) else str(doc.get("created_at", ""))
            }
            
            # Format details based on output type
            details = []
            if doc.get("output_type") == "Sports":
                if doc.get("sport_name"):
                    details.append(doc["sport_name"])
                if doc.get("level"):
                    details.append(f"Level: {doc['level']}")
                if doc.get("number_of_wins", 0) > 0:
                    details.append(f"Wins: {doc['number_of_wins']}")
                record["details"] = " • ".join(details) if details else "Sports Event"
                record["participants"] = doc.get(
                    "number_of_participants",
                    doc.get("number_of_students", 0)
                )
            elif normalize_output_type(doc.get("output_type")) == "Cultural":
                event_category = doc.get("event_category", "")
                # Use custom category if event_category is "Other"
                if event_category == "Other" and doc.get("cultural_custom_category"):
                    details.append(doc["cultural_custom_category"])
                elif event_category:
                    details.append(event_category)
                if doc.get("event_name"):
                    details.append(doc["event_name"])
                record["details"] = " • ".join(details) if details else "Cultural Event"
                record["participants"] = doc.get(
                    "number_of_students",
                    doc.get("number_of_participants", 0)
                )
            elif normalize_output_type(doc.get("output_type")) == "Technical":
                if doc.get("event_type"):
                    details.append(doc["event_type"])
                if doc.get("event_name"):
                    details.append(doc["event_name"])
                if doc.get("organizer"):
                    details.append(f"by {doc['organizer']}")
                if doc.get("level"):
                    details.append(f"({doc['level']})")
                if doc.get("achievement"):
                    details.append(f"- {doc['achievement']}")
                record["details"] = " • ".join(details) if details else "Technical Event"
                record["participants"] = doc.get(
                    "number_of_students",
                    doc.get("number_of_participants", 0)
                )
            elif normalize_output_type(doc.get("output_type")) == "Internship":
                if doc.get("internship_domain"):
                    details.append(doc["internship_domain"])
                if doc.get("organization_name"):
                    details.append(f"at {doc['organization_name']}")
                if doc.get("mode"):
                    details.append(f"Mode: {doc['mode']}")
                if doc.get("duration"):
                    details.append(f"Duration: {doc['duration']}")
                record["details"] = " • ".join(details) if details else "Internship"
                record["participants"] = doc.get("number_of_students", doc.get("number_of_participants", 0))
            elif normalize_output_type(doc.get("output_type")) == "Courses":
                if doc.get("course_name"):
                    details.append(doc["course_name"])
                if doc.get("platform"):
                    details.append(f"on {doc['platform']}")
                if doc.get("mode"):
                    details.append(f"({doc['mode']})")
                record["details"] = " • ".join(details) if details else "Course/Certification"
                record["participants"] = doc.get("number_of_students", doc.get("number_of_participants", 0))
            elif normalize_output_type(doc.get("output_type")) == "Industrial Visit":
                if doc.get("industrial_visit_details"):
                    details.append(doc.get("industrial_visit_details"))
                if doc.get("faculty"):
                    details.append(f"Faculty: {doc['faculty']}")
                date_val = doc.get("date")
                if date_val not in (None, ""):
                    display_date = date_val
                    if isinstance(date_val, (int, float)):
                        try:
                            parsed = pd.to_datetime(date_val, unit="D", origin="1899-12-30").date()
                            display_date = parsed.isoformat()
                        except Exception:
                            display_date = str(date_val)
                    details.append(f"Date: {display_date}")
                record["details"] = " • ".join(details) if details else "Industrial Visit"
                record["participants"] = doc.get("number_of_students", doc.get("number_of_participants", 0))
            else:
                # Legacy format
                if doc.get("category"):
                    details.append(doc["category"])
                if doc.get("type"):
                    details.append(doc["type"])
                if doc.get("level"):
                    details.append(f"({doc['level']})")
                record["details"] = " • ".join(details) if details else "Extracurricular Event"
                record["participants"] = doc.get("count_participants", 0)
            
            formatted_records.append(record)
        
        return jsonify({
            "records": formatted_records,
            "total": total,
            "page": page,
            "per_page": per_page,
            "total_pages": (total + per_page - 1) // per_page
        }), 200
        
    except Exception as e:
        return jsonify({"message": f"Error retrieving records: {str(e)}"}), 500

@app.route('/api/extracurricular/records/<record_id>', methods=['DELETE', 'OPTIONS'])
@login_required
def delete_extracurricular_record(record_id):
    """Soft delete or finalize delete an extracurricular record by ID."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    college = session.get('college')
    branch = session.get('branch')
    if not college or not branch:
        return jsonify({"message": "User college or branch not found"}), 400

    try:
        record = db.extracurricular_records.find_one({"_id": ObjectId(record_id)})
        if not record:
            return jsonify({"message": "Record not found"}), 404
        
        # Verify ownership
        batch = db.batches.find_one({"_id": record['batch_id']})
        if not batch or batch.get('college') != college or batch.get('branch') != branch:
             return jsonify({"message": "Unauthorized access"}), 403

        finalize = (request.args.get('finalize') or '').lower() == 'true'
        if not finalize:
            db.extracurricular_records.update_one(
                {"_id": ObjectId(record_id)},
                {"$set": {
                    "pending_delete": True,
                    "delete_requested_at": datetime.now(timezone.utc)
                }}
            )
            return jsonify({"message": "Record marked for deletion", "pending_delete": True}), 200
        else:
            fs = GridFS(db)
            file_ids = []
            if "evidence_file_ids" in record and record["evidence_file_ids"]:
                file_ids = record["evidence_file_ids"]
            elif "evidence_file_id" in record and record["evidence_file_id"]:
                file_ids = [record["evidence_file_id"]]
            for file_id_str in file_ids:
                try:
                    fs.delete(ObjectId(file_id_str))
                except Exception as e:
                    print(f"⚠️ Warning: Could not delete GridFS file {file_id_str}: {e}")
            result = db.extracurricular_records.delete_one({"_id": ObjectId(record_id)})
            if result.deleted_count == 0:
                return jsonify({"message": "Record not found"}), 404
            return jsonify({"message": "Record permanently deleted", "finalized": True}), 200
        
    except Exception as e:
        return jsonify({"message": f"Error deleting record: {str(e)}"}), 500

@app.route('/api/placement-outcomes/<outcome_id>/restore', methods=['POST', 'OPTIONS'])
@login_required
def restore_placement_outcome(outcome_id):
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    try:
        college = session.get('college')
        branch = session.get('branch')
        if not college or not branch:
            return jsonify({"message": "User college or branch not found"}), 400
        
        try:
            object_id = ObjectId(outcome_id)
        except Exception:
            return jsonify({"message": "Invalid record ID format"}), 400
        
        record = db.placement_outcomes.find_one({"_id": object_id})
        if not record:
            return jsonify({"message": "Record not found"}), 404
        
        # Verify ownership via batch/college/branch
        if record.get('college') != college or record.get('branch') != branch:
            batch_id = record.get('batch_id')
            if batch_id:
                batch = db.batches.find_one({"_id": batch_id, "college": college, "branch": branch})
                if not batch:
                    return jsonify({"message": "Unauthorized access to this record"}), 403
            else:
                return jsonify({"message": "Unauthorized access to this record"}), 403
        
        db.placement_outcomes.update_one(
            {"_id": object_id},
            {"$unset": {"pending_delete": "", "delete_requested_at": "", "deleted": ""}}
        )
        return jsonify({"message": "Record restored successfully"}), 200
    except Exception as e:
        return jsonify({"message": f"Error restoring record: {str(e)}"}), 500
@app.route('/api/extracurricular/records/<record_id>/restore', methods=['POST', 'OPTIONS'])
@login_required
def restore_extracurricular_record(record_id):
    """Restore a soft-deleted extracurricular record."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    college = session.get('college')
    branch = session.get('branch')
    if not college or not branch:
        return jsonify({"message": "User college or branch not found"}), 400
    
    try:
        record = db.extracurricular_records.find_one({"_id": ObjectId(record_id)})
        if not record:
            return jsonify({"message": "Record not found"}), 404
        
        # Verify ownership
        batch = db.batches.find_one({"_id": record['batch_id']})
        if not batch or batch.get('college') != college or batch.get('branch') != branch:
                return jsonify({"message": "Unauthorized access"}), 403
        
        db.extracurricular_records.update_one(
            {"_id": ObjectId(record_id)},
            {"$unset": {"pending_delete": "", "delete_requested_at": "", "deleted": ""}}
        )
        return jsonify({"message": "Record restored successfully"}), 200
        
    except Exception as e:
        return jsonify({"message": f"Error restoring record: {str(e)}"}), 500

@app.route('/api/extracurricular/records/delete_all', methods=['POST', 'OPTIONS'])
@login_required
def delete_all_extracurricular_records():
    """Delete all extracurricular records for the current college/branch/batch."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    college = session.get('college')
    branch = session.get('branch')
    if not college or not branch:
        return jsonify({"message": "User college or branch not found"}), 400

    try:
        # Build query
        query = {"college": college, "branch": branch}
        
        # Optional: Filter by batch if provided in request body
        data = request.get_json() or {}
        batch_id = data.get('batch_id')
        batch_name = data.get('batch_name')
        
        if batch_id:
            query['batch_id'] = ObjectId(batch_id)
        elif batch_name:
            # Look up batch by name (batch_range)
            batch = db.batches.find_one({"batch_range": batch_name, "college": college, "branch": branch})
            if batch:
                query['batch_id'] = batch["_id"]
            else:
                return jsonify({"message": f"Batch '{batch_name}' not found"}), 404
            
        # Find all records to delete their files first
        records = list(db.extracurricular_records.find(query))
        
        if not records:
             return jsonify({"message": "No records found to delete", "deleted_count": 0}), 200

        # Delete associated files from GridFS
        fs = GridFS(db)
        file_deletion_errors = 0
        for record in records:
            file_ids = []
            if "evidence_file_ids" in record and record["evidence_file_ids"]:
                file_ids.extend(record["evidence_file_ids"])
            if "evidence_file_id" in record and record["evidence_file_id"]:
                file_ids.append(record["evidence_file_id"])
            
            for file_id_str in file_ids:
                try:
                    fs.delete(ObjectId(file_id_str))
                except Exception:
                    file_deletion_errors += 1
        
        # Delete records
        result = db.extracurricular_records.delete_many(query)
        
        return jsonify({
            "message": "All records deleted successfully", 
            "deleted_count": result.deleted_count,
            "file_deletion_errors": file_deletion_errors
        }), 200
        
    except Exception as e:
        return jsonify({"message": f"Error deleting records: {str(e)}"}), 500

@app.route('/api/extracurricular/summary', methods=['GET', 'OPTIONS'])
@login_required
def get_extracurricular_summary():
    """Get aggregated summary for dashboard (no student names)."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    academic_year = request.args.get('year') or request.args.get('academic_year')
    batch_id = request.args.get('batch_id')
    
    college = session.get('college')
    branch = session.get('branch')
    if not college or not branch:
        return jsonify({"message": "User college or branch not found"}), 400

    try:
        # Build query
        query = {}
        if academic_year:
            query["academic_year"] = academic_year
        if batch_id:
            try:
                batch_obj_id = ObjectId(batch_id)
                # Verify ownership
                batch = db.batches.find_one({"_id": batch_obj_id, "college": college, "branch": branch})
                if not batch:
                     return jsonify({"message": "Invalid batch ID or unauthorized access"}), 400
                query["batch_id"] = batch_obj_id
            except Exception:
                pass
        else:
            # Filter by all batches for this college and branch
            branch_batches = list(db.batches.find({"college": college, "branch": branch}, {"_id": 1}))
            branch_batch_ids = [b["_id"] for b in branch_batches]
            query['batch_id'] = {"$in": branch_batch_ids}
        # Exclude soft-deleted records
        query['deleted'] = {"$ne": True}
        query['pending_delete'] = {"$ne": True}
        
        # Get sanctioned intake for participation index calculation
        dashboard_summary = db.dashboard_summary.find_one({"_id": "summary"}) or {}
        kpis = dashboard_summary.get("kpis", {})
        sanctioned_intake = kpis.get("sanctioned_intake", 0) or 0
        
        # Aggregate pipeline for summary - handle both old and new structures
        pipeline = [
            {"$match": query},
            {"$addFields": {
                # Normalize participant count from both old and new structures
                "normalized_participants": {
                    "$ifNull": [
                        {"$ifNull": ["$number_of_participants", "$number_of_students"]},
                        "$count_participants"
                    ]
                },
                # Normalize wins count from both structures
                "normalized_won": {
                    "$ifNull": [
                        "$number_of_wins",
                        "$count_won"
                    ]
                },
                # Normalize category/output_type
                "normalized_category": {
                    "$ifNull": [
                        "$output_type",
                        "$category"
                    ]
                }
            }},
            {"$group": {
                "_id": None,
                "total_events": {"$sum": 1},
                "total_participants": {"$sum": {"$ifNull": ["$normalized_participants", 0]}},
                "total_won": {"$sum": {"$ifNull": ["$normalized_won", 0]}},
                "by_category": {
                    "$push": {
                        "category": {"$ifNull": ["$normalized_category", "Unknown"]},
                        "participants": {"$ifNull": ["$normalized_participants", 0]}
                    }
                },
                "by_level": {
                    "$push": {
                        "level": {"$ifNull": ["$level", "Unknown"]},
                        "participants": {"$ifNull": ["$normalized_participants", 0]}
                    }
                }
            }}
        ]
        
        result = list(db.extracurricular_records.aggregate(pipeline))
        
        if not result:
            return jsonify({
                "total_events": 0,
                "total_participants": 0,
                "total_won": 0,
                "participation_index": 0,
                "distribution_by_category": {},
                "participants_by_level": {},
                "monthly_trend": []
            }), 200
        
        summary = result[0]
        
        # Calculate participation index in Python (not in MongoDB aggregation)
        total_participants = summary.get("total_participants", 0) or 0
        participation_index = 0
        if sanctioned_intake > 0 and total_participants > 0:
            participation_index = round((total_participants / sanctioned_intake) * 100, 2)
        
        # Process category distribution (normalize for backward compatibility)
        category_dist = {}
        for item in summary.get("by_category", []):
            cat = normalize_output_type(item.get("category", "Unknown"))
            category_dist[cat] = category_dist.get(cat, 0) + item.get("participants", 0)
        
        # Process level distribution with automatic inference
        # Fetch full records to infer missing levels from event names/types
        full_records = list(db.extracurricular_records.find(query, {
            "level": 1, "type": 1, "category": 1, "output_type": 1, 
            "count_participants": 1, "number_of_participants": 1, "number_of_students": 1,
            "event_name": 1, "event_type": 1,
            "organization_name": 1, "organizer": 1, "mode": 1, "is_paid": 1, "stipend": 1
        }))
        
        def infer_level(record):
            """Infer level from event name/type/category when level is missing."""
            level = record.get("level")
            if level and str(level).strip() and str(level).strip().lower() != "unknown":
                return str(level).strip()
            
            # Get event name/type for inference (handle both old and new structures)
            event_type = str(record.get("event_type") or record.get("type", "")).lower()
            event_name = str(record.get("event_name", "")).lower()
            event_category = str(record.get("category") or record.get("output_type", "")).lower()
            # Combine event_type and event_name for better inference
            event_text = f"{event_type} {event_name}".strip()
            
            # College-level indicators
            college_keywords = ["synergy", "college fest", "college festival", "intra", "intra-college", 
                              "college level", "campus", "institute", "college cultural"]
            if any(keyword in event_text for keyword in college_keywords):
                return "College"
            
            # National/International indicators
            if "national" in event_text or "nation" in event_text:
                return "National"
            if "international" in event_text or "internat" in event_text or "global" in event_text:
                return "International"
            if "state" in event_text or "state-level" in event_text:
                return "State"
            if "zonal" in event_text or "zone" in event_text:
                return "Zonal"
            if "university" in event_text or "univ" in event_text:
                return "University"
            
            # Category-based inference
            if event_category == "cultural":
                # Cultural events without level specified are often college-level
                return "College"
            elif event_category == "sports":
                # Sports events without level - could be college or higher, default to College
                return "College"
            elif event_category == "student activities":
                # Student activities are typically college-level
                return "College"
            
            # Default to College for missing levels (most common case)
            return "College"
        
        level_dist = {}
        internship_dist = {"Online": 0, "Offline": 0, "Paid": 0, "Unpaid": 0}
        participation_location_dist = {"College": 0, "Outside": 0}
        
        # Enhanced College Level Detection Logic
        def normalize_str_safe(s):
            if not s: return ""
            return " ".join(str(s).strip().lower().split())

        current_college_norm = normalize_str_safe(college)
        college_abbr_norm = ""
        
        # Fetch optional abbreviation from college record
        if college:
            try:
                col_doc = db.colleges.find_one({"name": college})
                if col_doc and col_doc.get("abbreviation"):
                    college_abbr_norm = normalize_str_safe(col_doc.get("abbreviation"))
            except Exception:
                pass

        for record in full_records:
            inferred_level = infer_level(record)
            # Normalize participants count
            participants = record.get("count_participants") or record.get("number_of_participants") or record.get("number_of_students") or 0
            
            level_dist[inferred_level] = level_dist.get(inferred_level, 0) + participants

            # Normalize category
            cat = record.get("output_type") or record.get("category", "")
            cat_normalized = normalize_output_type(cat) if 'normalize_output_type' in globals() else str(cat).capitalize()
            if cat_normalized == "Internships": cat_normalized = "Internship"

            # --- Internship Distribution Logic ---
            # Check if category/output_type is Internship
            if cat_normalized == "Internship" or str(cat).lower() == "internship":
                mode = str(record.get("mode", "")).strip().lower()
                
                # Determine Online vs Offline (ignore Hybrid)
                if mode == "online":
                    internship_dist["Online"] += 1
                elif mode == "offline":
                    internship_dist["Offline"] += 1
                # Hybrid is ignored as per requirements

            # --- Participation Location Logic (Standardized) ---
            # Categories: Internship, Sports, Technical, Cultural
            target_cats = ["Internship", "Sports", "Technical", "Cultural"]
            
            if cat_normalized in target_cats:
                is_loc_outside = False
                is_loc_college = False
                count_val = participants # Default to participants count
                
                if cat_normalized == "Internship":
                    # Rule 1: ALL internships are considered OUTSIDE COLLEGE
                    is_loc_outside = True
                    
                elif cat_normalized == "Sports":
                    # Rule 2: Sports Level Check
                    lvl = str(record.get("level", "")).lower()
                    
                    # Check College keywords
                    if "inter-department" in lvl or "department" in lvl:
                        is_loc_college = True
                    # Check Outside keywords
                    elif any(x in lvl for x in ["inter-college", "university", "state", "national"]):
                        is_loc_outside = True
                    else:
                        # Fallback to inferred level logic
                        if inferred_level in ["Inter-College", "State", "National", "International", "Zonal", "University"]:
                            is_loc_outside = True
                        else:
                            is_loc_college = True

                elif cat_normalized == "Technical":
                    # Rule 3: Technical -> Organizer Check (Count PARTICIPANTS)
                    # count_val remains participants (default)
                    org_name = normalize_str_safe(record.get("organization_name", ""))
                    organizer = normalize_str_safe(record.get("organizer", ""))
                    check_str = org_name + " " + organizer
                    
                    # If organizer contains college name (or abbreviation) or department -> College Level
                    is_match = False
                    if current_college_norm and current_college_norm in check_str:
                        is_match = True
                    elif college_abbr_norm and college_abbr_norm in check_str:
                        is_match = True
                        
                    if is_match or "department" in check_str:
                        is_loc_college = True
                    else:
                        is_loc_outside = True

                elif cat_normalized == "Cultural":
                    # Rule 4: Cultural -> Organizer Check (Count PARTICIPANTS)
                    org_name = normalize_str_safe(record.get("organization_name", ""))
                    organizer = normalize_str_safe(record.get("organizer", ""))
                    check_str = org_name + " " + organizer
                    
                    # If organizer contains college name (or abbreviation) or department -> College Level
                    is_match = False
                    if current_college_norm and current_college_norm in check_str:
                        is_match = True
                    elif college_abbr_norm and college_abbr_norm in check_str:
                        is_match = True
                        
                    if is_match or "department" in check_str:
                        is_loc_college = True
                    else:
                        is_loc_outside = True
                
                # Apply to distribution
                if is_loc_outside:
                    participation_location_dist["Outside"] += count_val
                elif is_loc_college:
                    participation_location_dist["College"] += count_val
        
        # Monthly trend - fetch records and process in Python for date parsing
        records = list(db.extracurricular_records.find(query, {"event_date": 1, "count_participants": 1}))
        monthly_data_dict = {}
        for record in records:
            try:
                event_date_str = record.get("event_date")
                if not event_date_str:
                    continue
                # Handle both string and date objects
                if isinstance(event_date_str, str):
                    event_date = pd.to_datetime(event_date_str).date()
                else:
                    event_date = pd.to_datetime(event_date_str).date()
                period_key = f"{event_date.year}-{event_date.month:02d}"
                if period_key not in monthly_data_dict:
                    monthly_data_dict[period_key] = {"participants": 0, "events": 0}
                monthly_data_dict[period_key]["participants"] += record.get("count_participants", 0) or 0
                monthly_data_dict[period_key]["events"] += 1
            except Exception as e:
                print(f"⚠️ Warning: Could not parse date for record: {record.get('_id')}, error: {e}")
                continue
        
        monthly_data = [
            {
                "period": period,
                "participants": data["participants"],
                "events": data["events"]
            }
            for period, data in sorted(monthly_data_dict.items())
        ]
        
        # Remove debug logs for Extracurricular Summary
        pass

        internship_dist_filtered = {
            "Online": internship_dist.get("Online", 0),
            "Offline": internship_dist.get("Offline", 0)
        }

        return jsonify({
            "total_events": summary.get("total_events", 0) or 0,
            "total_participants": summary.get("total_participants", 0) or 0,
            "total_won": summary.get("total_won", 0) or 0,
            "participation_index": participation_index,
            "distribution_by_category": category_dist,
            "participants_by_level": level_dist,
            "monthly_trend": monthly_data,
            "internship_distribution": internship_dist_filtered,
            "participation_location": participation_location_dist
        }), 200
        
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"❌ Error in extracurricular summary: {str(e)}")
        print(f"Traceback: {error_trace}")
        return jsonify({"message": f"Error generating summary: {str(e)}"}), 500

@app.route('/api/extracurricular/trend', methods=['GET', 'OPTIONS'])
def get_extracurricular_trend():
    """Get timeseries trend data."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    academic_year = request.args.get('year') or request.args.get('academic_year')
    category = request.args.get('category')
    period = request.args.get('period', 'monthly')  # monthly, quarterly, yearly
    
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
             return jsonify({"message": "User college or branch not found in session"}), 400

        query = {"college": college, "branch": branch}
        if academic_year:
            query["academic_year"] = academic_year
        if category:
            query["category"] = category
        
        # Fetch records and process in Python for flexible date parsing
        records = list(db.extracurricular_records.find(query, {"event_date": 1, "count_participants": 1, "count_won": 1}))
        trend_dict = {}
        
        for record in records:
            try:
                event_date = pd.to_datetime(record.get("event_date")).date()
                
                if period == "monthly":
                    period_key = f"{event_date.year}-{event_date.month:02d}"
                elif period == "quarterly":
                    quarter = (event_date.month - 1) // 3 + 1
                    period_key = f"{event_date.year}-Q{quarter}"
                else:  # yearly
                    period_key = str(event_date.year)
                
                if period_key not in trend_dict:
                    trend_dict[period_key] = {"participants": 0, "events": 0, "won": 0}
                
                trend_dict[period_key]["participants"] += record.get("count_participants", 0)
                trend_dict[period_key]["events"] += 1
                trend_dict[period_key]["won"] += record.get("count_won", 0)
            except Exception:
                continue
        
        trend_data = [
            {
                "period": period,
                "participants": data["participants"],
                "events": data["events"],
                "won": data["won"]
            }
            for period, data in sorted(trend_dict.items())
        ]
        
        return jsonify(trend_data), 200
        
    except Exception as e:
        return jsonify({"message": f"Error generating trend: {str(e)}"}), 500

@app.route('/api/extracurricular/top', methods=['GET', 'OPTIONS'])
def get_extracurricular_top():
    """Get top achievements by count_won."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    academic_year = request.args.get('year') or request.args.get('academic_year')
    limit = int(request.args.get('limit', 3))
    
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
             return jsonify({"message": "User college or branch not found in session"}), 400

        query = {"college": college, "branch": branch}
        if academic_year:
            query["academic_year"] = academic_year
        
        # Only include records with wins
        query["count_won"] = {"$gt": 0}
        
        results = list(
            db.extracurricular_records
            .find(query)
            .sort("count_won", -1)
            .limit(limit)
        )
        
        top_achievements = []
        for doc in results:
            top_achievements.append({
                "category": doc.get("category"),
                "level": doc.get("level"),
                "type": doc.get("type"),
                "count_won": doc.get("count_won", 0),
                "count_participants": doc.get("count_participants", 0),
                "event_date": doc.get("event_date"),
                "organizer": doc.get("organizer")
            })
        
        return jsonify(top_achievements), 200
        
    except Exception as e:
        return jsonify({"message": f"Error fetching top achievements: {str(e)}"}), 500

@app.route('/api/extracurricular/records', methods=['GET', 'OPTIONS'])
def get_extracurricular_records():
    """Get detailed records for admin/faculty (includes evidence links)."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    academic_year = request.args.get('year') or request.args.get('academic_year')
    category = request.args.get('category')
    verified = request.args.get('verified')
    approved = request.args.get('approved')
    batch_id = request.args.get('batch_id')
    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 50))
    
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
             return jsonify({"message": "User college or branch not found in session"}), 400

        # Run cleanup for this user's context
        cleanup_pending_deletes(college, branch)

        query = {"college": college, "branch": branch}
        if academic_year:
            query["academic_year"] = academic_year
        if category:
            query["category"] = category
        if verified is not None:
            query["verified"] = verified.lower() == 'true'
        if approved is not None:
            query["approved"] = approved.lower() == 'true'
        if batch_id:
            try:
                batch_obj_id = ObjectId(batch_id)
                query["batch_id"] = batch_obj_id
            except Exception:
                pass
        
        # Exclude soft-deleted records
        query["deleted"] = {"$ne": True}
        query["pending_delete"] = {"$ne": True}
        
        skip = (page - 1) * per_page
        
        # Get total count
        total = db.extracurricular_records.count_documents(query)
        
        # Get records
        records = list(
            db.extracurricular_records
            .find(query)
            .sort("event_date", -1)
            .skip(skip)
            .limit(per_page)
        )
        
        # Format records - include both old and new structure fields
        formatted_records = []
        for doc in records:
            record = {
                "id": str(doc["_id"]),
                "academic_year": doc.get("academic_year"),
                "event_date": doc.get("event_date"),
                # Legacy fields
                "category": doc.get("category"),
                "level": doc.get("level"),
                "type": doc.get("type"),
                "count_participants": doc.get("count_participants", 0),
                "count_won": doc.get("count_won", 0),
                "organizer": doc.get("organizer"),
                "remarks": doc.get("remarks"),
                # New structure fields
                "output_type": doc.get("output_type"),
                "sport_name": doc.get("sport_name"),
                "event_name": doc.get("event_name"),
                "event_type": doc.get("event_type"),
                "number_of_participants": doc.get("number_of_participants"),
                "number_of_students": doc.get("number_of_students"),
                "number_of_wins": doc.get("number_of_wins"),
                "organization_name": doc.get("organization_name"),
                "mode": doc.get("mode"),
                "duration": doc.get("duration"),
                "has_stipend": doc.get("has_stipend"),
                "stipend_amount": doc.get("stipend_amount"),
                "course_name": doc.get("course_name"),
                "platform": doc.get("platform"),
                "achievement": doc.get("achievement"),
                "industrial_visit_details": doc.get("industrial_visit_details"),
                "faculty": doc.get("faculty"),
                "date": doc.get("date"),
                # Common fields
                "evidence_file_id": str(doc.get("evidence_file_id")) if doc.get("evidence_file_id") else None,
                "evidence_file_ids": [str(fid) for fid in doc.get("evidence_file_ids", [])] if doc.get("evidence_file_ids") else [],
                "verified": doc.get("verified", False),
                "approved": doc.get("approved", False),
                "created_at": doc.get("created_at").isoformat() if doc.get("created_at") else None
            }
            formatted_records.append(record)
        
        return jsonify({
            "records": formatted_records,
            "total": total,
            "page": page,
            "per_page": per_page,
            "total_pages": (total + per_page - 1) // per_page
        }), 200
        
    except Exception as e:
        return jsonify({"message": f"Error fetching records: {str(e)}"}), 500

@app.route('/api/extracurricular/verify/<record_id>', methods=['POST', 'OPTIONS'])
def verify_extracurricular_record(record_id):
    """Verify an extracurricular record (admin/faculty only)."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    try:
        result = db.extracurricular_records.update_one(
            {"_id": ObjectId(record_id)},
            {"$set": {"verified": True, "verified_at": datetime.now(timezone.utc)}}
        )
        
        if result.matched_count == 0:
            return jsonify({"message": "Record not found"}), 404
        
        return jsonify({"message": "Record verified successfully"}), 200
        
    except Exception as e:
        return jsonify({"message": f"Error verifying record: {str(e)}"}), 500

@app.route('/api/extracurricular/approve/<record_id>', methods=['POST', 'OPTIONS'])
def approve_extracurricular_record(record_id):
    """Approve an extracurricular record (admin only)."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    try:
        result = db.extracurricular_records.update_one(
            {"_id": ObjectId(record_id)},
            {"$set": {"approved": True, "approved_at": datetime.now(timezone.utc)}}
        )
        
        if result.matched_count == 0:
            return jsonify({"message": "Record not found"}), 404
        
        return jsonify({"message": "Record approved successfully"}), 200
        
    except Exception as e:
        return jsonify({"message": f"Error approving record: {str(e)}"}), 500

@app.route('/api/extracurricular/download-evidence/<file_id>', methods=['GET', 'OPTIONS'])
def download_extracurricular_evidence(file_id):
    """Download evidence file for an extracurricular record."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    try:
        fs = GridFS(db)
        file_data = fs.get(ObjectId(file_id))
        
        if not file_data:
            return jsonify({"message": "File not found"}), 404
        
        from flask import Response
        return Response(
            file_data.read(),
            mimetype=file_data.content_type or 'application/octet-stream',
            headers={'Content-Disposition': f'attachment; filename="{file_data.filename}"'}
        )
        
    except Exception as e:
        return jsonify({"message": f"Error downloading file: {str(e)}"}), 500



# --- Placement & Higher Education Module APIs ---

# Initialize indexes for placement_records collection
def init_placement_indexes():
    """Create indexes for placement_records collection."""
    try:
        db.placement_records.create_index([("batch", 1)])
        db.placement_records.create_index([("student_id", 1)])
        db.placement_records.create_index([("student_name", 1)])
        db.placement_records.create_index([("placed", 1)])
        db.placement_records.create_index([("higher_studies", 1)])
        # Silent on success
        pass
    except Exception as e:
        print(f"⚠️ Warning: Could not create placement indexes: {e}")

# Initialize indexes on startup
init_placement_indexes()

# Validation constants
PLACEMENT_COMPANY_CATEGORIES = ["Core", "IT", "Non-IT", "Startup"]
PLACEMENT_TYPES = ["On-campus", "Off-campus"]
HIGHER_STUDIES_FIELDS = ["MBA", "MTech", "MS", "PhD", "MSc", "Other"]

def validate_placement_record(data, is_bulk=False):
    """Validate a single placement record.
    
    Returns: (is_valid, error_message, validated_data)
    """
    errors = []
    validated = {}
    
    # Required fields
    required_fields = ["student_id", "student_name", "batch"]
    for field in required_fields:
        if field not in data or (isinstance(data[field], str) and not data[field].strip()):
            errors.append(f"Missing required field: {field}")
        else:
            validated[field] = str(data[field]).strip()
    
    if errors:
        return False, "; ".join(errors), None
    
    # Boolean fields with Yes/No or True/False conversion
    def parse_yes_no(value):
        if value is None:
            return False
        if isinstance(value, bool):
            return value
        value_str = str(value).strip().lower()
        return value_str in ['yes', 'y', 'true', '1', 'placed']
    
    # Placement details
    validated["placed"] = parse_yes_no(data.get("placed"))
    if validated["placed"]:
        validated["company_name"] = str(data.get("company_name", "")).strip() if data.get("company_name") else None
        company_category = str(data.get("company_category", "")).strip()
        if company_category and company_category in PLACEMENT_COMPANY_CATEGORIES:
            validated["company_category"] = company_category
        elif company_category:
            # Try to match case-insensitive
            matched = None
            for cat in PLACEMENT_COMPANY_CATEGORIES:
                if cat.lower() == company_category.lower():
                    matched = cat
                    break
            validated["company_category"] = matched if matched else None
        else:
            validated["company_category"] = None
        
        # Standardize Package Fields
        # 1. ctc_package (Display String) - matches placement_outcomes & report
        # 2. package_lpa (Numeric Float) - for calculations/sorting
        
        raw_package = data.get("ctc_package") or data.get("ctc") or data.get("package_lpa") or data.get("package")
        validated["ctc_package"] = str(raw_package).strip() if raw_package is not None else None
        
        validated["package_lpa"] = None
        if raw_package:
            try:
                # Try to extract numeric value
                import re
                match = re.search(r"(\d+(\.\d+)?)", str(raw_package))
                if match:
                    pkg = float(match.group(1))
                    if pkg >= 0:
                        validated["package_lpa"] = round(pkg, 2)
            except (ValueError, TypeError):
                pass
        
        placement_type = str(data.get("placement_type", "")).strip()
        if placement_type and placement_type in PLACEMENT_TYPES:
            validated["placement_type"] = placement_type
        elif placement_type:
            # Try to match case-insensitive
            matched = None
            for pt in PLACEMENT_TYPES:
                if pt.lower() == placement_type.lower():
                    matched = pt
                    break
            validated["placement_type"] = matched if matched else None
        else:
            validated["placement_type"] = None
        
        validated["job_role"] = str(data.get("job_role", "")).strip() if data.get("job_role") else None
        
        # Date of offer
        validated["date_of_offer"] = None
        if data.get("date_of_offer"):
            try:
                offer_date = pd.to_datetime(data.get("date_of_offer")).date()
                validated["date_of_offer"] = offer_date.isoformat()
            except Exception:
                pass
    else:
        validated["company_name"] = None
        validated["company_category"] = None
        validated["package_lpa"] = None
        validated["ctc_package"] = None
        validated["placement_type"] = None
        validated["job_role"] = None
        validated["date_of_offer"] = None
    
    # Higher studies details
    validated["higher_studies"] = parse_yes_no(data.get("higher_studies"))
    if validated["higher_studies"]:
        field_of_study = str(data.get("field_of_study", "")).strip()
        validated["field_of_study"] = field_of_study if field_of_study else None
        validated["university_name"] = str(data.get("university_name", "")).strip() if data.get("university_name") else None
        validated["country"] = str(data.get("country", "")).strip() if data.get("country") else None
        validated["entrance_exams"] = str(data.get("entrance_exams", "")).strip() if data.get("entrance_exams") else None
    else:
        validated["field_of_study"] = None
        validated["university_name"] = None
        validated["country"] = None
        validated["entrance_exams"] = None
    
    # Entrepreneurship
    validated["entrepreneurship"] = parse_yes_no(data.get("entrepreneurship"))
    if validated["entrepreneurship"]:
        validated["startup_name"] = str(data.get("startup_name", "")).strip() if data.get("startup_name") else None
    else:
        validated["startup_name"] = None
    
    # Not placed / Not interested
    validated["not_placed"] = parse_yes_no(data.get("not_placed"))
    validated["not_interested"] = parse_yes_no(data.get("not_interested"))
    
    if errors:
        return False, "; ".join(errors), None
    
    return True, None, validated

@app.route('/api/placement/upload', methods=['POST', 'OPTIONS'])
def upload_placement_bulk():
    """Bulk upload placement records via CSV/Excel with validation report."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    if 'file' not in request.files:
        return jsonify({"message": "No file part"}), 400
    
    file = request.files['file']
    batch_id = request.form.get('batch_id')
    
    if not batch_id:
        return jsonify({"message": "Batch is required"}), 400
    
    # Validate batch_id
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
             return jsonify({"message": "User college or branch not found in session"}), 400

        batch = None
        batch_obj_id = None
        
        if ObjectId.is_valid(batch_id):
            batch_obj_id = ObjectId(batch_id)
            batch = db.batches.find_one({"_id": batch_obj_id, "college": college, "branch": branch})
        
        if not batch:
            # Try finding by batch_range string
            batch = db.batches.find_one({"batch_range": batch_id, "college": college, "branch": branch})
            if batch:
                batch_obj_id = batch["_id"]

        if not batch:
            return jsonify({"message": "Invalid batch ID or unauthorized access"}), 400
    except Exception:
        return jsonify({"message": "Invalid batch ID format"}), 400
    
    if file.filename == '':
        return jsonify({"message": "No file selected"}), 400
    
    try:
        # Read file content
        file_content = file.read()
        file.seek(0)
        
        # Determine file type and read
        filename_lower = file.filename.lower()
        if filename_lower.endswith('.csv'):
            df = pd.read_csv(BytesIO(file_content))
        elif filename_lower.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(BytesIO(file_content))
        else:
            return jsonify({"message": "Unsupported file format. Use CSV or Excel."}), 400
        
        # Normalize column names (case-insensitive, strip whitespace, handle newlines)
        df.columns = [str(c).replace('\n', ' ').strip() for c in df.columns]
        
        # Map common column name variations
        column_mapping = {
            'student_id': ['student_id', 'student id', 'Student ID', 'PRN', 'prn', 'PRN Number', 'prn_number', 'urn', 'URN', 'roll_number', 'roll number', 'Roll Number'],
            'student_name': ['student_name', 'student name', 'Student Name', 'Name', 'name', 'student_names', 'student names', 'Student Names', 'name_of_student', 'name of student', 'Name of Student', 'name_of_students', 'name of students', 'Name of Students', 'name of student(s)', 'student list', 'list of students', 'participants', 'participant names'],
            'batch': ['batch', 'Batch', 'BATCH', 'academic_year', 'academic year'],
            'placed': ['placed', 'Placed', 'PLACED', 'is_placed', 'is placed'],
            'company_name': ['company_name', 'company name', 'Company Name', 'Company', 'company'],
            'company_category': ['company_category', 'company category', 'Company Category', 'Category', 'category', 'company_type', 'company type'],
            'package_lpa': ['package_lpa', 'package lpa', 'Package LPA', 'Package', 'package', 'salary', 'Salary', 'CTC', 'ctc'],
            'placement_type': ['placement_type', 'placement type', 'Placement Type', 'Type', 'type'],
            'job_role': ['job_role', 'job role', 'Job Role', 'Role', 'role', 'designation', 'Designation'],
            'date_of_offer': ['date_of_offer', 'date of offer', 'Date of Offer', 'Offer Date', 'offer_date'],
            'higher_studies': ['higher_studies', 'higher studies', 'Higher Studies', 'pg', 'PG', 'post_graduation', 'post graduation'],
            'field_of_study': ['field_of_study', 'field of study', 'Field of Study', 'Course', 'course'],
            'university_name': ['university_name', 'university name', 'University Name', 'University', 'university', 'institute', 'Institute'],
            'country': ['country', 'Country', 'COUNTRY'],
            'entrance_exams': ['entrance_exams', 'entrance exams', 'Entrance Exams', 'Exams', 'exams'],
            'entrepreneurship': ['entrepreneurship', 'Entrepreneurship', 'startup', 'Startup'],
            'startup_name': ['startup_name', 'startup name', 'Startup Name'],
            'not_placed': ['not_placed', 'not placed', 'Not Placed'],
            'not_interested': ['not_interested', 'not interested', 'Not Interested']
        }
        
        # Normalize column names
        normalized_cols = {}
        for standard_name, variations in column_mapping.items():
            for col in df.columns:
                if col in variations or col.lower() == standard_name.lower():
                    normalized_cols[col] = standard_name
                    break
        
        # Rename columns
        df = df.rename(columns=normalized_cols)
        
        # Validation report
        valid_records = []
        invalid_records = []
        
        for idx, row in df.iterrows():
            row_dict = row.to_dict()
            # Remove NaN values
            row_dict = {k: v for k, v in row_dict.items() if pd.notna(v)}
            
            is_valid, error_msg, validated_data = validate_placement_record(row_dict, is_bulk=True)
            
            if is_valid:
                valid_records.append({
                    "row": int(idx) + 2,  # +2 because 0-indexed and header row
                    "data": validated_data
                })
            else:
                invalid_records.append({
                    "row": int(idx) + 2,
                    "data": row_dict,
                    "error": error_msg
                })
        
        # Return validation report (preview before commit)
        return jsonify({
            "message": f"Validation complete. {len(valid_records)} valid, {len(invalid_records)} invalid records.",
            "valid_count": len(valid_records),
            "invalid_count": len(invalid_records),
            "valid_records": valid_records,  # Return all valid records for commit
            "valid_records_preview": valid_records[:10],  # Preview first 10 for display
            "invalid_records": invalid_records,
            "batch_id": batch_id,  # Return batch_id for commit
            "preview": True
        }), 200
        
    except Exception as e:
        return jsonify({"message": f"Error processing file: {str(e)}"}), 500

@app.route('/api/placement/upload-commit', methods=['POST', 'OPTIONS'])
def commit_placement_bulk():
    """Commit validated records from bulk upload."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    data = request.get_json()
    # Support both formats: direct records array or wrapped in valid_records
    records = data.get('records', [])
    if not records and 'valid_records' in data:
        # Extract data from valid_records format (from upload endpoint)
        records = [item.get('data', item) for item in data.get('valid_records', [])]
    
    if not records:
        return jsonify({"message": "No records to commit"}), 400
    
    # Get batch_id from request
    batch_id = data.get('batch_id')
    if not batch_id:
        return jsonify({"message": "Batch is required"}), 400
    
    # Validate batch_id
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
             return jsonify({"message": "User college or branch not found in session"}), 400

        batch_obj_id = ObjectId(batch_id)
        batch = db.batches.find_one({"_id": batch_obj_id, "college": college, "branch": branch})
        if not batch:
            return jsonify({"message": "Invalid batch ID or unauthorized access"}), 400
    except Exception:
        return jsonify({"message": "Invalid batch ID format"}), 400
    
    try:
        inserted_count = 0
        errors = []
        
        for record_data in records:
            # Re-validate before insert
            is_valid, error_msg, validated_data = validate_placement_record(record_data, is_bulk=False)
            
            if not is_valid:
                errors.append({"record": record_data, "error": error_msg})
                continue
            
            # Create document
            doc = {
                "student_id": validated_data["student_id"],
                "student_name": validated_data["student_name"],
                "batch": validated_data["batch"],
                "batch_id": batch_obj_id,
                "college": college,
                "branch": branch,
                "placed": validated_data.get("placed", False),
                "company_name": validated_data.get("company_name"),
                "company_category": validated_data.get("company_category"),
                "package_lpa": validated_data.get("package_lpa"),
                "ctc_package": validated_data.get("ctc_package"), # Standardized field
                "placement_type": validated_data.get("placement_type"),
                "job_role": validated_data.get("job_role"),
                "date_of_offer": validated_data.get("date_of_offer"),
                "higher_studies": validated_data.get("higher_studies", False),
                "field_of_study": validated_data.get("field_of_study"),
                "university_name": validated_data.get("university_name"),
                "country": validated_data.get("country"),
                "entrance_exams": validated_data.get("entrance_exams"),
                "entrepreneurship": validated_data.get("entrepreneurship", False),
                "startup_name": validated_data.get("startup_name"),
                "startup_domain": validated_data.get("startup_domain"),
                "not_placed": validated_data.get("not_placed", False),
                "not_interested": validated_data.get("not_interested", False),
                "created_at": datetime.now(timezone.utc),
                "updated_at": datetime.now(timezone.utc)
            }
            
            # Upsert based on student_id and batch_id (update if exists, insert if new)
            db.placement_records.update_one(
                {"student_id": validated_data["student_id"], "batch_id": batch_obj_id},
                {"$set": doc},
                upsert=True
            )
            inserted_count += 1
        
        return jsonify({
            "message": f"Successfully inserted/updated {inserted_count} records",
            "inserted_count": inserted_count,
            "errors": errors
        }), 200
        
    except Exception as e:
        return jsonify({"message": f"Error committing records: {str(e)}"}), 500

@app.route('/api/placement/summary', methods=['GET', 'OPTIONS'])
def get_placement_summary():
    """Get aggregated summary for dashboard."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    batch = request.args.get('batch')
    batch_id = request.args.get('batch_id')
    
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
             return jsonify({"message": "User college or branch not found in session"}), 400

        # Build query
        query = {"college": college, "branch": branch}
        if batch:
            query["batch"] = batch
        if batch_id:
            try:
                batch_obj_id = ObjectId(batch_id)
                query["batch_id"] = batch_obj_id
            except Exception:
                pass
        
        # Get total students (from dashboard summary or count all records)
        total_students = db.placement_records.count_documents(query)
        
        if total_students == 0:
            return jsonify({
                "total_students": 0,
                "placed_count": 0,
                "placed_percentage": 0,
                "higher_studies_count": 0,
                "higher_studies_percentage": 0,
                "entrepreneurship_count": 0,
                "not_placed_count": 0,
                "not_interested_count": 0,
                "by_company_category": {},
                "by_placement_type": {},
                "by_field_of_study": {},
                "average_package": 0,
                "max_package": 0,
                "min_package": 0
            }), 200
        
        # Aggregate statistics
        pipeline = [
            {"$match": query},
            {"$group": {
                "_id": None,
                "placed_count": {"$sum": {"$cond": ["$placed", 1, 0]}},
                "higher_studies_count": {"$sum": {"$cond": ["$higher_studies", 1, 0]}},
                "entrepreneurship_count": {"$sum": {"$cond": ["$entrepreneurship", 1, 0]}},
                "not_placed_count": {"$sum": {"$cond": ["$not_placed", 1, 0]}},
                "not_interested_count": {"$sum": {"$cond": ["$not_interested", 1, 0]}},
                "packages": {"$push": "$package_lpa"},
                "by_category": {"$push": "$company_category"},
                "by_placement_type": {"$push": "$placement_type"},
                "by_field": {"$push": "$field_of_study"}
            }}
        ]
        
        result = list(db.placement_records.aggregate(pipeline))
        
        if not result:
            return jsonify({
                "total_students": total_students,
                "placed_count": 0,
                "placed_percentage": 0,
                "higher_studies_count": 0,
                "higher_studies_percentage": 0,
                "entrepreneurship_count": 0,
                "not_placed_count": 0,
                "not_interested_count": 0,
                "by_company_category": {},
                "by_placement_type": {},
                "by_field_of_study": {},
                "average_package": 0,
                "max_package": 0,
                "min_package": 0
            }), 200
        
        summary = result[0]
        placed_count = summary.get("placed_count", 0)
        higher_studies_count = summary.get("higher_studies_count", 0)
        
        # Calculate percentages
        placed_percentage = round((placed_count / total_students) * 100, 2) if total_students > 0 else 0
        higher_studies_percentage = round((higher_studies_count / total_students) * 100, 2) if total_students > 0 else 0
        
        # Process company categories
        category_dist = {}
        for cat in summary.get("by_category", []):
            if cat:
                category_dist[cat] = category_dist.get(cat, 0) + 1
        
        # Process placement types
        type_dist = {}
        for pt in summary.get("by_placement_type", []):
            if pt:
                type_dist[pt] = type_dist.get(pt, 0) + 1
        
        # Process fields of study
        field_dist = {}
        for field in summary.get("by_field", []):
            if field:
                field_dist[field] = field_dist.get(field, 0) + 1
        
        # Calculate package statistics
        packages = [p for p in summary.get("packages", []) if p is not None and p > 0]
        avg_package = round(sum(packages) / len(packages), 2) if packages else 0
        max_package = max(packages) if packages else 0
        min_package = min(packages) if packages else 0
        
        return jsonify({
            "total_students": total_students,
            "placed_count": placed_count,
            "placed_percentage": placed_percentage,
            "higher_studies_count": higher_studies_count,
            "higher_studies_percentage": higher_studies_percentage,
            "entrepreneurship_count": summary.get("entrepreneurship_count", 0),
            "not_placed_count": summary.get("not_placed_count", 0),
            "not_interested_count": summary.get("not_interested_count", 0),
            "by_company_category": category_dist,
            "by_placement_type": type_dist,
            "by_field_of_study": field_dist,
            "average_package": avg_package,
            "max_package": max_package,
            "min_package": min_package
        }), 200
        
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"❌ Error in placement summary: {str(e)}")
        print(f"Traceback: {error_trace}")
        return jsonify({"message": f"Error generating summary: {str(e)}"}), 500

@app.route('/api/placement/records', methods=['GET', 'OPTIONS'])
def get_placement_records():
    """Get detailed placement records."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    batch = request.args.get('batch')
    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 50))
    
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
             return jsonify({"message": "User college or branch not found in session"}), 400

        # Run cleanup for this user's context
        cleanup_pending_deletes(college, branch)

        query = {"college": college, "branch": branch}
        query["pending_delete"] = {"$ne": True}
        
        if batch:
            query["batch"] = batch
        
        skip = (page - 1) * per_page
        
        # Get total count
        total = db.placement_records.count_documents(query)
        
        # Get records
        records = list(
            db.placement_records
            .find(query)
            .sort("student_name", 1)
            .skip(skip)
            .limit(per_page)
        )
        
        # Format records
        formatted_records = []
        for doc in records:
            record = {
                "id": str(doc["_id"]),
                "student_id": doc.get("student_id"),
                "student_name": doc.get("student_name"),
                "batch": doc.get("batch"),
                "placed": doc.get("placed", False),
                "company_name": doc.get("company_name"),
                "company_category": doc.get("company_category"),
                "package_lpa": doc.get("package_lpa"),
                "placement_type": doc.get("placement_type"),
                "job_role": doc.get("job_role"),
                "date_of_offer": doc.get("date_of_offer"),
                "higher_studies": doc.get("higher_studies", False),
                "field_of_study": doc.get("field_of_study"),
                "university_name": doc.get("university_name"),
                "country": doc.get("country"),
                "entrance_exams": doc.get("entrance_exams"),
                "entrepreneurship": doc.get("entrepreneurship", False),
                "startup_name": doc.get("startup_name"),
                "startup_domain": doc.get("startup_domain"),
                "not_placed": doc.get("not_placed", False),
                "not_interested": doc.get("not_interested", False),
                "created_at": doc.get("created_at").isoformat() if doc.get("created_at") else None,
                "updated_at": doc.get("updated_at").isoformat() if doc.get("updated_at") else None
            }
            formatted_records.append(record)
        
        return jsonify({
            "records": formatted_records,
            "total": total,
            "page": page,
            "per_page": per_page,
            "total_pages": (total + per_page - 1) // per_page
        }), 200
        
    except Exception as e:
        return jsonify({"message": f"Error fetching records: {str(e)}"}), 500

# --- Aggregate Placement & Higher Education Outcomes Module APIs ---

# Initialize indexes for placement_outcomes collection
def init_placement_outcomes_indexes():
    """Create indexes for placement_outcomes collection."""
    try:
        db.placement_outcomes.create_index([("outcome_type", 1)])
        db.placement_outcomes.create_index([("year", 1)])
        db.placement_outcomes.create_index([("created_at", -1)])
        # Silent on success
        pass
    except Exception as e:
        print(f"⚠️ Warning: Could not create placement outcomes indexes: {e}")

# Initialize indexes on startup
init_placement_outcomes_indexes()

# Validation constants
OUTCOME_TYPES = ["Placed", "Higher Studies", "Entrepreneurship"]
PLACEMENT_TYPES = ["On-campus", "Off-campus"]
DEGREE_TYPES = ["MS", "MTech", "MBA", "PhD"]
COUNTRIES = ["India", "Abroad"]
STARTUP_STATUSES = ["Active", "Early-stage", "Closed"]

def validate_placement_outcome(data):
    """Validate an aggregate placement outcome record.
    
    Returns: (is_valid, error_message, validated_data)
    """
    errors = []
    validated = {}
    
    # Required: outcome_type
    outcome_type = str(data.get("outcome_type", "")).strip()
    if not outcome_type:
        errors.append("outcome_type is required")
    else:
        validated["outcome_type"] = outcome_type
    
    # Required: number_of_students
    try:
        num_students = int(data.get("number_of_students", 0))
        if num_students <= 0:
            errors.append("number_of_students must be a positive integer")
        else:
            validated["number_of_students"] = num_students
    except (ValueError, TypeError):
        errors.append("number_of_students must be a positive integer")
    
    # Required: year
    year_input = data.get("year")
    year_str = str(year_input).strip() if year_input is not None else ""
    
    # Normalize dashes (en-dash, em-dash) to hyphen
    year_str = year_str.replace('–', '-').replace('—', '-')

    if not year_str:
        errors.append("year is required")
    else:
        # Check if it's a simple integer year (YYYY)
        if year_str.isdigit() and len(year_str) == 4:
            y = int(year_str)
            if 2000 <= y <= 2100:
                validated["year"] = y
            else:
                errors.append("Year must be between 2000 and 2100")
        # Check for academic year format (YYYY-YY or YYYY-YYYY)
        elif '-' in year_str:
            parts = year_str.split('-')
            is_valid_format = False
            if len(parts) == 2:
                # 2023-24 or 2023-2024
                p1, p2 = parts[0].strip(), parts[1].strip()
                if p1.isdigit() and p2.isdigit():
                    if len(p1) == 4 and (len(p2) == 2 or len(p2) == 4):
                        validated["year"] = year_str
                        is_valid_format = True
            
            if not is_valid_format:
                errors.append("Invalid year format. Use YYYY, YYYY-YY, or YYYY-YYYY")
        else:
            errors.append("Invalid year format. Expected YYYY or YYYY-YY")
    
    if errors:
        return False, "; ".join(errors), None
    
    # Optional: student_names
    student_names = []
    raw_names = data.get("student_names")
    if raw_names:
        if isinstance(raw_names, str):
            student_names = [n.strip() for n in raw_names.split(',') if n.strip()]
        elif isinstance(raw_names, list):
            student_names = [str(n).strip() for n in raw_names if str(n).strip()]
    validated["student_names"] = student_names
    
    # Relaxed validation for free-text inputs - no strict type checks
    # Map all potential fields if present
    
    # Company / Job / Placed specific
    if data.get("company_name"): validated["company_name"] = str(data.get("company_name")).strip()
    if data.get("job_role"): validated["job_role"] = str(data.get("job_role")).strip()
    if data.get("location"): validated["location"] = str(data.get("location")).strip()
    
    placement_type_raw = str(data.get("placement_type", "")).strip()
    if placement_type_raw:
        validated["placement_type"] = placement_type_raw # Allow free text
    else:
        validated["placement_type"] = None
    
    validated["ctc_package"] = None
    ctc_val = data.get("ctc_package") or data.get("package")
    if ctc_val is not None:
        # Try to parse as float first
        try:
            pkg = float(ctc_val)
            if pkg >= 0:
                validated["ctc_package"] = round(pkg, 2)
        except (ValueError, TypeError):
            # Fallback to string if it's not a pure number (e.g. "10 LPA")
            validated["ctc_package"] = str(ctc_val).strip()
            
    # Higher Studies specific
    if data.get("course_name"): validated["course_name"] = str(data.get("course_name")).strip()
    if data.get("university_institute"): validated["university_institute"] = str(data.get("university_institute")).strip()
    
    degree_type_raw = str(data.get("degree_type", "")).strip()
    if degree_type_raw:
        validated["degree_type"] = degree_type_raw # Allow free text
    else:
        validated["degree_type"] = None
    
    country_raw = str(data.get("country", "")).strip()
    if country_raw:
         validated["country"] = country_raw
    else:
         validated["country"] = None
         
    if data.get("country_name"): validated["country_name"] = str(data.get("country_name")).strip()
    
    # Entrepreneurship specific
    if data.get("startup_name"): validated["startup_name"] = str(data.get("startup_name")).strip()
    if data.get("startup_status"): validated["startup_status"] = str(data.get("startup_status")).strip()
    
    if data.get("year_started"):
        try:
             validated["year_started"] = int(data.get("year_started"))
        except:
             pass
             
    # Evidence Link (for bulk uploads mostly)
    if data.get("evidence_link"):
        validated["evidence_link"] = str(data.get("evidence_link")).strip()
    
    # Basic required checks for obvious types (optional, can be removed if user wants TOTAL flexibility)
    # Keeping them loose.
    
    if outcome_type == "Placed" and (not validated.get("company_name") or not validated.get("job_role")):
        # Only enforce if type is EXACTLY "Placed". If user types "Placement", we let it slide?
        # Better to rely on UI required attributes. Backend is safety net.
        # I'll keep this check but only if it matches exactly.
        if not validated.get("company_name"): errors.append("company_name is required for Placed outcome")
        if not validated.get("job_role"): errors.append("job_role is required for Placed outcome")

    if errors:
        return False, "; ".join(errors), None
    
    return True, None, validated

@app.route('/api/placement-outcomes', methods=['POST', 'OPTIONS'])
@login_required
def add_placement_outcome():
    """Add a single aggregate placement outcome record."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    try:
        # Handle both JSON and form-data
        if request.content_type and 'application/json' in request.content_type:
            data = request.get_json() or {}
        else:
            data = request.form.to_dict()
            # Convert numeric fields from form data strings
            if 'number_of_students' in data:
                try: data['number_of_students'] = int(data['number_of_students'])
                except: pass
            if 'year' in data:
                try: data['year'] = int(data['year'])
                except: pass
            if 'ctc_package' in data:
                try: data['ctc_package'] = float(data['ctc_package'])
                except: pass
            if 'year_started' in data:
                try: data['year_started'] = int(data['year_started'])
                except: pass

        # Remove debug logs for placement data
        pass
        
        batch_id = data.get('batch_id')
        
        if not batch_id:
            return jsonify({"message": "Batch is required"}), 400
        
        # Validate batch_id and ownership
        try:
            college = session.get('college')
            branch = session.get('branch')
            
            if not college or not branch:
                 return jsonify({"message": "User college or branch not found in session"}), 400

            batch = None
            batch_obj_id = None
            
            if ObjectId.is_valid(batch_id):
                batch_obj_id = ObjectId(batch_id)
                batch = db.batches.find_one({"_id": batch_obj_id, "college": college, "branch": branch})
            
            if not batch:
                # Try finding by batch_range string
                batch = db.batches.find_one({"batch_range": batch_id, "college": college, "branch": branch})
                if batch:
                    batch_obj_id = batch["_id"]
            
            if not batch:
                return jsonify({"message": "Invalid batch ID or unauthorized access"}), 400
        except Exception as e:
            print(f"❌ Batch ID validation error: {e}")
            return jsonify({"message": "Invalid batch ID format"}), 400
        
        # Validate
        is_valid, error_msg, validated_data = validate_placement_outcome(data)
        
        if not is_valid:
            print(f"❌ Validation failed: {error_msg}")
            return jsonify({"message": error_msg}), 400
            
        # Handle evidence file upload
        evidence_file_id = None
        
        # Check for evidence file in various keys (Placement, Higher Studies, Entrepreneurship)
        file_keys = ['evidence_file', 'higher_studies_evidence_file', 'entrepreneurship_evidence_file']
        file = None
        
        for key in file_keys:
            if key in request.files:
                f = request.files[key]
                if f and f.filename != '':
                    file = f
                    break
        
        if file:
            try:
                fs = GridFS(db)
                evidence_file_id = fs.put(
                    file,
                    filename=file.filename,
                    file_type='placement_evidence',
                    uploaded_at=datetime.now(timezone.utc),
                    batch_id=batch_obj_id
                )
            except Exception as e:
                print(f"❌ Error uploading evidence file: {e}")
                    # Continue without file or fail? Let's continue but log it.
                    # Or maybe better to fail if upload fails?
                    # For now, let's just log.
        
        # Create document
        doc = {
            "outcome_type": validated_data["outcome_type"],
            "number_of_students": validated_data["number_of_students"],
            "student_names": validated_data.get("student_names", []),
            "year": validated_data["year"],
            "batch_id": batch_obj_id,
            "college": college,
            "branch": branch,
            "evidence_file_id": evidence_file_id,
            # Placed fields
            "company_name": validated_data.get("company_name"),
            "job_role": validated_data.get("job_role"),
            "location": validated_data.get("location"),
            "placement_type": validated_data.get("placement_type"),
            "ctc_package": validated_data.get("ctc_package"),
            # Higher Studies fields
            "course_name": validated_data.get("course_name"),
            "degree_type": validated_data.get("degree_type"),
            "university_institute": validated_data.get("university_institute"),
            "country": validated_data.get("country"),
            "country_name": validated_data.get("country_name"),
            # Entrepreneurship fields
            "startup_name": validated_data.get("startup_name"),
            "year_started": validated_data.get("year_started"),
            "startup_status": validated_data.get("startup_status"),
            "evidence_link": validated_data.get("evidence_link"),
            "created_at": datetime.now(timezone.utc),
            "updated_at": datetime.now(timezone.utc)
        }
        
        result = db.placement_outcomes.insert_one(doc)
        
        return jsonify({
            "message": "Outcome record added successfully",
            "id": str(result.inserted_id)
        }), 201
        
    except Exception as e:
        return jsonify({"message": f"Error adding outcome record: {str(e)}"}), 500

@app.route('/api/placement-outcomes/delete_all', methods=['POST', 'OPTIONS'])
@login_required
def delete_all_placement_outcomes():
    """Delete all placement outcome records for the current college/branch/batch."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    college = session.get('college')
    branch = session.get('branch')
    if not college or not branch:
        return jsonify({"message": "User college or branch not found"}), 400

    try:
        # Build query
        query = {"college": college, "branch": branch}
        
        # Optional: Filter by batch if provided in request body
        data = request.get_json() or {}
        batch_id = data.get('batch_id')
        batch_name = data.get('batch_name')
        
        if batch_id:
            query['batch_id'] = ObjectId(batch_id)
        elif batch_name:
            # Look up batch by name (batch_range)
            batch = db.batches.find_one({"batch_range": batch_name, "college": college, "branch": branch})
            if batch:
                query['batch_id'] = batch["_id"]
            else:
                return jsonify({"message": f"Batch '{batch_name}' not found"}), 404
            
        # Find all records to delete their files first
        records = list(db.placement_outcomes.find(query))
        
        if not records:
             return jsonify({"message": "No records found to delete", "deleted_count": 0}), 200

        # Delete associated files from GridFS
        fs = GridFS(db)
        file_deletion_errors = 0
        for record in records:
            if "evidence_file_id" in record and record["evidence_file_id"]:
                try:
                    fs.delete(ObjectId(record["evidence_file_id"]))
                except Exception:
                    file_deletion_errors += 1
        
        # Delete records
        result = db.placement_outcomes.delete_many(query)
        
        msg = f"Deleted {result.deleted_count} records"
        if file_deletion_errors > 0:
            msg += f" (with {file_deletion_errors} file deletion errors)"
            
        return jsonify({
            "message": msg,
            "deleted_count": result.deleted_count
        }), 200
        
    except Exception as e:
        return jsonify({"message": f"Error deleting records: {str(e)}"}), 500

@app.route('/api/placement-outcomes/template', methods=['GET', 'OPTIONS'])
@login_required
def download_placement_outcomes_template():
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    try:
        outcome_type = (request.args.get('outcome_type') or '').strip()
        # Removed strict OUTCOME_TYPES check to allow free text

        templates = {
            "Placed": [
                "year",
                "number_of_students",
                "student_names",
                "company_name",
                "job_role",
                "location",
                "placement_type",
                "ctc_package",
                "evidence_link"
            ],
            "Higher Studies": [
                "year",
                "number_of_students",
                "student_names",
                "course_name",
                "degree_type",
                "university_institute",
                "country",
                "country_name",
                "evidence_link"
            ],
            "Entrepreneurship": [
                "year",
                "number_of_students",
                "student_names",
                "startup_name",
                "year_started",
                "startup_status",
                "evidence_link"
            ]
        }

        columns = templates.get(outcome_type, [
            "year",
            "number_of_students",
            "student_names",
            "remarks"
        ])
        df = pd.DataFrame(columns=columns)

        output = BytesIO()
        writer = pd.ExcelWriter(output, engine='xlsxwriter')
        df.to_excel(writer, index=False, sheet_name='Template')

        workbook = writer.book
        inst_sheet = workbook.add_worksheet('Instructions')
        inst_sheet.write(0, 0, 'Placement Outcomes Bulk Template')
        inst_sheet.write(2, 0, '1) Use separate template per outcome_type.')
        inst_sheet.write(3, 0, '2) Names of students are optional; comma-separated if provided.')
        inst_sheet.write(4, 0, '3) Year and number_of_students are required.')
        writer.close()
        output.seek(0)

        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f'Placement_{outcome_type}_Template.xlsx'
        )
    except Exception as e:
        return jsonify({"message": f"Error generating template: {str(e)}"}), 500

@app.route('/api/placement-outcomes/upload-bulk', methods=['POST', 'OPTIONS'])
@login_required
def upload_placement_outcomes_bulk():
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200

    if 'file' not in request.files:
        return jsonify({"message": "No file part"}), 400

    file = request.files['file']
    batch_id = request.form.get('batch_id')
    outcome_type = (request.form.get('outcome_type') or '').strip()

    if not batch_id:
        return jsonify({"message": "Batch is required"}), 400
    if outcome_type not in OUTCOME_TYPES:
        return jsonify({"message": "Invalid or missing outcome_type for bulk upload"}), 400


    try:
        college = session.get('college')
        branch = session.get('branch')
        if not college or not branch:
            missing = []
            if not college: missing.append("college")
            if not branch: missing.append("branch")
            return jsonify({"message": f"User profile incomplete: missing {', '.join(missing)}. Please logout and login again."}), 400

        batch = None
        batch_obj_id = None
        
        if ObjectId.is_valid(batch_id):
            batch_obj_id = ObjectId(batch_id)
            batch = db.batches.find_one({"_id": batch_obj_id, "college": college, "branch": branch})
        
        if not batch:
            # Try finding by batch_range string
            batch = db.batches.find_one({"batch_range": batch_id, "college": college, "branch": branch})
            if batch:
                batch_obj_id = batch["_id"]

        if not batch:
            return jsonify({"message": "Invalid batch ID or unauthorized access"}), 400
    except Exception:
        return jsonify({"message": "Invalid batch ID format"}), 400

    if file.filename == '':
        return jsonify({"message": "No file selected"}), 400

    # Check for duplicate file (Option B: Duplicate Upload Handling)
    existing_file = db.placement_outcome_files.find_one({
        "batch_id": batch_obj_id,
        "filename": file.filename,
        "outcome_type": outcome_type
    })
    
    warning_msg = ""
    if existing_file:
        warning_msg = " Warning: Data from a file with this name may already exist. Re-uploading can cause duplicate entries."

    try:
        content = file.read()
        file.seek(0)

        filename_lower = file.filename.lower()
        if filename_lower.endswith('.csv'):
            df = pd.read_csv(BytesIO(content))
        elif filename_lower.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(BytesIO(content))
        else:
            return jsonify({"message": "Unsupported file format. Use CSV or Excel."}), 400

        df.columns = [str(c).replace('\n', ' ').strip().lower() for c in df.columns]

        valid_records = []
        invalid_records = []

        for idx, row in df.iterrows():
            row_dict = {k: (str(v).strip() if isinstance(v, str) else v) for k, v in row.to_dict().items()}
            row_dict = {k: v for k, v in row_dict.items() if pd.notna(v)}

            year_val = get_first_value(row_dict, ["year", "academic_year", "academic year", "batch"])
            payload = {"outcome_type": outcome_type}

            if year_val in (None, ""):
                invalid_records.append({
                    "row": int(idx) + 2,
                    "data": row_dict,
                    "error": "Missing required field: year"
                })
                continue

            # Pass year validation to validate_placement_outcome
            payload["year"] = year_val

            # Support aliases for number_of_students
            num_val = get_first_value(row_dict, [
                "number_of_students", "number of students", 
                "students", "count", "participants", "total_students"
            ])
            if num_val is not None:
                payload["number_of_students"] = num_val

            # Use robust student name extraction
            payload["student_names"] = get_student_names(row_dict)

            def get_val(keys):
                for k in keys:
                    if k in row_dict and pd.notna(row_dict[k]):
                        return row_dict[k]
                return None

            if outcome_type in ["Placed", "Placement"]:
                payload["company_name"] = get_val(["company_name", "company name", "company", "organization"])
                payload["job_role"] = get_val(["job_role", "job role", "role", "designation", "position"])
                payload["location"] = get_val(["location", "city", "place"])
                payload["placement_type"] = get_val(["placement_type", "placement type", "type"])
                payload["ctc_package"] = get_val(["ctc_package", "ctc package", "ctc", "package", "salary", "package_lpa", "package lpa"])
                payload["evidence_link"] = get_val(["evidence_link", "evidence link", "link", "proof"])
                
            elif outcome_type in ["Higher Studies", "HigherStudies"]:
                payload["course_name"] = get_val(["course_name", "course name", "course", "program"])
                payload["degree_type"] = get_val(["degree_type", "degree type", "degree"])
                payload["university_institute"] = get_val(["university_institute", "university institute", "university", "institute", "college"])
                payload["country"] = get_val(["country", "location"])
                payload["country_name"] = get_val(["country_name", "country name"])
                payload["evidence_link"] = get_val(["evidence_link", "evidence link", "link", "proof"])

            elif outcome_type == "Entrepreneurship":
                payload["startup_name"] = get_val(["startup_name", "startup name", "startup", "company"])
                payload["year_started"] = get_val(["year_started", "year started", "founded"])
                payload["startup_status"] = get_val(["startup_status", "startup status", "status"])
                payload["evidence_link"] = get_val(["evidence_link", "evidence link", "link", "proof"])
            else:
                # Generic handler for custom outcome types
                # Capture all other fields
                standard_keys = ["outcome_type", "year", "number_of_students", "student_names"]
                for k, v in row_dict.items():
                    if k not in standard_keys:
                        payload[k] = v

            is_valid, error_msg, validated_data = validate_placement_outcome(payload)

            if is_valid:
                valid_records.append({
                    "row": int(idx) + 2,
                    "data": validated_data
                })
            else:
                invalid_records.append({
                    "row": int(idx) + 2,
                    "data": payload,
                    "error": error_msg
                })

        fs = GridFS(db)
        grid_id = fs.put(
            BytesIO(content),
            filename=file.filename,
            file_type='placement_outcome',
            batch_id=batch_obj_id,
            year=None,
            outcome_type=outcome_type,
            uploaded_at=datetime.now(timezone.utc),
            college=college,
            branch=branch
        )

        total_records = len(valid_records) + len(invalid_records)

        db.placement_outcome_files.insert_one({
            "file_id": grid_id,
            "filename": file.filename,
            "batch_id": batch_obj_id,
            "college": college,
            "branch": branch,
            "year": None,
            "outcome_type": outcome_type,
            "uploaded_at": datetime.now(timezone.utc),
            "file_size": len(content),
            "record_count": total_records,
            "valid_count": len(valid_records),
            "invalid_count": len(invalid_records)
        })

        return jsonify({
            "message": f"Validation complete. {len(valid_records)} valid, {len(invalid_records)} invalid records.{warning_msg}",
            "valid_count": len(valid_records),
            "invalid_count": len(invalid_records),
            "valid_records": valid_records,
            "valid_records_preview": valid_records[:10],
            "invalid_records": invalid_records,
            "batch_id": batch_id,
            "year": None,
            "preview": True
        }), 200
    except Exception as e:
        return jsonify({"message": f"Error processing file: {str(e)}"}), 500

@app.route('/api/placement-outcomes/upload-commit', methods=['POST', 'OPTIONS'])
@login_required
def commit_placement_outcomes_bulk():
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200

    data = request.get_json()
    raw_records = data.get('records', [])
    if not raw_records and 'valid_records' in data:
        raw_records = data.get('valid_records', [])

    records = []
    for item in raw_records:
        if isinstance(item, dict) and 'data' in item and isinstance(item['data'], dict):
            records.append(item['data'])
        else:
            records.append(item)

    if not records:
        return jsonify({"message": "No records to commit"}), 400

    batch_id = data.get('batch_id')
    if not batch_id:
        return jsonify({"message": "Batch is required"}), 400

    try:
        college = session.get('college')
        branch = session.get('branch')
        if not college or not branch:
            missing = []
            if not college: missing.append("college")
            if not branch: missing.append("branch")
            return jsonify({"message": f"User profile incomplete: missing {', '.join(missing)}. Please logout and login again."}), 400

        batch = None
        batch_obj_id = None
        
        if ObjectId.is_valid(batch_id):
            batch_obj_id = ObjectId(batch_id)
            batch = db.batches.find_one({"_id": batch_obj_id, "college": college, "branch": branch})
        
        if not batch:
            # Try finding by batch_range string
            batch = db.batches.find_one({"batch_range": batch_id, "college": college, "branch": branch})
            if batch:
                batch_obj_id = batch["_id"]

        if not batch:
            return jsonify({"message": "Invalid batch ID or unauthorized access"}), 400
    except Exception:
        return jsonify({"message": "Invalid batch ID format"}), 400

    try:
        inserted_count = 0
        errors = []

        for record_data in records:
            is_valid, error_msg, validated_data = validate_placement_outcome(record_data)

            if not is_valid:
                errors.append({"record": record_data, "error": error_msg})
                continue

            doc = {
                "outcome_type": validated_data["outcome_type"],
                "number_of_students": validated_data["number_of_students"],
                "student_names": validated_data.get("student_names", []),
                "year": validated_data["year"],
                "batch_id": batch_obj_id,
                "college": college,
                "branch": branch,
                "company_name": validated_data.get("company_name"),
                "job_role": validated_data.get("job_role"),
                "location": validated_data.get("location"),
                "placement_type": validated_data.get("placement_type"),
                "ctc_package": validated_data.get("ctc_package"),
                "course_name": validated_data.get("course_name"),
                "degree_type": validated_data.get("degree_type"),
                "university_institute": validated_data.get("university_institute"),
                "country": validated_data.get("country"),
                "country_name": validated_data.get("country_name"),
                "startup_name": validated_data.get("startup_name"),
                "startup_domain": validated_data.get("startup_domain"),
                "year_started": validated_data.get("year_started"),
                "startup_status": validated_data.get("startup_status"),
                "evidence_link": validated_data.get("evidence_link"),
                "created_at": datetime.now(timezone.utc),
                "updated_at": datetime.now(timezone.utc)
            }

            db.placement_outcomes.insert_one(doc)
            inserted_count += 1

        return jsonify({
            "message": f"Successfully inserted {inserted_count} records",
            "inserted_count": inserted_count,
            "errors": errors
        }), 200
    except Exception as e:
        return jsonify({"message": f"Error committing records: {str(e)}"}), 500

@app.route('/api/placement-outcomes', methods=['GET', 'OPTIONS'])
@login_required
def get_placement_outcomes():
    """Get aggregate placement outcome records."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    outcome_type = request.args.get('outcome_type')
    year = request.args.get('year')
    batch_id = request.args.get('batch_id')
    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 50))
    
    try:
        query = {}

        college = session.get('college')
        branch = session.get('branch')
        if not college or not branch:
             return jsonify({"message": "User college or branch not found"}), 400
        
        # Run cleanup for this user's context
        cleanup_pending_deletes(college, branch)
        
        # Exclude soft-deleted records
        query["deleted"] = {"$ne": True}
        query["pending_delete"] = {"$ne": True}
        
        # Batch ID is optional - if provided, filter by it; if not, show all records (for data upload page)
        if batch_id:
            try:
                batch_obj_id = ObjectId(batch_id)
                # Verify batch exists and belongs to branch
                batch = db.batches.find_one({"_id": batch_obj_id, "college": college, "branch": branch})
                if not batch:
                    return jsonify({"message": "Invalid batch ID or unauthorized access"}), 400
                query["batch_id"] = batch_obj_id
            except Exception as e:
                print(f"❌ Invalid batch_id format: {batch_id}, error: {e}")
                return jsonify({"message": "Invalid batch ID format"}), 400
        else:
            # Filter by all batches for this college and branch
            branch_batches = list(db.batches.find({"college": college, "branch": branch}, {"_id": 1}))
            branch_batch_ids = [b["_id"] for b in branch_batches]
            query['batch_id'] = {"$in": branch_batch_ids}
        
        if outcome_type:
            query["outcome_type"] = outcome_type
        if year:
            try:
                query["year"] = int(year)
            except ValueError:
                pass
        
        skip = (page - 1) * per_page
        
        # Get total count
        total = db.placement_outcomes.count_documents(query)
        
        # Get records
        records = list(
            db.placement_outcomes
            .find(query)
            .sort("year", -1)
            .sort("created_at", -1)
            .skip(skip)
            .limit(per_page)
        )
        
        # Format records
        formatted_records = []
        for doc in records:
            # Get batch information
            batch_range = "N/A"
            if doc.get("batch_id"):
                try:
                    batch = db.batches.find_one({"_id": doc["batch_id"]})
                    if batch:
                        batch_range = batch.get("batch_range", "N/A")
                except Exception:
                    pass
            
            record = {
                "id": str(doc["_id"]),
                "outcome_type": doc.get("outcome_type"),
                "number_of_students": doc.get("number_of_students", 0),
                "student_names": doc.get("student_names", []),
                "year": doc.get("year"),
                "batch_range": batch_range,
                # Placed fields
                "company_name": doc.get("company_name"),
                "job_role": doc.get("job_role"),
                "location": doc.get("location"),
                "placement_type": doc.get("placement_type"),
                "ctc_package": doc.get("ctc_package"),
                # Higher Studies fields
                "course_name": doc.get("course_name"),
                "degree_type": doc.get("degree_type"),
                "university_institute": doc.get("university_institute"),
                "country": doc.get("country"),
                "country_name": doc.get("country_name"),
                # Entrepreneurship fields
                "startup_name": doc.get("startup_name"),
                "startup_domain": doc.get("startup_domain"),
                "year_started": doc.get("year_started"),
                "startup_status": doc.get("startup_status"),
                "created_at": doc.get("created_at").isoformat() if doc.get("created_at") else None,
                "updated_at": doc.get("updated_at").isoformat() if doc.get("updated_at") else None
            }
            formatted_records.append(record)
        
        return jsonify({
            "records": formatted_records,
            "total": total,
            "page": page,
            "per_page": per_page,
            "total_pages": (total + per_page - 1) // per_page
        }), 200
        
    except Exception as e:
        return jsonify({"message": f"Error fetching records: {str(e)}"}), 500

@app.route('/api/placement-outcomes/<outcome_id>', methods=['DELETE', 'OPTIONS'])
@login_required
def delete_placement_outcome(outcome_id):
    """Delete an aggregate placement outcome record by ID."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    try:
        college = session.get('college')
        branch = session.get('branch')
        if not college or not branch:
            return jsonify({"message": "User college or branch not found"}), 400

        # Validate ObjectId format
        try:
            object_id = ObjectId(outcome_id)
        except Exception:
            return jsonify({"message": "Invalid record ID format"}), 400
        
        # Find the record first to verify ownership
        record = db.placement_outcomes.find_one({"_id": object_id})
        if not record:
            return jsonify({"message": "Record not found"}), 404
            
        # Verify ownership via batch
        if record.get('college') != college or record.get('branch') != branch:
            # Fallback: check batch_id if college/branch not directly on record
            batch_id = record.get('batch_id')
            if batch_id:
                batch = db.batches.find_one({"_id": batch_id, "college": college, "branch": branch})
                if not batch:
                     return jsonify({"message": "Unauthorized access to this record"}), 403
            else:
                 return jsonify({"message": "Unauthorized access to this record"}), 403
        
        # Check for finalize flag
        finalize = (request.args.get('finalize') or '').lower() == 'true'
        
        if not finalize:
            # Soft delete - Pending Delete State
            db.placement_outcomes.update_one(
                {"_id": object_id},
                {"$set": {
                    "pending_delete": True,
                    "delete_requested_at": datetime.now(timezone.utc)
                }}
            )
            return jsonify({"message": "Record marked for deletion", "pending_delete": True}), 200
        else:
            # Permanent delete
            # Delete associated files if any (though placement outcomes currently don't have direct file links in the same way, 
            # or if they do, we should clean them up. Assuming standard delete for now.)
            result = db.placement_outcomes.delete_one({"_id": object_id})
            
            if result.deleted_count == 0:
                return jsonify({"message": "Record not found"}), 404
            
            return jsonify({"message": "Record permanently deleted", "finalized": True}), 200
        
    except Exception as e:
        return jsonify({"message": f"Error deleting record: {str(e)}"}), 500

@app.route('/api/placement-outcomes/summary', methods=['GET', 'OPTIONS'])
@login_required
def get_placement_outcomes_summary():
    """Get aggregated summary for dashboard."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    batch_id = request.args.get('batch_id')
    
    try:
        college = session.get('college')
        branch = session.get('branch')
        if not college or not branch:
            return jsonify({"message": "User college or branch not found"}), 400

        # Run cleanup for this user's context
        cleanup_pending_deletes(college, branch)

        # Build match stage for batch filtering - batch_id is required for dashboard
        match_stage = {}
        if batch_id:
            try:
                batch_obj_id = ObjectId(batch_id)
                # Verify batch exists and belongs to user
                batch = db.batches.find_one({"_id": batch_obj_id, "college": college, "branch": branch})
                if not batch:
                    return jsonify({
                        "placed": {"total_students": 0, "count": 0},
                        "higher_studies": {"total_students": 0, "count": 0},
                        "entrepreneurship": {"total_students": 0, "count": 0}
                    }), 200
                match_stage["batch_id"] = batch_obj_id
                # Exclude soft-deleted records
                match_stage["deleted"] = {"$ne": True}
                match_stage["pending_delete"] = {"$ne": True}
            except Exception as e:
                print(f"❌ Invalid batch_id format in summary: {batch_id}, error: {e}")
                return jsonify({
                    "placed": {"total_students": 0, "count": 0},
                    "higher_studies": {"total_students": 0, "count": 0},
                    "entrepreneurship": {"total_students": 0, "count": 0}
                }), 200
        else:
            # If no batch_id provided for dashboard summary, return empty (dashboard always provides batch_id)
            # Or could return summary for all user's batches? Dashboard seems to be batch-specific.
            return jsonify({
                "placed": {"total_students": 0, "count": 0},
                "higher_studies": {"total_students": 0, "count": 0},
                "entrepreneurship": {"total_students": 0, "count": 0}
            }), 200
        
        # Aggregate by outcome type
        pipeline = []
        if match_stage:
            pipeline.append({"$match": match_stage})
        pipeline.append({
            "$group": {
                "_id": "$outcome_type",
                "total_students": {"$sum": "$number_of_students"},
                "count": {"$sum": 1}
            }
        })
        
        result = list(db.placement_outcomes.aggregate(pipeline))
        
        summary = {
            "placed": {"total_students": 0, "count": 0},
            "higher_studies": {"total_students": 0, "count": 0},
            "entrepreneurship": {"total_students": 0, "count": 0}
        }
        
        for item in result:
            outcome_type = item.get("_id", "").lower().replace(" ", "_")
            if outcome_type in summary:
                summary[outcome_type] = {
                    "total_students": item.get("total_students", 0),
                    "count": item.get("count", 0)
                }
        
        # Get package statistics for placed
        placed_query = {
            "outcome_type": "Placed", 
            "ctc_package": {"$exists": True, "$ne": None},
            "deleted": {"$ne": True},
            "pending_delete": {"$ne": True}
        }
        if batch_id:
            try:
                batch_obj_id = ObjectId(batch_id)
                placed_query["batch_id"] = batch_obj_id
            except Exception:
                pass
        placed_records = list(db.placement_outcomes.find(
            placed_query,
            {"ctc_package": 1, "number_of_students": 1}
        ))
        
        total_package_weighted = 0
        total_students_placed = 0
        packages = []
        import re
        for record in placed_records:
            pkg = record.get("ctc_package")
            students = record.get("number_of_students", 0)
            
            numeric_pkg = None
            if isinstance(pkg, (int, float)):
                numeric_pkg = pkg
            elif isinstance(pkg, str):
                try:
                    # Try to extract the first number found in the string (e.g. "5.5 LPA" -> 5.5)
                    match = re.search(r"(\d+(\.\d+)?)", pkg)
                    if match:
                        numeric_pkg = float(match.group(1))
                except (ValueError, TypeError):
                    pass
            
            if numeric_pkg is not None and numeric_pkg > 0 and students > 0:
                total_package_weighted += numeric_pkg * students
                total_students_placed += students
                packages.append(numeric_pkg)
        
        avg_package = round(total_package_weighted / total_students_placed, 2) if total_students_placed > 0 else 0
        max_package = max(packages) if packages else 0
        min_package = min(packages) if packages else 0
        
        summary["placed"]["average_package"] = avg_package
        summary["placed"]["max_package"] = max_package
        summary["placed"]["min_package"] = min_package
        
        return jsonify(summary), 200
        
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"❌ Error in placement outcomes summary: {str(e)}")
        print(f"Traceback: {error_trace}")
        return jsonify({"message": f"Error generating summary: {str(e)}"}), 500

# --- Batch Management APIs ---

def validate_batch_range(batch_range):
    """Validate batch range format: must be YYYY–YYYY with exactly 4 years difference."""
    if not batch_range or not isinstance(batch_range, str):
        return False, "Batch range must be a string"
    
    # Check format: YYYY–YYYY or YYYY-YYYY (en dash, em dash, or hyphen), allowing spaces
    import re
    # Allow optional spaces around the separator
    pattern = r'^(\d{4})\s*[–— -]\s*(\d{4})$'
    match = re.match(pattern, batch_range.strip())
    if not match:
        return False, "Batch range must be in format YYYY-YYYY (e.g., 2022-2026)"
    
    start_year = int(match.group(1))
    end_year = int(match.group(2))
    
    # Check year difference is exactly 4
    if end_year - start_year != 4:
        return False, "Batch range must span exactly 4 years (e.g., 2022-2026)"
    
    # Check years are reasonable (e.g., 2000-2100)
    if start_year < 2000 or start_year > 2100 or end_year < 2000 or end_year > 2100:
        return False, "Years must be between 2000 and 2100"
    
    return True, None

@app.route('/api/batches', methods=['GET', 'OPTIONS'])
@login_required
def get_batches():
    """Get all batches for the logged-in user's college and branch."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200

    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
            print(f"❌ get_batches: User profile incomplete in session: {session}")
            missing = []
            if not college: missing.append("college")
            if not branch: missing.append("branch")
            return jsonify({"message": f"User profile incomplete: missing {', '.join(missing)}. Please logout and login again."}), 400

        # Cleanup pending deletes
        cleanup_pending_deletes(college, branch)

        batches = list(db.batches.find({"college": college, "branch": branch}).sort("batch_range", 1))
        batches_list = []
        for batch in batches:
            batches_list.append({
                "id": batch.get("_id"),
                "batch_range": batch.get("batch_range"),
                "branch": batch.get("branch"),
                "college": batch.get("college"),
                "created_at": batch.get("created_at")
            })
        batches_list = convert_objectid_to_str(batches_list)
        return jsonify(batches_list), 200
    except Exception as e:
        return jsonify({"message": f"Error fetching batches: {str(e)}"}), 500

@app.route('/api/batches', methods=['POST', 'OPTIONS'])
@login_required
def create_batch():
    """Create a new batch for the logged-in user's college and branch."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    data = request.get_json()
    batch_range = data.get('batch_range', '').strip()
    
    if not batch_range:
        return jsonify({"message": "Batch range is required"}), 400
    
    # Validate format
    is_valid, error_msg = validate_batch_range(batch_range)
    if not is_valid:
        print(f"❌ Invalid batch range format: '{batch_range}'. Error: {error_msg}")
        return jsonify({"message": error_msg}), 400
    
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        # Log creation quietly
        pass
        
        if not college or not branch:
            missing = []
            if not college: missing.append("college")
            if not branch: missing.append("branch")
            msg = f"User profile incomplete: missing {', '.join(missing)}. Please logout and login again."
            print(f"❌ {msg}")
            return jsonify({"message": msg}), 400

        # Check if batch already exists for this college + branch
        existing = db.batches.find_one({
            "batch_range": batch_range, 
            "college": college, 
            "branch": branch
        })
        if existing:
            return jsonify({"message": "Batch with this range already exists for your college and branch"}), 400
        
        # Create batch
        batch_doc = {
            "batch_range": batch_range,
            "college": college,
            "branch": branch,
            "created_by": session.get('email'),
            "user_id": session.get('user_id'),
            "created_at": datetime.now(timezone.utc)
        }
        result = db.batches.insert_one(batch_doc)
        
        return jsonify({
            "message": "Batch created successfully",
            "id": str(result.inserted_id),
            "batch_range": batch_range,
            "college": college,
            "branch": branch
        }), 201
    except Exception as e:
        return jsonify({"message": f"Error creating batch: {str(e)}"}), 500

@app.route('/api/batches/<batch_id>', methods=['DELETE', 'OPTIONS'])
@login_required
def delete_batch(batch_id):
    """Delete a batch and all related data."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    try:
        college = session.get('college')
        branch = session.get('branch')
        
        if not college or not branch:
            missing = []
            if not college: missing.append("college")
            if not branch: missing.append("branch")
            return jsonify({"message": f"User profile incomplete: missing {', '.join(missing)}. Please logout and login again."}), 400

        # Find batch to get batch_range for data deletion
        batch = db.batches.find_one({"_id": ObjectId(batch_id)})
        if not batch:
            return jsonify({"message": "Batch not found"}), 404
        
        # Verify ownership
        if batch.get('college') != college or batch.get('branch') != branch:
            return jsonify({"message": "Unauthorized to delete this batch"}), 403
        
        batch_range = batch.get("batch_range")

        
        # Delete all related data
        # Note: MongoDB doesn't support CASCADE, so we delete manually
        fs = GridFS(db)
        
        # Delete GridFS files for result_files before deleting metadata
        result_files_deleted = 0
        for file_doc in db.result_files.find({"batch_id": ObjectId(batch_id)}):
            try:
                fs.delete(file_doc.get('file_id'))
                result_files_deleted += 1
            except Exception as e:
                print(f"⚠️ Warning: Could not delete GridFS file for result_file {file_doc.get('_id')}: {e}")
        
        # Delete GridFS files for intake_files before deleting metadata
        intake_files_deleted = 0
        for file_doc in db.intake_files.find({"batch_id": ObjectId(batch_id)}):
            try:
                if file_doc.get('file_id'):
                    fs.delete(file_doc.get('file_id'))
                    intake_files_deleted += 1
            except Exception as e:
                print(f"⚠️ Warning: Could not delete GridFS file for intake_file {file_doc.get('_id')}: {e}")
        
        # Delete GridFS files for extracurricular_files before deleting metadata
        extracurricular_files_deleted = 0
        for file_doc in db.extracurricular_files.find({"batch_id": ObjectId(batch_id)}):
            try:
                if file_doc.get('file_id'):
                    fs.delete(file_doc.get('file_id'))
                    extracurricular_files_deleted += 1
            except Exception as e:
                print(f"⚠️ Warning: Could not delete GridFS file for extracurricular_file {file_doc.get('_id')}: {e}")
        
        # Delete GridFS files for extracurricular records (evidence files)
        for record in db.extracurricular_records.find({"batch_id": ObjectId(batch_id)}):
            file_ids = []
            if "evidence_file_ids" in record and record["evidence_file_ids"]:
                file_ids = record["evidence_file_ids"]
            elif "evidence_file_id" in record and record["evidence_file_id"]:
                file_ids = [record["evidence_file_id"]]
            for file_id_str in file_ids:
                try:
                    fs.delete(ObjectId(file_id_str))
                except Exception as e:
                    print(f"⚠️ Warning: Could not delete GridFS evidence file {file_id_str}: {e}")
        
        # Now delete metadata records
        deleted_counts = {
            "result_files": db.result_files.delete_many({"batch_id": ObjectId(batch_id)}).deleted_count,
            "intake_files": db.intake_files.delete_many({"batch_id": ObjectId(batch_id)}).deleted_count,
            "extracurricular_files": db.extracurricular_files.delete_many({"batch_id": ObjectId(batch_id)}).deleted_count,
            "extracurricular_records": db.extracurricular_records.delete_many({"batch_id": ObjectId(batch_id)}).deleted_count,
            "placement_records": db.placement_records.delete_many({"batch_id": ObjectId(batch_id)}).deleted_count,
            "placement_outcomes": db.placement_outcomes.delete_many({"batch_id": ObjectId(batch_id)}).deleted_count
        }
        
        # Delete batch itself
        result = db.batches.delete_one({"_id": ObjectId(batch_id)})
        
        if result.deleted_count == 0:
            return jsonify({"message": "Batch not found"}), 404
        
        return jsonify({
            "message": "Batch and all related data deleted successfully",
            "deleted_counts": deleted_counts
        }), 200
    except Exception as e:
        return jsonify({"message": f"Error deleting batch: {str(e)}"}), 500

# --- Migration Endpoint ---

@app.route('/api/migrate-data-to-batch', methods=['POST', 'OPTIONS'])
@login_required
def migrate_data_to_batch():
    """Migrate existing data to a specific batch (for initial setup)."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    college = session.get('college')
    branch = session.get('branch')
    if not college or not branch:
        missing = []
        if not college: missing.append("college")
        if not branch: missing.append("branch")
        return jsonify({"message": f"User profile incomplete: missing {', '.join(missing)}. Please logout and login again."}), 400

    data = request.get_json()
    batch_range = data.get('batch_range', '2022-2026')
    
    try:
        # Find or create the batch for this college/branch
        batch = db.batches.find_one({"batch_range": batch_range, "college": college, "branch": branch})
        if not batch:
            # Create the batch
            batch_doc = {
                "batch_range": batch_range,
                "college": college,
                "branch": branch,
                "created_at": datetime.now(timezone.utc)
            }
            result = db.batches.insert_one(batch_doc)
            batch_id = result.inserted_id
        else:
            batch_id = batch["_id"]
        
        # Update all existing records to have this batch_id
        # STRICT ISOLATION: Only update records belonging to this college/branch
        updated_counts = {
            "result_files": db.result_files.update_many(
                {"batch_id": {"$exists": False}, "college": college, "branch": branch},
                {"$set": {"batch_id": batch_id}}
            ).modified_count,
            "intake_files": db.intake_files.update_many(
                {"batch_id": {"$exists": False}, "college": college, "branch": branch},
                {"$set": {"batch_id": batch_id}}
            ).modified_count,
            "extracurricular_records": db.extracurricular_records.update_many(
                {"batch_id": {"$exists": False}, "college": college, "branch": branch},
                {"$set": {"batch_id": batch_id}}
            ).modified_count,
            "placement_records": db.placement_records.update_many(
                {"batch_id": {"$exists": False}, "college": college, "branch": branch},
                {"$set": {"batch_id": batch_id}}
            ).modified_count,
            "placement_outcomes": db.placement_outcomes.update_many(
                {"batch_id": {"$exists": False}, "college": college, "branch": branch},
                {"$set": {"batch_id": batch_id}}
            ).modified_count
        }
        
        # Also update dashboard_summary to be batch-specific
        # Note: We can't easily migrate "summary" (global) to a specific batch/college without knowing which one it belongs to.
        # Assuming legacy summary is abandoned or manually handled. 
        # But if we want to try, we should check if it belongs to us.
        # Since legacy summary didn't have college/branch, we skip it to avoid conflicts.
        
        return jsonify({
            "message": f"Migration completed. Data assigned to batch {batch_range}",
            "batch_id": str(batch_id),
            "batch_range": batch_range,
            "updated_counts": updated_counts
        }), 200
        
    except Exception as e:
        return jsonify({"message": f"Error during migration: {str(e)}"}), 500

# --- Reports Module APIs ---

def calculate_nba_metrics(batch_obj_id, college=None, branch=None):
    """
    Calculate NBA metrics for a specific batch.
    Returns a dictionary with all metrics or None if batch_id is invalid.
    """
    try:
        # Validate batch exists
        if college and branch:
             batch = db.batches.find_one({"_id": batch_obj_id, "college": college, "branch": branch})
        else:
             batch = db.batches.find_one({"_id": batch_obj_id})
             
        if not batch:
            return None
        
        batch_id = str(batch_obj_id)
        metrics = {}
        
        # Fetch batch-specific intake data
        summary_id = f"summary_{batch_id}"
        intake_doc = db.dashboard_summary.find_one({"_id": summary_id}) or {}
        kpis = intake_doc.get("kpis", {})
        
        # Fetch academic data
        semesters_data = intake_doc.get("semesters", [])
        
        # Fetch placement data
        query = {"batch_id": batch_obj_id}
        if college: query["college"] = college
        if branch: query["branch"] = branch
        placement_outcomes = list(db.placement_outcomes.find(query))
        
        # 1. Enrolment Ratio = (Total Admitted / Sanctioned Intake) * 100
        sanctioned_intake = kpis.get("sanctioned_intake", 0) or 0
        total_admitted = kpis.get("total_admitted", 0) or 0
        if sanctioned_intake > 0:
            metrics["enrolment_ratio"] = round((total_admitted / sanctioned_intake) * 100, 2)
        else:
            metrics["enrolment_ratio"] = None
        
        # 2. Success Index (Without Backlogs) = (Graduated without backlog / Total admitted)
        # CORRECT NBA Logic: Count unique students who completed program without backlog in any semester
        # Numerator: Unique students who completed (Sem 8) - each student counted only once
        # Denominator: Total admitted (first year + lateral entry + separate division)
        graduated_without_backlog = 0.0
        
        try:
            # Get master DataFrame to evaluate students at program level
            master, _ = build_master_and_summary(batch_id=batch_id, college=college, branch=branch)
            if master is not None and not master.empty:
                # Find students who completed the program (have data in Sem 8)
                if "Sem8" in master.columns:
                    # Get unique students who completed (have non-NA value in Sem8)
                    completed_students = master[master["Sem8"].notna()]
                    
                    if not completed_students.empty:
                        # Count unique students who completed
                        # Each student is counted only once (DISTINCT)
                        unique_completed = completed_students["Name"].nunique()
                        graduated_without_backlog = float(unique_completed)
                        # Remove debug logs for Success Index
                        pass
                    else:
                        # Silent on failure
                        pass
                else:
                    print(f"⚠️ Warning: Sem8 column not found in master DataFrame")
            else:
                print(f"⚠️ Warning: Master DataFrame is empty or None")
        except Exception as e:
            print(f"⚠️ Warning: Could not calculate unique students without backlog: {e}")
            import traceback
            traceback.print_exc()
        
        # Denominator: Total admitted (first year + lateral entry + separate division)
        # This is already available as total_admitted from intake KPIs
        if total_admitted > 0:
            # Apply floating-point division, round only at final step
            raw_result = float(graduated_without_backlog) / float(total_admitted)
            metrics["success_index_without_backlogs"] = round(raw_result, 2)
            # Log success quietly
            pass
        else:
            metrics["success_index_without_backlogs"] = None
            print(f"⚠️ Warning: total_admitted is 0, cannot calculate Success Index")
        
        # 3. Success Index (Stipulated Period) = Students who completed in stipulated period / Total students
        # For now, use the same calculation as without backlogs (can be refined based on specific requirements)
        # This typically means students who passed all semesters without repeating years
        # Using final semester data if available, otherwise same as without backlogs
        final_sem_data = None
        for sem_data in semesters_data:
            sem_name = sem_data.get("semester", "").lower()
            if sem_name in ['sem8', 'sem7']:  # Check last semesters
                if not final_sem_data or sem_name == 'sem8':
                    final_sem_data = sem_data
        
        if final_sem_data:
            final_total = final_sem_data.get("total", 0) or 0
            final_without_backlog = final_sem_data.get("without_backlog", 0) or 0
            if final_total > 0:
                metrics["success_index_stipulated_period"] = round(final_without_backlog / final_total, 2)
            else:
                metrics["success_index_stipulated_period"] = None
        else:
            # Fallback to without backlogs calculation
            metrics["success_index_stipulated_period"] = metrics["success_index_without_backlogs"]
        
        # 4. Academic Performance Index (API) = Average CGPA across all semesters
        cgpa_values = []
        for sem_data in semesters_data:
            avg_cgpa = sem_data.get("avg_cgpa")
            if avg_cgpa is not None:
                try:
                    cgpa_float = float(avg_cgpa)
                    if cgpa_float > 0:
                        cgpa_values.append(cgpa_float)
                except (ValueError, TypeError):
                    pass
        
        if cgpa_values:
            metrics["academic_performance_index"] = round(sum(cgpa_values) / len(cgpa_values), 2)
        else:
            metrics["academic_performance_index"] = None
        
        # 4a. Second Year API = Average of Sem3 and Sem4
        sem3_cgpa = None
        sem4_cgpa = None
        for sem_data in semesters_data:
            sem_name = sem_data.get("semester", "").lower()
            avg_cgpa = sem_data.get("avg_cgpa")
            if sem_name == "sem3" and avg_cgpa is not None:
                try:
                    sem3_cgpa = float(avg_cgpa)
                except (ValueError, TypeError):
                    pass
            elif sem_name == "sem4" and avg_cgpa is not None:
                try:
                    sem4_cgpa = float(avg_cgpa)
                except (ValueError, TypeError):
                    pass
        
        if sem3_cgpa is not None and sem4_cgpa is not None:
            metrics["second_year_api"] = round((sem3_cgpa + sem4_cgpa) / 2, 2)
        elif sem3_cgpa is not None:
            metrics["second_year_api"] = round(sem3_cgpa, 2)
        elif sem4_cgpa is not None:
            metrics["second_year_api"] = round(sem4_cgpa, 2)
        else:
            metrics["second_year_api"] = None
        
        # 4b. Third Year API = Average of Sem5 and Sem6
        sem5_cgpa = None
        sem6_cgpa = None
        for sem_data in semesters_data:
            sem_name = sem_data.get("semester", "").lower()
            avg_cgpa = sem_data.get("avg_cgpa")
            if sem_name == "sem5" and avg_cgpa is not None:
                try:
                    sem5_cgpa = float(avg_cgpa)
                except (ValueError, TypeError):
                    pass
            elif sem_name == "sem6" and avg_cgpa is not None:
                try:
                    sem6_cgpa = float(avg_cgpa)
                except (ValueError, TypeError):
                    pass
        
        if sem5_cgpa is not None and sem6_cgpa is not None:
            metrics["third_year_api"] = round((sem5_cgpa + sem6_cgpa) / 2, 2)
        elif sem5_cgpa is not None:
            metrics["third_year_api"] = round(sem5_cgpa, 2)
        elif sem6_cgpa is not None:
            metrics["third_year_api"] = round(sem6_cgpa, 2)
        else:
            metrics["third_year_api"] = None
        
        # 5. Placement Index = (Placed Only) / Total Eligible Students
        # User Instruction: Count ONLY records from the placement module.
        # Do NOT include internships, higher studies, or entrepreneurship.
        
        placed_count = 0.0
        
        for outcome in placement_outcomes:
            outcome_type = outcome.get("outcome_type", "").lower()
            num_students = outcome.get("number_of_students", 0) or 0
            # Convert to float to ensure floating-point arithmetic
            num_students = float(num_students) if num_students else 0.0
            
            if outcome_type == "placed":
                placed_count += num_students
                
        # Calculate numerator (no rounding here) - ensure all are floats
        total_placed_outcomes = float(placed_count)
        
        # Total eligible students = Total Final Year Students (students who appeared in Sem 7 OR Sem 8)
        # Get unique count of students from Sem 7 and Sem 8
        total_eligible = 0.0
        sem7_data = None
        sem8_data = None
        
        for sem_data in semesters_data:
            sem_name = sem_data.get("semester", "").lower()
            if sem_name == "sem7":
                sem7_data = sem_data
                # Log data found quietly
                pass
            elif sem_name == "sem8":
                sem8_data = sem_data
                # Log data found quietly
                pass
        
        if not sem7_data and not sem8_data:
            print(f"⚠️ WARNING: No Sem 7 or Sem 8 data found in semesters_data. Available semesters: {[s.get('semester') for s in semesters_data]}")
        
        # Get unique student count: students who appeared in Sem 7 OR Sem 8
        # CRITICAL: Must use final year students (Sem 7 OR Sem 8), NOT total_admitted
        total_eligible = 0.0
        
        # First priority: Use Sem 8 total (most recent, represents final year students)
        # Only use if total > 0 (data actually exists)
        if sem8_data:
            sem8_total = float(sem8_data.get("total", 0) or 0)
            if sem8_total > 0:
                total_eligible = sem8_total
                # Log data usage quietly
                pass
        
        # Second priority: Use Sem 7 total if Sem 8 not available or has 0 total
        if total_eligible == 0.0 and sem7_data:
            sem7_total = float(sem7_data.get("total", 0) or 0)
            if sem7_total > 0:
                total_eligible = sem7_total
                # Log data usage quietly
                pass
        
        # Try to get unique count from master DataFrame for more accuracy (if both Sem 7 and Sem 8 exist with data)
        if sem7_data and sem8_data and total_eligible > 0:
            try:
                # Get master DataFrame to count unique students
                master, _ = build_master_and_summary(batch_id=batch_id)
                if master is not None and not master.empty and "Sem7" in master.columns and "Sem8" in master.columns:
                    # Count students who appeared in Sem 7 OR Sem 8
                    sem7_students = set(master[master["Sem7"].notna()]["Name"].unique())
                    sem8_students = set(master[master["Sem8"].notna()]["Name"].unique())
                    unique_final_year_students = len(sem7_students | sem8_students)  # Union
                    if unique_final_year_students > 0:
                        total_eligible = float(unique_final_year_students)
                        # Log unique calculation quietly
                        pass
            except Exception as e:
                print(f"⚠️ Warning: Could not get unique count from master, using semester total: {e}")
        
        # If Sem 7/8 not available OR have total=0, use the most recent semester available as proxy
        if total_eligible == 0.0:
            # Find the highest semester number available with total > 0
            latest_sem_data = None
            latest_sem_num = 0
            for sem_data in semesters_data:
                sem_name = sem_data.get("semester", "").lower()
                sem_total = float(sem_data.get("total", 0) or 0)
                if sem_name.startswith("sem") and sem_total > 0:  # Only consider semesters with actual data
                    try:
                        sem_num = int(sem_name.replace("sem", ""))
                        if sem_num > latest_sem_num:
                            latest_sem_num = sem_num
                            latest_sem_data = sem_data
                    except:
                        pass
            
            if latest_sem_data and latest_sem_num >= 5:  # Use if Sem 5 or higher
                total_eligible = float(latest_sem_data.get("total", 0) or 0)
                if total_eligible > 0:
                    # Log proxy usage quietly
                    pass
        
        # Last resort: Only use total_admitted if NO semester data exists at all
        if total_eligible == 0.0:
            print(f"⚠️ WARNING: No semester data with total > 0 found, falling back to total_admitted (this may be incorrect for Placement Index)")
            total_eligible = float(total_admitted) if total_admitted else 0.0
            # Log fallback usage quietly
            pass
        
        # Apply rounding ONLY at the final step - use explicit float division
        if total_eligible > 0:
            # Ensure floating-point division - no intermediate rounding
            raw_result = float(total_placed_outcomes) / float(total_eligible)
            metrics["placement_index"] = round(raw_result, 2)
            # Comprehensive debug output
            # Log Placement Index Calculation quietly
            pass
        else:
            metrics["placement_index"] = None
            print(f"⚠️ Placement Index: total_eligible is 0, cannot calculate")
        
        return metrics
        
    except Exception as e:
        print(f"❌ Error calculating NBA metrics: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

@app.route('/api/reports/generate', methods=['POST', 'OPTIONS'])
@login_required
def generate_report():
    """Generate a static NBA-ready report snapshot."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    data = request.get_json() or {}
    batch_id = data.get('batch_id')
    
    if not batch_id:
        return jsonify({"message": "Batch ID is required"}), 400
    
    try:
        # Validate batch_id and ownership
        try:
            college = session.get('college')
            branch = session.get('branch')
            if not college or not branch:
                missing = []
                if not college: missing.append("college")
                if not branch: missing.append("branch")
                return jsonify({"message": f"User profile incomplete: missing {', '.join(missing)}. Please logout and login again."}), 400

            batch_obj_id = ObjectId(batch_id)
            batch = db.batches.find_one({"_id": batch_obj_id, "college": college, "branch": branch})
            if not batch:
                return jsonify({"message": "Invalid batch ID or unauthorized access"}), 400
            batch_name = f"{batch.get('branch', 'Unknown')} – {batch.get('batch_range', 'Unknown')}"
        except Exception:
            return jsonify({"message": "Invalid batch ID format"}), 400
        
        # Capture current system data
        report_id = ObjectId()
        
        report_data = {
            "_id": report_id,
            "batch_id": batch_obj_id,
            "college": college, # Store college for isolation
            "branch": branch,  # Store branch for isolation
            "batch_name": batch_name,
            "generated_at": datetime.now(timezone.utc),
            "report_id": str(report_id),
            "status": "generated"
        }
        
        # Fetch intake data (batch-specific)
        summary_id = f"summary_{batch_id}"
        intake_doc = db.dashboard_summary.find_one({"_id": summary_id}) or {}
        kpis = intake_doc.get("kpis", {})
        # N1 = Total Admitted (from dashboard), N2 = DSE, Total = N1 + N2
        n1 = kpis.get("total_admitted", 0) or 0
        n2 = kpis.get("dse", 0) or 0
        report_data["intake"] = {
            "sanctioned_intake": kpis.get("sanctioned_intake", 0),
            "n1": n1,
            "n2": n2,
            "n3": 0,  # Additional intake if any
            "total_admitted": n1 + n2
        }
        
        # Fetch academic performance data
        semesters_data = intake_doc.get("semesters", [])
        report_data["academic"] = []
        for sem_data in semesters_data:
            report_data["academic"].append({
                "semester": sem_data.get("semester", ""),
                "avg_sgpa": sem_data.get("avg_sgpa", 0),
                "avg_cgpa": sem_data.get("avg_cgpa", 0),
                "pass_percentage": sem_data.get("pass_percentage", 0),
                "total_students": sem_data.get("total", 0)
            })
        
        # Fetch extracurricular data - handle both old (category) and new (output_type) structures
        extracurricular_records = list(db.extracurricular_records.find({"batch_id": batch_obj_id}))
        
        report_data["extracurricular"] = {
            "sports": [],
            "technical": [],
            "cultural": [],
            "internships": [],
            "courses": [],
            "industrial_visits": []
        }
        for record in extracurricular_records:
            # Handle both old structure (category) and new structure (output_type)
            output_type = record.get("output_type", "").strip()
            category = record.get("category", "").lower()
            
            # Get participant count from either structure
            participants = record.get("number_of_participants") or record.get("number_of_students") or record.get("count_participants", 0)
            wins = record.get("number_of_wins") or record.get("count_won", 0)
            
            # Determine category based on output_type or category
            # Normalize output_type for backward compatibility
            normalized_output_type = normalize_output_type(output_type)
            
            if normalized_output_type == "Sports" or category == "sports":
                report_data["extracurricular"]["sports"].append({
                    "event_date": record.get("event_date"),
                    "level": record.get("level", ""),
                    "sports_name": record.get("sports_name") or record.get("sport_name") or record.get("type", ""),
                    "participants": participants,
                    "wins": wins,
                    "student_names": record.get("student_names", []),
                    "evidence_file_id": record.get("file_id"),
                    "evidence_link": record.get("evidence_link")
                })
            elif normalized_output_type == "Technical" or category == "technical":
                report_data["extracurricular"]["technical"].append({
                    "event_date": record.get("event_date"),
                    "level": record.get("level", ""),
                    "event_type": record.get("event_type") or record.get("type", ""),
                    "event_name": record.get("event_name") or record.get("type", ""),
                    "organizer": record.get("organizer") or record.get("organization") or record.get("organizer_name") or "",
                    "achievement": record.get("achievement") or record.get("outcome") or "Participated",
                    "participants": participants,
                    "student_names": record.get("student_names", []),
                    "evidence_file_id": record.get("file_id"),
                    "evidence_link": record.get("evidence_link")
                })
            elif normalized_output_type == "Cultural" or category == "cultural" or category == "student activities":
                # Use custom category if event_category is "Other" and custom category is provided
                event_category = record.get("event_category") or record.get("category", "")
                if event_category == "Other":
                    custom_cat = record.get("cultural_custom_category", "").strip()
                    display_category = custom_cat if custom_cat else "Other"
                else:
                    display_category = event_category
                
                report_data["extracurricular"]["cultural"].append({
                    "event_date": record.get("event_date"),
                    "event_category": display_category,
                    "event_name": record.get("event_name", ""),
                    "organizer": record.get("organizer", ""),
                    "participants": record.get("number_of_participants", 0),
                    "student_names": record.get("student_names", []),
                    "evidence_file_id": record.get("file_id"),
                    "evidence_link": record.get("evidence_link")
                })
            elif normalized_output_type == "Courses" or category == "courses" or category == "certifications":
                report_data["extracurricular"]["courses"].append({
                    "event_date": record.get("event_date"),
                    "course_name": record.get("course_name") or record.get("type", ""),
                    "platform": record.get("platform", ""),
                    "mode": record.get("mode", ""),
                    "duration": record.get("duration", ""),
                    "participants": participants,
                    "student_names": record.get("student_names", []),
                    "evidence_file_id": record.get("file_id"),
                    "evidence_link": record.get("evidence_link")
                })
            elif normalized_output_type == "Industrial Visit":
                report_data["extracurricular"]["industrial_visits"].append({
                    "academic_year": record.get("academic_year", ""),
                    "details": record.get("industrial_visit_details", ""),
                    "date": record.get("date", "") or record.get("iv_date", ""),
                    "total_students": record.get("number_of_students", participants),
                    "number_of_faculty": record.get("number_of_faculty") or record.get("faculty", ""),
                    "evidence_file_id": record.get("file_id"),
                    "evidence_link": record.get("evidence_link")
                })
            # Check for internships - both output_type and level (separate if to handle records that might have both)
            if normalized_output_type == "Internship" or record.get("level", "").lower() == "internship":
                report_data["extracurricular"]["internships"].append({
                    "event_date": record.get("event_date"),
                    "organization_name": record.get("organization_name") or record.get("type", ""),
                    "internship_domain": record.get("internship_domain", ""),
                    "participants": participants,
                    "has_stipend": record.get("has_stipend", "No"),
                    "stipend_amount": record.get("stipend_amount", "0"),
                    "student_names": record.get("student_names", []),
                    "evidence_file_id": record.get("file_id"),
                    "evidence_link": record.get("evidence_link")
                })
        
        # Fetch placement data
        placement_outcomes = list(db.placement_outcomes.find({"batch_id": batch_obj_id}))
        report_data["placement"] = {
            "placed": 0,
            "higher_studies_india": 0,
            "higher_studies_abroad": 0,
            "entrepreneurship": 0,
            "company_wise": {},
            "country_wise": {}
        }
        
        # Initialize detailed lists for DOCX generation
        report_data["career_outcomes"] = {
            "placement": [],
            "higher_studies": [],
            "entrepreneurship": []
        }
        
        for outcome in placement_outcomes:
            outcome_type = outcome.get("outcome_type", "").lower()
            num_students = outcome.get("number_of_students", 0)
            
            # Common fields
            outcome_data = {
                "year": outcome.get("year"),
                "student_names": outcome.get("student_names", []),
                "number_of_students": num_students,
                "evidence_file_id": outcome.get("evidence_file_id"),
                "evidence_link": outcome.get("evidence_link")
            }
            
            if outcome_type == "placed":
                report_data["placement"]["placed"] += num_students
                company = outcome.get("company_name")
                if company:
                    report_data["placement"]["company_wise"][company] = \
                        report_data["placement"]["company_wise"].get(company, 0) + num_students
                
                # Add to detailed list
                outcome_data.update({
                    "company_name": outcome.get("company_name", ""),
                    "job_role": outcome.get("job_role", ""),
                    "ctc_package": outcome.get("ctc_package"),
                    "placement_type": outcome.get("placement_type", ""),
                    "location": outcome.get("location", "")
                })
                report_data["career_outcomes"]["placement"].append(outcome_data)
                
            elif outcome_type == "higher_studies":
                country = outcome.get("country", "").lower()
                if country == "india":
                    report_data["placement"]["higher_studies_india"] += num_students
                else:
                    report_data["placement"]["higher_studies_abroad"] += num_students
                    country_name = outcome.get("country_name") or country
                    report_data["placement"]["country_wise"][country_name] = \
                        report_data["placement"]["country_wise"].get(country_name, 0) + num_students
                
                # Add to detailed list
                outcome_data.update({
                    "course_name": outcome.get("course_name", ""),
                    "degree_type": outcome.get("degree_type", ""),
                    "university_institute": outcome.get("university_institute", ""),
                    "country": outcome.get("country", ""),
                    "country_name": outcome.get("country_name", "")
                })
                report_data["career_outcomes"]["higher_studies"].append(outcome_data)
                
            elif outcome_type == "entrepreneurship":
                report_data["placement"]["entrepreneurship"] += num_students
                
                # Add to detailed list
                outcome_data.update({
                    "startup_name": outcome.get("startup_name", ""),
                    "startup_status": outcome.get("startup_status", ""),
                    "year_started": outcome.get("year_started", "")
                })
                report_data["career_outcomes"]["entrepreneurship"].append(outcome_data)
        
        # Calculate NBA Metrics for this batch
        nba_metrics = calculate_nba_metrics(batch_obj_id, college=college, branch=branch)
        if nba_metrics:
            report_data["nba_metrics"] = nba_metrics
        else:
            report_data["nba_metrics"] = {
                "enrolment_ratio": None,
                "success_index_without_backlogs": None,
                "success_index_stipulated_period": None,
                "academic_performance_index": None,
                "placement_index": None
            }
        
        # Store report snapshot in database
        db.reports.insert_one(report_data)
        
        return jsonify({
            "message": "Report generated successfully",
            "report_id": str(report_id),
            "generated_at": report_data["generated_at"].isoformat()
        }), 201
        
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"❌ Error generating report: {str(e)}")
        print(f"Traceback: {error_trace}")
        return jsonify({"message": f"Error generating report: {str(e)}"}), 500

@app.route('/api/reports', methods=['GET', 'OPTIONS'])
@login_required
def list_reports():
    """List all generated reports."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    try:
        college = session.get('college')
        branch = session.get('branch')
        if not college or not branch:
            missing = []
            if not college: missing.append("college")
            if not branch: missing.append("branch")
            return jsonify({"message": f"User profile incomplete: missing {', '.join(missing)}. Please logout and login again."}), 400

        # Run cleanup for this user's context
        cleanup_pending_deletes(college, branch)

        # Find batches for this branch
        branch_batches = list(db.batches.find({"college": college, "branch": branch}, {"_id": 1}))
        branch_batch_ids = [b["_id"] for b in branch_batches]
        
        # Query reports that either have the college+branch field explicitly or belong to one of the branch's batches
        reports_query = {
            "$and": [
                {
                    "$or": [
                        {"college": college, "branch": branch},
                        {"batch_id": {"$in": branch_batch_ids}}
                    ]
                },
                {"deleted": {"$ne": True}}
            ]
        }
        reports = list(db.reports.find(reports_query).sort("generated_at", -1))

        reports_list = []
        for report in reports:
            reports_list.append({
                "report_id": str(report.get("_id")),
                "batch_name": report.get("batch_name", "Unknown Batch"),
                "generated_at": report.get("generated_at").isoformat() if report.get("generated_at") else None,
                "status": report.get("status", "generated")
            })
        reports_list = convert_objectid_to_str(reports_list)
        return jsonify(reports_list), 200
    except Exception as e:
        return jsonify({"message": f"Error listing reports: {str(e)}"}), 500

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    A function that places a hyperlink within a paragraph object.
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def _add_evidence_to_cell(cell, rec, endpoint_name):
    # 1. External Links (Drive, OneDrive, etc.)
    bulk_link = rec.get("evidence_link") or rec.get("proof_link_or_file") or rec.get("report_or_photos")
    
    if bulk_link:
        bulk_link_str = str(bulk_link).strip()
        # Check if it looks like a URL
        if bulk_link_str.lower().startswith(('http://', 'https://')):
            # Add as clickable hyperlink
            # Use existing paragraph if cell has text but new line, or create new one
            if cell.text:
                p = cell.add_paragraph()
            else:
                p = cell.paragraphs[0]
            
            add_hyperlink(p, bulk_link_str, "View External Link")
        else:
            # Not a URL, display as text
            if cell.text:
                cell.add_paragraph(bulk_link_str)
            else:
                cell.text = bulk_link_str
    
    # 2. Internal Files (Do not expose backend URLs)
    evidence_ids = rec.get("evidence_file_ids")
    if not evidence_ids and rec.get("evidence_file_id"):
        evidence_ids = [rec.get("evidence_file_id")]
    
    if evidence_ids:
        msg = "Uploaded internally (not shareable)"
        if cell.text:
            cell.add_paragraph(msg)
        else:
            cell.text = msg
def generate_docx_report(report, college_name):
    """Generate a DOCX report from the report data."""
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
    except ImportError:
        pass
        
    doc = Document()
    
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Helvetica'
    font.size = Pt(11)
    
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Helvetica'
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Helvetica'
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(24)
    h2.paragraph_format.space_after = Pt(12)
    
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Helvetica'
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(16)
    h3.paragraph_format.space_after = Pt(8)

    # --- Cover Page ---
    doc.add_paragraph("\n\n")
    doc.add_paragraph("REPORT ON", style='Heading 1')
    
    branch_name = "Unknown Branch"
    if report.get('batch_name'):
        parts = report.get('batch_name').split('–')
        if parts:
            branch_name = parts[0].strip().upper()
            
    doc.add_paragraph(branch_name, style='Heading 1')
    doc.add_paragraph("\n")
    
    p = doc.add_paragraph()
    p.add_run("Batch: ").bold = True
    p.add_run(f"{report.get('batch_name', '')}")
    
    p = doc.add_paragraph()
    p.add_run("Institution: ").bold = True
    p.add_run(f"{college_name}")
    
    # Date Formatting
    generated_time = report.get('generated_at')
    formatted_time = 'N/A'
    if generated_time:
        if isinstance(generated_time, datetime):
             if generated_time.tzinfo is None:
                generated_time = generated_time.replace(tzinfo=timezone.utc)
             ist_time = generated_time.astimezone(timezone(timedelta(hours=5, minutes=30)))
             formatted_time = ist_time.strftime('%B %d, %Y at %I:%M %p IST')
        elif isinstance(generated_time, str):
             formatted_time = generated_time
             
    p = doc.add_paragraph()
    p.add_run("Generated on: ").bold = True
    p.add_run(f"{formatted_time}")
    
    doc.add_page_break()

    # --- Section 1: Intake Data ---
    doc.add_paragraph("Section 1 – Intake Data", style='Heading 2')
    
    intake = report.get("intake", {})
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    # Bold headers
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                
    metrics = [
        ('Sanctioned Intake (N)', str(intake.get("sanctioned_intake", 0))),
        ('N1 (Regular Admission)', str(intake.get("n1", 0))),
        ('N2 (DSE)', str(intake.get("n2", 0))),
        ('N3 (Additional)', str(intake.get("n3", 0))),
        ('Total Admitted (N1 + N2 + N3)', str(intake.get("total_admitted", 0)))
    ]
    
    for metric, value in metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value
        
    doc.add_page_break()

    # --- Section 2: Academic Performance ---
    doc.add_paragraph("Section 2 – Academic Performance", style='Heading 2')
    
    academic = report.get("academic", [])
    filtered_academic = [sem for sem in academic if sem.get("semester", "").lower() in ['sem1', 'sem2', 'sem3', 'sem4', 'sem5', 'sem6', 'sem7', 'sem8']]
    filtered_academic.sort(key=lambda x: int(x.get("semester", "").lower().replace('sem', '')))
    
    if filtered_academic:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        headers = ['Semester', 'Avg SGPA', 'Avg CGPA', 'Pass %', 'Total Students']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    
        for sem in filtered_academic:
            row_cells = table.add_row().cells
            
            avg_sgpa = float(sem.get('avg_sgpa') or 0)
            avg_cgpa = float(sem.get('avg_cgpa') or 0)
            total = float(sem.get("total_students") or sem.get("total") or 0)
            without = float(sem.get("without_backlog") or sem.get("success_without_backlog") or 0)
            pass_pct = round((without / total) * 100, 2) if total > 0 else 0.0
            
            row_cells[0].text = sem.get("semester", "").upper()
            row_cells[1].text = f"{avg_sgpa:.2f}"
            row_cells[2].text = f"{avg_cgpa:.2f}"
            row_cells[3].text = f"{pass_pct:.2f}%"
            row_cells[4].text = str(int(total))
            
        # Charts
        labels = [s.get("semester", "").upper() for s in filtered_academic]
        cgpa_data = [float(s.get("avg_cgpa") or 0) for s in filtered_academic]
        pass_pct_data = []
        for s in filtered_academic:
            total = float(s.get("total_students") or s.get("total") or 0)
            without = float(s.get("without_backlog") or s.get("success_without_backlog") or 0)
            pass_pct_data.append(round((without / total) * 100, 2) if total > 0 else 0.0)

        # CGPA Chart
        fig, ax = plt.subplots(figsize=(6, 4))
        ax.plot(labels, cgpa_data, marker='o', color='#2c3e50')
        ax.set_title('Average CGPA Trend')
        ax.set_ylabel('CGPA')
        ax.grid(True, alpha=0.3)
        
        chart_buffer = BytesIO()
        plt.savefig(chart_buffer, format='png', dpi=100, bbox_inches='tight')
        plt.close()
        chart_buffer.seek(0)
        doc.add_picture(chart_buffer, width=Inches(6))
        doc.add_paragraph("")
        
        # Pass % Chart
        fig, ax = plt.subplots(figsize=(6, 4))
        ax.bar(labels, pass_pct_data, color='#2c3e50', alpha=0.8)
        ax.set_title('Pass Percentage Trend')
        ax.set_ylabel('Pass Percentage (%)')
        ax.set_ylim(0, 100)
        
        chart_buffer = BytesIO()
        plt.savefig(chart_buffer, format='png', dpi=100, bbox_inches='tight')
        plt.close()
        chart_buffer.seek(0)
        doc.add_picture(chart_buffer, width=Inches(6))
        
    doc.add_page_break()

    # --- Section 3: Extracurricular ---
    doc.add_paragraph("Section 3 – Extracurricular Activities", style='Heading 2')
    extracurricular = report.get("extracurricular", {})
    
    # FIX: Fetch Cultural records directly from DB to bypass data loss in preprocessing
    batch_id = report.get("batch_id")
    if batch_id:
        # Ensure batch_id is ObjectId if needed, but it should be passed correctly
        raw_cultural = list(db.extracurricular_records.find({
            "batch_id": batch_id,
            "output_type": "Cultural"
        }))
        if raw_cultural:
            if "cultural" not in extracurricular:
                extracurricular["cultural"] = []
            extracurricular["cultural"] = raw_cultural

    # Helper to format student names
    def fmt_names(r):
        names = r.get("student_names", [])
        if isinstance(names, list):
            valid_names = [str(n) for n in names if n]
            return ", ".join(valid_names) if valid_names else "—"
        
        # Backward compatibility: Parse string to array then join for consistent formatting
        val_str = str(names).strip() if names else ""
        if val_str:
            name_list = [n.strip() for n in val_str.split(',') if n.strip()]
            return ", ".join(name_list) if name_list else "—"
            
        return "—"

    # Helper functions for row mapping
    def map_sports_row(r):
        return [
             r.get("sports_name") or r.get("sport_name") or r.get("type", ""), 
             r.get("level", ""), 
             str(r.get("number_of_participants") or r.get("participants") or r.get("number_of_students") or ""), 
             str(r.get("number_of_wins") or r.get("wins") or "0"),
             fmt_names(r)
        ]

    def map_cultural_row(r):
        row_data = [
             r.get("event_category", "") if r.get("event_category") != "Other" else r.get("cultural_custom_category", "Other"), 
             r.get("event_name", ""), 
             r.get("organizer", ""),
             str(r.get("number_of_participants", "")),
             fmt_names(r)
        ]
        
        # Validate column count (6 headers - 1 report column = 5 data columns)
        expected_len = 5
        if len(row_data) != expected_len:
             # Log mismatch quietly or handle error if critical
             pass
             
        return row_data

    def map_technical_row(r):
        return [
             r.get("event_type") or r.get("type", ""), 
             r.get("event_name") or r.get("type", ""), 
             r.get("organizer") or r.get("organization") or r.get("organizer_name") or "",
             r.get("level", ""), 
             str(r.get("number_of_participants") or r.get("participants") or r.get("number_of_students") or ""),
             fmt_names(r),
             r.get("achievement") or r.get("outcome") or "Participated"
        ]

    def map_internship_row(r):
        return [
             r.get("internship_title") or r.get("internship_domain") or r.get("type", ""), 
             r.get("organization_name") or r.get("type", ""), 
             r.get("mode", ""), 
             str(r.get("number_of_students") or r.get("participants") or ""),
             fmt_names(r),
             r.get("duration", ""),
             "Yes" if str(r.get("has_stipend", "")).lower() == 'yes' else "No",
             str(r.get("stipend_amount", "0"))
        ]

    def map_course_row(r):
        return [
             r.get("course_name") or r.get("type", ""), 
             r.get("mode", ""), 
             r.get("platform", ""), 
             str(r.get("number_of_students") or r.get("participants") or ""),
             fmt_names(r),
             r.get("duration", "")
        ]

    categories_map = [
        ("Sports Activities", "sports", ["Sport Name", "Level", "No. of Participants", "No. of Wins", "Name of Students", "Report"], 
         map_sports_row),
        ("Cultural Activities", "cultural", ["Event Category", "Event Name", "Organizer", "No. of Participants", "Name of Students", "Report"],
         map_cultural_row),
        ("Technical Activities", "technical", ["Event Type", "Event Name", "Organizer", "Level", "No. of Participants", "Name of Students", "Achievement", "Report"],
         map_technical_row),
        ("Internships", "internships", ["Internship Title", "Company Name", "Mode", "No. of Students", "Name of Students", "Duration", "Stipend (Yes / No)", "Stipend Amount", "Report"],
         map_internship_row),
        ("Courses & Certifications", "courses", ["Course Name", "Mode", "Platform", "No. of Students", "Name of Students", "Duration", "Report"],
         map_course_row),
    ]
    
    for title, key, headers, row_mapper in categories_map:
        records = extracurricular.get(key, [])
        if records:
            doc.add_paragraph(title, style='Heading 3')
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            
            for rec in records:
                row_cells = table.add_row().cells
                values = row_mapper(rec)
                for i, val in enumerate(values):
                    row_cells[i].text = str(val)
                _add_evidence_to_cell(row_cells[-1], rec, 'download_extracurricular_file')
            doc.add_paragraph("")

    # Industrial Visits
    ivs = extracurricular.get("industrial_visits", [])
    if ivs:
        doc.add_paragraph("Industrial Visits", style='Heading 3')
        iv_by_year = {}
        for iv in ivs:
            year = iv.get("academic_year", "Unknown")
            iv_by_year.setdefault(year, []).append(iv)
            
        for year in sorted(iv_by_year.keys()):
            doc.add_paragraph(f"Academic Year: {year}")
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            headers = ['Sr. No', 'Details', 'Date', 'Total Students', 'Faculty', 'Evidence']
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                
            for idx, rec in enumerate(iv_by_year[year], start=1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = str(rec.get("details", ""))
                row_cells[2].text = str(rec.get("date", ""))
                row_cells[3].text = str(rec.get("total_students", ""))
                row_cells[4].text = str(rec.get("number_of_faculty") or rec.get("faculty", ""))
                _add_evidence_to_cell(row_cells[5], rec, 'download_extracurricular_file')
            doc.add_paragraph("")
            
    # Participation Summary Table (College vs Outside)
    extracurricular_summary = extracurricular.get("summary", {})
    if extracurricular_summary:
        doc.add_paragraph("Participation Summary (College vs Outside)", style='Heading 3')
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Category'
        hdr_cells[1].text = 'Total Participants'
        
        # Bold headers
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        summary_data = [
            ("College-Level Participation", str(extracurricular_summary.get("college_level_participants", 0))),
            ("Outside-College Participation", str(extracurricular_summary.get("outside_college_participants", 0)))
        ]
        
        for cat, count in summary_data:
            row_cells = table.add_row().cells
            row_cells[0].text = cat
            row_cells[1].text = count
            
        doc.add_paragraph("")

    # Extracurricular Charts
    doc.add_paragraph("Extracurricular Visualizations", style='Heading 3')
    categories = ['Sports', 'Technical', 'Cultural', 'Internships', 'Courses']
    participation = [
        sum(int(r.get("participants") or 0) for r in extracurricular.get("sports", [])),
        sum(int(r.get("participants") or 0) for r in extracurricular.get("technical", [])),
        sum(int(r.get("participants") or 0) for r in extracurricular.get("cultural", [])),
        sum(int(r.get("participants") or 0) for r in extracurricular.get("internships", [])),
        sum(int(r.get("participants") or 0) for r in extracurricular.get("courses", []))
    ]
    if sum(participation) > 0:
        fig, ax = plt.subplots(figsize=(8, 5))
        ax.bar(categories, participation, color='#34495e', alpha=0.8)
        ax.set_title('Student Participation by Activity Category')
        ax.set_ylabel('Participants')
        chart_buffer = BytesIO()
        plt.savefig(chart_buffer, format='png', dpi=100, bbox_inches='tight')
        plt.close()
        chart_buffer.seek(0)
        doc.add_picture(chart_buffer, width=Inches(6))
    doc.add_page_break()
    
    # Section 4: Career Outcomes
    doc.add_paragraph("Section 4 – Career Outcomes", style='Heading 2')
    
    # Get placement summary data (used for fallback and summary table)
    placement = report.get("placement", {})

    # Try to get pre-processed career outcomes (new structure)
    career_outcomes = report.get("career_outcomes", {})
    placements = career_outcomes.get("placement", [])
    higher_studies = career_outcomes.get("higher_studies", [])
    entrepreneurship = career_outcomes.get("entrepreneurship", [])
    
    # Fallback to old structure if new structure is empty (e.g. legacy reports or direct DB access)
    if not placements and not higher_studies and not entrepreneurship:
        all_outcomes = placement.get("all_outcomes", [])
        
        for outcome in all_outcomes:
            o_type = str(outcome.get("outcome_type", "")).strip().lower()
            if o_type == "placed":
                placements.append(outcome)
            elif o_type in ["higher studies", "higher_studies"]:
                higher_studies.append(outcome)
            elif o_type == "entrepreneurship":
                entrepreneurship.append(outcome)

    # 7. Placement Table
    if placements:
        doc.add_paragraph("Placement Details", style='Heading 3')
        # Group by Year
        year_groups = {}
        for p in placements:
            year = str(p.get("year") or "Unknown")
            year_groups.setdefault(year, []).append(p)
            
        for year in sorted(year_groups.keys()):
            doc.add_paragraph(f"Year: {year}")
            # Columns: No. of Students, Year, Company Name, Job Role, Name of Students, CTC Package, Placement Type, Location, Report
            headers = ["No. of Students", "Year", "Company Name", "Job Role", "Name of Students", "CTC Package", "Placement Type", "Location", "Report"]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                
            for rec in year_groups[year]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(rec.get("number_of_students", ""))
                row_cells[1].text = str(rec.get("year", ""))
                row_cells[2].text = str(rec.get("company_name", ""))
                row_cells[3].text = str(rec.get("job_role") or rec.get("role", ""))
                row_cells[4].text = fmt_names(rec)
                row_cells[5].text = str(rec.get("ctc_package") or rec.get("ctc") or "")
                row_cells[6].text = str(rec.get("placement_type", ""))
                row_cells[7].text = str(rec.get("location", ""))
                _add_evidence_to_cell(row_cells[8], rec, 'download_placement_file')
        doc.add_paragraph("")

    # 8. Higher Studies Table
    if higher_studies:
        doc.add_paragraph("Higher Studies Details", style='Heading 3')
        # Group by Year
        year_groups = {}
        for h in higher_studies:
            year = str(h.get("year") or "Unknown")
            year_groups.setdefault(year, []).append(h)
            
        for year in sorted(year_groups.keys()):
            doc.add_paragraph(f"Year: {year}")
            # Columns: No. of Students, Year, Name of Students, Course Name, Degree Type, University Name, Country (India / Abroad), If Abroad – Country Name, Report
            headers = ["No. of Students", "Year", "Name of Students", "Course Name", "Degree Type", "University Name", "Country (India / Abroad)", "If Abroad – Country Name", "Report"]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                
            for rec in year_groups[year]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(rec.get("number_of_students", ""))
                row_cells[1].text = str(rec.get("year", ""))
                row_cells[2].text = fmt_names(rec)
                row_cells[3].text = str(rec.get("course_name") or "")
                row_cells[4].text = str(rec.get("degree_type") or rec.get("degree") or "")
                row_cells[5].text = str(rec.get("university_institute") or "")
                row_cells[6].text = str(rec.get("country") or "")
                row_cells[7].text = str(rec.get("country_name") or "")
                _add_evidence_to_cell(row_cells[8], rec, 'download_placement_file')
        doc.add_paragraph("")

    # 9. Entrepreneurship Table
    if entrepreneurship:
        doc.add_paragraph("Entrepreneurship Details", style='Heading 3')
        # Group by Year
        year_groups = {}
        for e in entrepreneurship:
            year = str(e.get("year") or "Unknown")
            year_groups.setdefault(year, []).append(e)
            
        for year in sorted(year_groups.keys()):
            doc.add_paragraph(f"Year: {year}")
            # Columns: No. of Students, Year, Name of Students, Startup Name, Status, Year Started, Report
            headers = ["No. of Students", "Year", "Name of Students", "Startup Name", "Status", "Year Started", "Report"]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                
            for rec in year_groups[year]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(rec.get("number_of_students", ""))
                row_cells[1].text = str(rec.get("year", ""))
                row_cells[2].text = fmt_names(rec)
                row_cells[3].text = str(rec.get("startup_name", ""))
                row_cells[4].text = str(rec.get("startup_status") or rec.get("entrepreneurship_stage") or rec.get("status", ""))
                row_cells[5].text = str(rec.get("year_started") or rec.get("year", ""))
                _add_evidence_to_cell(row_cells[6], rec, 'download_placement_file')
        doc.add_paragraph("")

    # Summary Tables
    doc.add_paragraph("Career Outcome Summary", style='Heading 3')
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    rows = [
        ['Outcome', 'Count'],
        ['Placed', str(placement.get("placed", 0))],
        ['Higher Studies', str(placement.get("higher_studies_india", 0) + placement.get("higher_studies_abroad", 0))],
        ['Entrepreneurship', str(placement.get("entrepreneurship", 0))]
    ]
    for i, row_data in enumerate(rows):
        row_cells = table.rows[i].cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        if i == 0:
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[1].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph("")
    
    # Visualizations
    doc.add_paragraph("Placement Visualizations", style='Heading 3')
    labels = ['Placed', 'Higher Studies', 'Entrepreneurship']
    sizes = [placement.get("placed", 0), 
             placement.get("higher_studies_india", 0) + placement.get("higher_studies_abroad", 0), 
             placement.get("entrepreneurship", 0)]
    
    if sum(sizes) > 0:
        fig, ax = plt.subplots(figsize=(6, 6))
        ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90, colors=['#34495e', '#7f8c8d', '#95a5a6'])
        ax.set_title('Placement vs Higher Studies vs Entrepreneurship')
        chart_buffer = BytesIO()
        plt.savefig(chart_buffer, format='png', dpi=100, bbox_inches='tight')
        plt.close()
        chart_buffer.seek(0)
        doc.add_picture(chart_buffer, width=Inches(5))
        
    doc.add_page_break()
    
    # --- Section 5: NBA Accreditation Metrics ---
    doc.add_paragraph("Section 5 – NBA Accreditation Metrics", style='Heading 2')
    
    nba_metrics = report.get("nba_metrics", {})
    
    # Table 1: Academic & Progression Metrics
    doc.add_paragraph("Academic & Progression Metrics", style='Heading 3')
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    # Bold headers
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                
    progression_metrics = [
        ('Enrolment Ratio', f"{nba_metrics.get('enrolment_ratio', 'N/A')}%" if nba_metrics.get('enrolment_ratio') is not None else "N/A"),
        ('Success Index (Without Backlogs)', str(nba_metrics.get('success_index_without_backlogs', 'N/A'))),
        ('Success Index (Stipulated Period)', str(nba_metrics.get('success_index_stipulated_period', 'N/A'))),
        ('Academic Performance Index (API)', str(nba_metrics.get('academic_performance_index', 'N/A'))),
        ('Second Year API', str(nba_metrics.get('second_year_api', 'N/A'))),
        ('Third Year API', str(nba_metrics.get('third_year_api', 'N/A'))),
        ('Placement Index', str(nba_metrics.get('placement_index', 'N/A')))
    ]
    
    for metric, value in progression_metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value

    doc.add_paragraph("")

    # Table 2: Placement & Career Outcomes
    doc.add_paragraph("Placement & Career Outcomes", style='Heading 3')
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Count'
    
    # Bold headers
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    placement_data = report.get("placement", {})
    career_metrics = [
        ('Placed', str(placement_data.get("placed", 0))),
        ('Higher Studies', str(placement_data.get("higher_studies_india", 0) + placement_data.get("higher_studies_abroad", 0))),
        ('Entrepreneurship', str(placement_data.get("entrepreneurship", 0)))
    ]
    
    for category, count in career_metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = count

    doc.add_paragraph("")

    # Table 3: Extracurricular Participation
    doc.add_paragraph("Extracurricular Participation Summary", style='Heading 3')
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Total Records'
    
    # Bold headers
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    extracurricular = report.get("extracurricular", {})
    extra_metrics = [
        ('Sports', str(len(extracurricular.get("sports", [])))),
        ('Technical', str(len(extracurricular.get("technical", [])))),
        ('Cultural', str(len(extracurricular.get("cultural", [])))),
        ('Internships', str(len(extracurricular.get("internships", [])))),
        ('Courses', str(len(extracurricular.get("courses", [])))),
        ('Industrial Visits', str(len(extracurricular.get("industrial_visits", []))))
    ]
    
    for category, count in extra_metrics:
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = count

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

@app.route('/api/reports/<report_id>/download', methods=['GET', 'OPTIONS'])
def download_report(report_id):
    """Download a report as PDF. Fetches fresh data each time."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    try:
        # Check if report exists (for metadata)
        report_meta = db.reports.find_one({"_id": ObjectId(report_id)})
        if not report_meta:
            return jsonify({"message": "Report not found"}), 404
        
        # Extract batch_id from report metadata - MANDATORY for all queries
        batch_obj_id = report_meta.get("batch_id")
        if not batch_obj_id:
            return jsonify({"message": "Report missing batch_id. Cannot generate PDF without batch filter."}), 400
        
        # Validate batch exists and belongs to branch
        college = session.get('college')
        branch = session.get('branch')
        if not college or not branch:
             return jsonify({"message": "User college or branch not found"}), 400

        batch = db.batches.find_one({"_id": batch_obj_id, "college": college, "branch": branch})
        if not batch:
            return jsonify({"message": "Batch associated with report no longer exists or unauthorized access"}), 400
        
        batch_id = str(batch_obj_id)
        batch_name = report_meta.get("batch_name", f"{batch.get('branch', 'Unknown')} – {batch.get('batch_range', 'Unknown')}")
        
        # Build fresh report data
        report = {
            "batch_name": batch_name,
            "batch_id": batch_obj_id, # Added for direct DB access in DOCX generation
            "generated_at": report_meta.get("generated_at", datetime.now(timezone.utc))
        }
        
        # Fetch fresh intake data - MUST use batch-specific summary
        summary_id = f"summary_{batch_id}"
        intake_doc = db.dashboard_summary.find_one({"_id": summary_id}) or {}
        kpis = intake_doc.get("kpis", {})
        # N1 = Total Admitted (from dashboard), N2 = DSE, Total = N1 + N2
        n1 = kpis.get("total_admitted", 0) or 0
        n2 = kpis.get("dse", 0) or 0
        report["intake"] = {
            "sanctioned_intake": kpis.get("sanctioned_intake", 0),
            "n1": n1,
            "n2": n2,
            "n3": 0,
            "total_admitted": n1 + n2
        }
        
        # Fetch fresh academic performance data - from batch-specific summary
        semesters_data = intake_doc.get("semesters", [])
        report["academic"] = []
        for sem_data in semesters_data:
            report["academic"].append({
                "semester": sem_data.get("semester", ""),
                "avg_sgpa": sem_data.get("avg_sgpa", 0),
                "avg_cgpa": sem_data.get("avg_cgpa", 0),
                "pass_percentage": sem_data.get("pass_percentage", 0),
                "total_students": sem_data.get("total", 0),
                "total": sem_data.get("total", 0),  # Also include as 'total' for compatibility
                "without_backlog": sem_data.get("without_backlog", 0),
                "success_without_backlog": sem_data.get("success_without_backlog", 0)
            })
        
        # Fetch fresh extracurricular data - MUST filter by batch_id
        extracurricular_records = list(db.extracurricular_records.find({"batch_id": batch_obj_id}))
        # Remove debug logs for PDF Generation
        pass
        report["extracurricular"] = {
            "sports": [],
            "technical": [],
            "cultural": [],
            "internships": [],
            "courses": [],
            "industrial_visits": [],
            "all_records": extracurricular_records  # Keep full records for detailed processing
        }
        
        # Initialize participation summary
        total_college_level = 0
        total_outside_college = 0
        current_college = str(session.get('college') or "").lower().strip()

        def _norm_names(val):
            if isinstance(val, str):
                return [n.strip() for n in val.split(",") if n and str(n).strip()]
            if isinstance(val, list):
                return [str(n).strip() for n in val if str(n).strip()]
            return []
        for record in extracurricular_records:
            output_type = record.get("output_type", "").strip()
            category = record.get("category", "").strip().lower()
            participants = record.get("number_of_participants") or record.get("number_of_students") or record.get("count_participants", 0)
            wins = record.get("number_of_wins") or record.get("count_won", 0)
            
            # Common fields to preserve for ALL categories (Evidence & Student Names)
            common_fields = {
                "student_names": _norm_names(record.get("student_names", [])),
                "evidence_file_ids": record.get("evidence_file_ids"),
                "evidence_file_id": record.get("evidence_file_id"),
                "evidence_link": record.get("evidence_link"),
                "proof_link_or_file": record.get("proof_link_or_file"),
                "report_or_photos": record.get("report_or_photos"),
                "event_date": record.get("event_date")
            }
            
            # Normalize output_type for backward compatibility
            normalized_output_type = normalize_output_type(output_type)
            
            # Strict Classification Logic
            is_sports = (
                normalized_output_type == "Sports" or 
                "sport" in category or 
                "game" in category or 
                "tournament" in category or
                "championship" in category
            )
            is_internship = (
                normalized_output_type == "Internship" or 
                "internship" in category or 
                record.get("level", "").lower() == "internship"
            )
            is_iv = (
                normalized_output_type == "Industrial Visit" or 
                "industrial visit" in category or 
                "iv" in category or
                "visit" in category
            )
            is_course = (
                normalized_output_type == "Courses" or 
                "course" in category or 
                "certification" in category or 
                "nptel" in category or 
                "coursera" in category or
                "udemy" in category
            )
            is_tech = (
                normalized_output_type == "Technical" or 
                "technical" in category or 
                "hackathon" in category or 
                "coding" in category or
                "competition" in category or
                "paper presentation" in category or
                "project" in category
            )
            
            # --- Participation Location Logic (Standardized) ---
            # Categories: Internship, Sports, Technical, Cultural
            is_loc_outside = False
            is_loc_college = False
            should_count_loc = False
            count_val = participants # Default to participants
            
            # Implied Cultural (everything else that isn't excluded)
            is_cultural_implied = not (is_sports or is_internship or is_iv or is_course or is_tech)
            
            if is_internship:
                should_count_loc = True
                # Rule 1: ALL internships are considered OUTSIDE COLLEGE
                is_loc_outside = True
                
            elif is_sports:
                should_count_loc = True
                # Rule 2: Sports Level Check
                lvl = str(record.get("level", "")).lower()
                if "inter-department" in lvl or "department" in lvl:
                    is_loc_college = True
                elif any(x in lvl for x in ["inter-college", "university", "state", "national", "international", "zonal"]):
                    is_loc_outside = True
                else:
                    is_loc_college = True # Default to college if ambiguous

            elif is_tech:
                should_count_loc = True
                # Rule 3: Technical -> Organizer Check (Count PARTICIPANTS)
                # count_val remains participants (default)
                org_name = str(record.get("organization_name", "")).strip().lower()
                organizer = str(record.get("organizer", "")).strip().lower()
                check_str = org_name + " " + organizer
                
                if (current_college and current_college in check_str) or "department" in check_str:
                    is_loc_college = True
                else:
                    is_loc_outside = True

            elif is_cultural_implied:
                should_count_loc = True
                # Rule 4: Cultural -> Organizer Check
                org_name = str(record.get("organization_name", "")).strip().lower()
                organizer = str(record.get("organizer", "")).strip().lower()
                check_str = org_name + " " + organizer
                
                if (current_college and current_college in check_str) or "department" in check_str:
                    is_loc_college = True
                else:
                    is_loc_outside = True
            
            if should_count_loc:
                if is_loc_outside:
                    total_outside_college += count_val
                elif is_loc_college:
                    total_college_level += count_val

            # Priority Handling (Strict if/elif chain to prevent duplicates)
            if is_sports:
                report["extracurricular"]["sports"].append({
                    "sport_name": record.get("sports_name") or record.get("sport_name") or record.get("event_name") or record.get("activity_name") or record.get("type", ""),
                    "level": record.get("level", ""),
                    "participants": participants,
                    "wins": wins,
                    **common_fields
                })
            elif is_internship:
                report["extracurricular"]["internships"].append({
                    "internship_title": record.get("internship_title") or record.get("internship_domain") or record.get("role") or record.get("position") or record.get("type", ""),
                    "internship_type": record.get("organization_name") or record.get("company_name") or record.get("type", ""),
                    "participants": participants,
                    "has_stipend": record.get("has_stipend", False),
                    "stipend_amount": record.get("stipend_amount", 0),
                    "organization_name": record.get("organization_name", ""),
                    "internship_domain": record.get("internship_domain", ""),
                    "mode": record.get("mode", ""),
                    "duration": record.get("duration", ""),
                    **common_fields
                })
            elif is_iv:
                report["extracurricular"]["industrial_visits"].append({
                    "academic_year": record.get("academic_year", ""),
                    "details": record.get("industrial_visit_details") or record.get("details") or record.get("place_of_visit", ""),
                    "date": record.get("date", ""),
                    "total_students": record.get("number_of_students", participants),
                    "faculty": record.get("faculty") or record.get("number_of_faculty", ""),
                    **common_fields
                })
            elif is_course:
                report["extracurricular"]["courses"].append({
                    "course_name": record.get("course_name") or record.get("skill_area") or record.get("event_name") or record.get("type", ""),
                    "platform": record.get("platform") or record.get("organizer", ""),
                    "mode": record.get("mode", ""),
                    "participants": participants,
                    "duration": record.get("duration", ""),
                    **common_fields
                })
            elif is_tech:
                report["extracurricular"]["technical"].append({
                    "event_type": record.get("event_type") or record.get("type", ""),
                    "event_name": record.get("event_name") or record.get("project_title") or record.get("type", ""),
                    "level": record.get("level", ""),
                    "achievement": record.get("achievement") or record.get("outcome", ""),
                    "participants": participants,
                    "organizer": record.get("organizer", ""),
                    **common_fields
                })
            else:
                # Fallback to Cultural/Other to ensure NO DATA IS DROPPED
                # This catches "Cultural", "Student Activities", and any uncategorized records
                event_category = record.get("event_category", "")
                
                # Check for "Other" custom category
                if event_category == "Other":
                    custom_cat = record.get("cultural_custom_category", "").strip()
                    display_category = custom_cat if custom_cat else "Other"
                else:
                    display_category = event_category or category or "General Event"
                
                # Report Requirement: Event Category must show exact user-entered value
                # (Do NOT merge "Student Activities" into "Cultural" for the report table)
                
                report["extracurricular"]["cultural"].append({
                    "event_category": display_category,
                    "event_name": record.get("event_name") or record.get("activity_name") or record.get("title", ""),
                    "participants": participants,
                    "achievement": record.get("achievement") or record.get("outcome", ""),
                    **common_fields
                })
        
        # Store Participation Summary
        report["extracurricular"]["summary"] = {
            "college_level_participants": total_college_level,
            "outside_college_participants": total_outside_college
        }
        
        # Fetch fresh placement data - MUST filter by batch_id
        placement_outcomes = list(db.placement_outcomes.find({"batch_id": batch_obj_id}))
        # Remove debug logs for PDF Generation
        pass
        
        report["placement"] = {
            "placed": 0,
            "higher_studies_india": 0,
            "higher_studies_abroad": 0,
            "entrepreneurship": 0,
            "company_wise": {},
            "sector_wise": {},
            "country_wise": {},
            "all_outcomes": placement_outcomes  # Keep full records for backward compatibility
        }
        
        # Initialize detailed lists for DOCX generation (New Structure)
        report["career_outcomes"] = {
            "placement": [],
            "higher_studies": [],
            "entrepreneurship": []
        }
        
        for outcome in placement_outcomes:
            outcome_type = str(outcome.get("outcome_type", "")).strip()
            outcome_type_lower = outcome_type.lower()
            num_students = outcome.get("number_of_students", 0)
            
            # Common fields
            outcome_data = {
                "year": outcome.get("year"),
                "student_names": _norm_names(outcome.get("student_names", [])),
                "number_of_students": num_students,
                "evidence_file_id": outcome.get("evidence_file_id"),
                "evidence_link": outcome.get("evidence_link")
            }
            
            if outcome_type_lower == "placed":
                report["placement"]["placed"] += num_students
                company = outcome.get("company_name")
                sector = outcome.get("sector") or outcome.get("company_type", "")
                if company:
                    report["placement"]["company_wise"][company] = \
                        report["placement"]["company_wise"].get(company, 0) + num_students
                if sector:
                    report["placement"]["sector_wise"][sector] = \
                        report["placement"]["sector_wise"].get(sector, 0) + num_students
                
                # Add to detailed list
                outcome_data.update({
                    "company_name": outcome.get("company_name", ""),
                    "job_role": outcome.get("job_role", ""),
                    "ctc_package": outcome.get("ctc_package"),
                    "placement_type": outcome.get("placement_type", ""),
                    "location": outcome.get("location", "")
                })
                report["career_outcomes"]["placement"].append(outcome_data)

            elif outcome_type_lower in ["higher studies", "higher_studies"]:
                country = str(outcome.get("country", "")).strip().lower()
                degree = outcome.get("degree_type") or outcome.get("degree") or outcome.get("course_name") or outcome.get("course", "")
                country_name = outcome.get("country_name") or (country.title() if country else "Unknown")
                
                if country == "india":
                    report["placement"]["higher_studies_india"] += num_students
                else:
                    report["placement"]["higher_studies_abroad"] += num_students
                    if country_name and country_name != "Unknown":
                        report["placement"]["country_wise"][country_name] = \
                            report["placement"]["country_wise"].get(country_name, 0) + num_students
                
                # Add to detailed list
                outcome_data.update({
                    "course_name": outcome.get("course_name", ""),
                    "degree_type": outcome.get("degree_type", ""),
                    "university_institute": outcome.get("university_institute", ""),
                    "country": outcome.get("country", ""),
                    "country_name": outcome.get("country_name", "")
                })
                report["career_outcomes"]["higher_studies"].append(outcome_data)
                
            elif outcome_type_lower == "entrepreneurship":
                report["placement"]["entrepreneurship"] += num_students
                
                # Add to detailed list
                outcome_data.update({
                    "startup_name": outcome.get("startup_name", ""),
                    "startup_status": outcome.get("startup_status", ""),
                    "year_started": outcome.get("year_started", "")
                })
                report["career_outcomes"]["entrepreneurship"].append(outcome_data)
        
        # Calculate NBA Metrics for this batch (fresh calculation)
        nba_metrics = calculate_nba_metrics(batch_obj_id, college=college, branch=branch)
        if nba_metrics:
            report["nba_metrics"] = nba_metrics
        else:
            report["nba_metrics"] = {
                "enrolment_ratio": None,
                "success_index_without_backlogs": None,
                "success_index_stipulated_period": None,
                "academic_performance_index": None,
                "placement_index": None
            }
        
        # Generate DOCX Report
        try:
            # generate_docx_report returns a BytesIO object with the file content
            buffer = generate_docx_report(report, college_name=session.get('college', 'Institute'))
            
            # Create safe filename
            safe_batch = "".join(c for c in batch_name if c.isalnum() or c in (' ', '-', '_')).strip()
            filename = f"NBA_Report_{safe_batch}_{datetime.now().strftime('%Y%m%d')}.docx"
            
            return send_file(
                buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            print(f"❌ Error generating DOCX report: {e}")
            print(f"Error details: {error_details}")
            return jsonify({
                "message": f"Error generating report: {str(e)}"
            }), 500


    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"❌ Error downloading report: {str(e)}")
        print(f"Traceback: {error_trace}")
        return jsonify({"message": f"Error generating report: {str(e)}"}), 500

@app.route('/api/reports/<report_id>', methods=['DELETE', 'OPTIONS'])
def delete_report(report_id):
    """Delete a report."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    try:
        # Fetch report first to check permissions/existence
        report = db.reports.find_one({"_id": ObjectId(report_id)})
        if not report:
             return jsonify({"message": "Report not found"}), 404

        # Verify ownership via session if available
        college = session.get('college')
        branch = session.get('branch')
        
        if college and branch:
            # If logged in, enforce ownership
            if report.get('college') != college or report.get('branch') != branch:
                 # Check if it belongs to a batch in this branch
                 batch_id = report.get('batch_id')
                 batch = db.batches.find_one({"_id": batch_id, "college": college, "branch": branch})
                 if not batch:
                     return jsonify({"message": "Unauthorized access to this report"}), 403

        finalize = (request.args.get('finalize') or '').lower() == 'true'
        if not finalize:
            # Soft delete - Pending Delete State
            db.reports.update_one(
                {"_id": ObjectId(report_id)},
                {"$set": {
                    "pending_delete": True,
                    "delete_requested_at": datetime.now(timezone.utc)
                }}
            )
            return jsonify({"message": "Report marked for deletion", "pending_delete": True}), 200

        # Permanent delete
        result = db.reports.delete_one({"_id": ObjectId(report_id)})
        if result.deleted_count == 0:
            return jsonify({"message": "Report not found"}), 404
        return jsonify({"message": "Report deleted successfully"}), 200
    except Exception as e:
        return jsonify({"message": f"Error deleting report: {str(e)}"}), 500

@app.route('/api/reports/<report_id>/restore', methods=['POST', 'OPTIONS'])
def restore_report(report_id):
    """Restore a soft-deleted report."""
    if request.method == 'OPTIONS':
        return jsonify(status='ok'), 200
    
    try:
        # Fetch report first to check permissions/existence
        report = db.reports.find_one({"_id": ObjectId(report_id)})
        if not report:
             return jsonify({"message": "Report not found"}), 404

        # Verify ownership via session if available
        college = session.get('college')
        branch = session.get('branch')
        
        if college and branch:
            # If logged in, enforce ownership
            if report.get('college') != college or report.get('branch') != branch:
                 batch_id = report.get('batch_id')
                 batch = db.batches.find_one({"_id": batch_id, "college": college, "branch": branch})
                 if not batch:
                     return jsonify({"message": "Unauthorized access to this report"}), 403

        db.reports.update_one(
            {"_id": ObjectId(report_id)},
            {"$unset": {"pending_delete": "", "delete_requested_at": "", "deleted": ""}}
        )
        return jsonify({"message": "Report restored successfully"}), 200
    except Exception as e:
        return jsonify({"message": f"Error restoring report: {str(e)}"}), 500

# --- Admin Panel Logic ---

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return jsonify({"message": "Authentication required"}), 401
        
        user = db.users.find_one({"_id": ObjectId(session['user_id'])})
        if not user or user.get('role') != 'super_admin':
            return jsonify({"message": "Access denied: Super Admin only"}), 403
            
        return f(*args, **kwargs)
    return decorated_function

@app.route('/admin', methods=['GET'])
def admin_panel():
    if 'user_id' not in session:
        return jsonify({"message": "Please login first"}), 401
    
    user = db.users.find_one({"_id": ObjectId(session['user_id'])})
    if not user or user.get('role') != 'super_admin':
        return "Access Denied: Super Admin role required.", 403
        
    return send_file('admin.html')

@app.route('/api/admin/colleges', methods=['GET', 'POST'])
@admin_required
def manage_colleges():
    if request.method == 'GET':
        # --- SELF-HEALING: Ensure Gharda exists if collection is empty ---
        if db.colleges.count_documents({}) == 0:
            # print("⚠️ Colleges collection empty. Seeding default college via API...")
            db.colleges.insert_one({
                "name": "Gharda Institute of Technology",
                "status": "active",
                "subscription_active": True,
                "subscription_expiry": None,
                "max_users_per_branch": 2,
                "created_at": datetime.now(timezone.utc)
            })

        colleges = list(db.colleges.find({}))
        # Format for frontend
        results = []
        for c in colleges:
            c['_id'] = str(c['_id'])
            results.append(c)
        
        stats = {
            "total_colleges": db.colleges.count_documents({}),
            "active_colleges": db.colleges.count_documents({"status": "active"}),
            "total_users": db.users.count_documents({
                "role": {"$ne": "super_admin"}
            })
        }
        return jsonify({"colleges": results, "stats": stats})

    if request.method == 'POST':
        data = request.get_json()
        name = data.get('name')
        if not name:
            return jsonify({"message": "College name is required"}), 400
            
        if db.colleges.find_one({"name": name}):
            return jsonify({"message": "College already exists"}), 400
            
        new_college = {
            "name": name,
            "status": data.get('status', 'active'),
            "subscription_active": data.get('subscription_active', True),
            "max_users_per_branch": int(data.get('max_users_per_branch', 2)),
            "allowed_domains": data.get('allowed_domains', []),
            "subscription_expiry": None, # Default unlimited
            "created_at": datetime.now(timezone.utc)
        }
        
        if data.get('subscription_expiry'):
             # Parse ISO string
             try:
                 expiry = datetime.fromisoformat(data['subscription_expiry'].replace('Z', '+00:00'))
                 new_college['subscription_expiry'] = expiry
             except:
                 pass

        db.colleges.insert_one(new_college)
        return jsonify({"message": "College created successfully"}), 201

@app.route('/api/admin/colleges/<college_id>', methods=['PUT'])
@admin_required
def update_college(college_id):
    data = request.get_json()
    
    update_fields = {}
    if 'name' in data: update_fields['name'] = data['name']
    if 'status' in data: update_fields['status'] = data['status']
    if 'subscription_active' in data: update_fields['subscription_active'] = data['subscription_active']
    if 'max_users_per_branch' in data: update_fields['max_users_per_branch'] = int(data['max_users_per_branch'])
    if 'allowed_domains' in data: update_fields['allowed_domains'] = data['allowed_domains']
    
    if 'subscription_expiry' in data:
        if data['subscription_expiry']:
            try:
                update_fields['subscription_expiry'] = datetime.fromisoformat(data['subscription_expiry'].replace('Z', '+00:00'))
            except:
                pass
        else:
            update_fields['subscription_expiry'] = None

    db.colleges.update_one(
        {"_id": ObjectId(college_id)},
        {"$set": update_fields}
    )
    return jsonify({"message": "College updated"}), 200


# --- Feedback Module (Beta) ---

@app.route('/api/feedback', methods=['POST'])
@login_required
def submit_feedback():
    """Submit feedback from faculty users."""
    try:
        data = request.json
        message = data.get('message', '').strip()
        feedback_category = data.get('feedback_category')

        if not message or len(message) < 10:
            return jsonify({"message": "Message must be at least 10 characters long."}), 400
        
        if not feedback_category:
            return jsonify({"message": "Feedback category is required."}), 400

        feedback_entry = {
            "college": session.get('college'),
            "branch": session.get('branch'),
            "user_email": session.get('email'), 
            "user_name": session.get('user_name', 'Faculty User'), 
            "message": message,
            "feedback_category": feedback_category,
            "created_at": datetime.now(timezone(timedelta(hours=5, minutes=30))) # IST
        }

        # Check if user_email is available in session, otherwise fetch from DB
        # We also need to fetch full_name if possible since it's not in session
        if 'email' not in session or not session.get('user_name') or not session.get('college'):
            try:
                user = db.users.find_one({"_id": ObjectId(session['user_id'])})
                if user:
                    if not feedback_entry['user_email']:
                        feedback_entry['user_email'] = user.get('email')
                    if not feedback_entry['college']:
                        feedback_entry['college'] = user.get('college')
                    if not feedback_entry['branch']:
                        feedback_entry['branch'] = user.get('branch')
                    feedback_entry['user_name'] = user.get('full_name') or 'Faculty User'
            except Exception as db_err:
                print(f"Error fetching user details for feedback: {db_err}")

        db.feedback.insert_one(feedback_entry)
        return jsonify({"success": True, "message": "Feedback submitted successfully."}), 201

    except Exception as e:
        print(f"Error submitting feedback: {e}")
        return jsonify({"message": "Internal server error"}), 500

@app.route('/api/admin/feedback', methods=['GET'])
@login_required
def get_all_feedback():
    """Get all feedback (Super Admin only)."""
    try:
        # Check for super admin role
        user_id = session.get('user_id')
        user = db.users.find_one({"_id": ObjectId(user_id)})
        
        if not user or user.get('role') != 'super_admin':
            return jsonify({"message": "Unauthorized access"}), 403

        feedback_list = list(db.feedback.find().sort("created_at", -1))
        
        # Enrich feedback with college info if missing
        for feedback in feedback_list:
            if not feedback.get('college') or not feedback.get('branch'):
                try:
                    # Try to find user by email to get college details
                    if feedback.get('user_email'):
                        user_info = db.users.find_one({"email": feedback['user_email']})
                        if user_info:
                            if not feedback.get('college'):
                                feedback['college'] = user_info.get('college') or 'Unknown'
                            if not feedback.get('branch'):
                                feedback['branch'] = user_info.get('branch') or ''
                except Exception as inner_e:
                    print(f"Error enriching feedback: {inner_e}")

        # Convert ObjectId to string for JSON serialization
        return jsonify(convert_objectid_to_str(feedback_list)), 200

    except Exception as e:
        print(f"Error fetching feedback: {e}")
        return jsonify({"message": "Internal server error"}), 500

@app.route('/api/admin/feedback/<feedback_id>', methods=['DELETE'])
@login_required
def delete_feedback(feedback_id):
    """Delete feedback (Super Admin only)."""
    try:
        # Check for super admin role
        user_id = session.get('user_id')
        user = db.users.find_one({"_id": ObjectId(user_id)})
        
        if not user or user.get('role') != 'super_admin':
            return jsonify({"message": "Unauthorized access"}), 403

        result = db.feedback.delete_one({"_id": ObjectId(feedback_id)})
        
        if result.deleted_count == 0:
            return jsonify({"message": "Feedback not found"}), 404

        return jsonify({"success": True, "message": "Feedback deleted successfully"}), 200

    except Exception as e:
        print(f"Error deleting feedback: {e}")
        return jsonify({"message": "Internal server error"}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', '5001'))
    app.run(host='0.0.0.0', debug=False, port=port, use_reloader=False)
